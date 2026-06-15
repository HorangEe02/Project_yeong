"""초안 생성 라우터.

v3.0: 감사 로깅 추가 — 누가 어떤 문서를 생성/내보내기 했는지 추적
Day 8 Phase 1: 5 신규 엔드포인트 (doc-types / cc-recommend / quality-score / diff / stream-v2)
"""

import json
import logging
import os
from typing import Literal

from fastapi import APIRouter, Depends, File, HTTPException, Request, UploadFile
from fastapi.responses import Response

from backend.schemas.draft import (
    CCGroup,
    CCRecRequest,
    CCRecResponse,
    DiagnoseCheck,
    DiagnoseResponse,
    DiffLine,
    DiffRequest,
    DiffResponse,
    DiffStats,
    DocTypeListResponse,
    DocTypeMeta,
    DraftExportRequest,
    DraftGenerateRequest,
    DraftResponse,
    DraftReviseRequest,
    DraftStreamRequest,
    DraftStreamV2Request,
    AttachmentRecommendResponse,
    AttachmentSuggestionItem,
    MailSendRequest,
    MailSendResponse,
    PartialEditRequest,
    PartialEditResponse,
    PartialEditScanRequest,
    PartialEditScanResponse,
    PartialEditSection,
    QualityRequest,
    QualityResponse,
    QualityScoresDetail,
    UploadReferenceResponse,
    VersionDetailResponse,
    VersionListItem,
    VersionListResponse,
    VersionReviewRequest,
    VersionSaveRequest,
    VersionSaveResponse,
)
from backend.dependencies import get_current_user, require_role_level
from backend.sse import create_sse_response, sse_from_sync_generator

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/draft", tags=["draft"])


def _attach_watermark_marker(file_bytes: bytes, fmt: str, watermark_id: str) -> bytes:
    """Attach a traceable AI marker for formats without a document footer.

    Args:
        file_bytes: Exported document bytes from the format shaper.
        fmt: Lower-case export format.
        watermark_id: Stable ID computed from the source content.

    Returns:
        bytes: Export bytes with the marker attached where the format supports it.

    Raises:
        ImportError: If the XLSX marker path cannot import openpyxl.
    """
    if fmt not in {"txt", "csv", "xlsx"}:
        return file_bytes

    from features.draft.watermark import watermark_text

    marker = watermark_text(watermark_id)
    if fmt in {"txt", "csv"}:
        suffix = f"\n\n{marker}\n".encode("utf-8")
        return file_bytes + suffix

    from io import BytesIO
    from openpyxl import load_workbook

    src = BytesIO(file_bytes)
    wb = load_workbook(src)
    wb.properties.keywords = "AJIN_AI_ASSISTED"
    wb.properties.description = marker
    out = BytesIO()
    wb.save(out)
    return out.getvalue()


@router.post("/generate")
async def generate_draft_stream(req: DraftGenerateRequest, user=Depends(get_current_user)):
    """초안을 SSE 스트리밍으로 생성한다."""
    from core.llm_client import auto_select_model, stream_generate

    model = req.model or auto_select_model("draft")

    tone_map = {"공식적": "격식체", "친근한": "존댓말 친근체", "긴급": "긴급 요청 톤"}
    tone_str = tone_map.get(req.tone, req.tone)

    # v2.0: context에 따른 프롬프트 분기
    context_hint = ""
    if req.context == "internal":
        context_hint = "\n[맥락] 사내 문서 — 간결하게, 존댓말 친근체로 작성하세요."
    else:
        context_hint = f"\n[수신처] {req.recipient}\n[언어] {req.language}"

    prompt = f"""당신은 아진산업의 업무 문서 작성 전문가입니다.
아래 요청에 따라 문서 초안을 작성하세요.

[문서유형] {req.doc_type or '자동 분류'}
[어조] {tone_str}{context_hint}
[요청] {req.user_input}

한국어로 작성하세요. 제목, 수신/발신, 본문을 포함하세요."""

    # v3.0: 감사 로깅
    if user:
        from backend.auth_middleware import log_api_access
        log_api_access(
            endpoint="/api/draft/generate",
            method="POST",
            detail=f"doc_type={req.doc_type}, tone={req.tone}",
            user=user,
        )

    return create_sse_response(
        sse_from_sync_generator(
            stream_generate,
            prompt=prompt,
            model=model,
            feature="draft",
        )
    )


@router.post("/generate-pipeline", response_model=DraftResponse)
async def generate_draft_pipeline(req: DraftGenerateRequest, user=Depends(get_current_user)):
    """DraftPipeline을 사용한 전체 초안 생성 (분류→생성→렌더링)."""
    try:
        import asyncio
        from features.draft import DraftPipeline

        pipeline = DraftPipeline()
        rendered, session = asyncio.run(pipeline.create_draft(req.user_input))

        return DraftResponse(
            session_id=session.session_id,
            doc_type=session.request.doc_type.value,
            content=rendered,
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/export")
async def export_draft(req: DraftExportRequest, user=Depends(get_current_user)):
    """초안을 지정 형식으로 내보낸다.

    Day 8 Phase 3: 7포맷 graceful 지원
    - txt:  utf-8
    - docx: python-docx
    - pdf:  reportlab/fpdf
    - xlsx: openpyxl (markdown 표 자동 파싱)
    - csv:  utf-8 BOM (Excel 한글)
    - odt:  HWPX OPF 패키지 활용 (ODT 호환 포맷)
    - hwpx: HwpxExporter (HWPX 정식)
    """
    try:
        # v3.0: 감사 로깅
        if user:
            from backend.auth_middleware import log_api_access
            log_api_access(
                endpoint="/api/draft/export",
                method="POST",
                detail=f"format={req.format}, doc_type={req.doc_type}",
                user=user,
            )

        content = req.content
        fmt = req.format.lower()

        # ── Plan v1.0 §4.2 — format_shaper 단일 진입점 ──
        from features.draft.format_shaper import shape_for_format

        # MIME / 확장자 / 파일명 매핑
        media_map = {
            "txt": ("text/plain", "draft.txt"),
            "docx": (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "draft.docx",
            ),
            "pdf": ("application/pdf", "draft.pdf"),
            "csv": ("text/csv; charset=utf-8", "draft.csv"),
            "xlsx": (
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                "draft.xlsx",
            ),
            "hwpx": ("application/vnd.hancom.hwpx", "draft.hwpx"),
            "odt": ("application/vnd.oasis.opendocument.text", "draft.odt"),
        }
        if fmt not in media_map:
            raise HTTPException(status_code=400, detail=f"지원하지 않는 형식: {fmt}")

        media_type, filename = media_map[fmt]

        try:
            file_bytes = shape_for_format(content, req.doc_type or "general", fmt)  # type: ignore[arg-type]
        except ImportError as e:
            raise HTTPException(status_code=503, detail=f"{fmt} 변환 라이브러리 미설치: {e}")
        except Exception as e:
            logger.warning("[draft] %s shaper 실패: %s", fmt, e)
            # CSV 는 BOM 평문 fallback (Excel 한글 호환)
            if fmt == "csv":
                file_bytes = "﻿".encode("utf-8") + content.encode("utf-8")
            else:
                raise HTTPException(status_code=503, detail=f"{fmt.upper()} 변환 실패: {e}")

        from features.draft.watermark import compute_watermark_id

        watermark_id = compute_watermark_id(content)
        try:
            file_bytes = _attach_watermark_marker(file_bytes, fmt, watermark_id)
        except ImportError as e:
            raise HTTPException(status_code=503, detail=f"{fmt} 워터마크 메타데이터 라이브러리 미설치: {e}")
        except Exception as e:
            logger.warning("[draft] %s watermark marker 실패: %s", fmt, e)
            raise HTTPException(status_code=503, detail=f"{fmt.upper()} 워터마크 부착 실패: {e}")

        headers = {
            "Content-Disposition": f"attachment; filename={filename}",
            "X-AJIN-Watermark-Id": watermark_id,
            "X-AJIN-AI-Assisted": "true",
        }
        if fmt == "odt":
            headers["X-AJIN-Fallback"] = "docx-as-odt"

        return Response(content=file_bytes, media_type=media_type, headers=headers)
    except HTTPException:
        raise
    except ImportError as e:
        raise HTTPException(status_code=500, detail=f"내보내기 모듈 없음: {e}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ──────────────────────────────────────────────────────────────────────────
# v4.7 Sprint 2 P0 (축 ③) — 템플릿 prefill (메일 초안 자동 채우기)
# ──────────────────────────────────────────────────────────────────────────


from pydantic import BaseModel as _BaseModel


class TemplatePrefillRecipient(_BaseModel):
    email: str = ""
    name: str = ""


class TemplatePrefillResponse(_BaseModel):
    subject: str = ""
    body: str = ""
    to: list[TemplatePrefillRecipient] = []
    cc: list[TemplatePrefillRecipient] = []


@router.get("/templates/{template_id}/prefill", response_model=TemplatePrefillResponse)
def get_template_prefill(template_id: str, context_id: str):
    """템플릿 ID + 컨텍스트 ID 로 메일 초안 placeholder 를 채워 반환.

    지원 템플릿: 'spc_violation' (context_id = spc_violations.db row id).
    """
    from features.draft.prefill_engine import prefill_template

    payload = prefill_template(template_id, context_id)
    if payload is None:
        raise HTTPException(status_code=404, detail="prefill_unavailable")
    return TemplatePrefillResponse(
        subject=payload.get("subject", ""),
        body=payload.get("body", ""),
        to=[TemplatePrefillRecipient(**r) for r in payload.get("to", [])],
        cc=[TemplatePrefillRecipient(**r) for r in payload.get("cc", [])],
    )


@router.get("/templates")
async def list_templates():
    """v1.6: 사용 가능한 문서 템플릿 목록을 반환한다."""
    from pathlib import Path
    from config import KNOWLEDGE_BASE_DIR

    templates_dir = KNOWLEDGE_BASE_DIR / "templates"
    templates = []
    if templates_dir.exists():
        for j2 in sorted(templates_dir.glob("*.j2")):
            name = j2.stem
            # 파일명에서 언어/유형 추론
            lang = "en" if "_en_" in name else "ko"
            category = "email" if "email" in name else "quality_doc"
            templates.append({
                "id": name,
                "name": name.replace("_", " ").title(),
                "language": lang,
                "category": category,
                "filename": j2.name,
            })
    return {"templates": templates, "count": len(templates)}


# ═══════════════════════════════════════════════════════════
# Day 8 Phase 1 — 5 신규 엔드포인트
# ═══════════════════════════════════════════════════════════


@router.get("/doc-types", response_model=DocTypeListResponse)
async def get_doc_types(user=Depends(get_current_user)):
    """13 문서 유형 메타 — internal 7 + external 6 (canonical Draft.jsx 일치).

    features/draft/doc_type_config.py 의 INTERNAL_DOC_TYPES + EXTERNAL_DOC_TYPES
    를 frontend 친화 포맷으로 변환.
    """
    try:
        from features.draft.doc_type_config import get_doc_types as _get_doc_types

        items: list[DocTypeMeta] = []

        for category in ("internal", "external"):
            type_dict = _get_doc_types(category)
            for ko_name, config in type_dict.items():
                # 영문 매핑은 향후 doc_type_config 에 추가될 수 있도록 fallback
                en_name = config.get("name_en", ko_name)
                # required_fields 추출 (config 에 fields 또는 required_fields 키)
                fields = config.get("required_fields") or config.get("fields") or []
                if isinstance(fields, dict):
                    fields = list(fields.keys())

                # id: 영문 식별자 (없으면 ko 이름의 slug)
                doc_id = config.get("id") or _slugify(ko_name)

                items.append(
                    DocTypeMeta(
                        id=doc_id,
                        category=category,  # type: ignore
                        name_ko=ko_name,
                        name_en=en_name if isinstance(en_name, str) else ko_name,
                        required_fields=[str(f) for f in fields],
                    )
                )

        # B1 — 템플릿 헤더 주석에서 메타데이터 추출 후 합성
        items = _enrich_with_template_metadata(items)

        internal_count = sum(1 for i in items if i.category == "internal")
        external_count = sum(1 for i in items if i.category == "external")

        return DocTypeListResponse(
            items=items,
            internal_count=internal_count,
            external_count=external_count,
        )
    except Exception as e:
        logger.warning("[draft] doc-types 로드 실패: %s — fallback 16종 반환", e)
        return _fallback_doc_types()


def _slugify(name: str) -> str:
    import re
    s = re.sub(r"[^\w\s-]", "", name).strip().lower()
    return re.sub(r"[\s_]+", "_", s) or "unknown"


def _enrich_with_template_metadata(items: list[DocTypeMeta]) -> list[DocTypeMeta]:
    """B1 — 각 doc_type 의 .j2 템플릿 헤더 주석에서 usage_hint·dept_recommend·var_metadata
    를 읽어와 DocTypeMeta 에 합성. 템플릿 없거나 메타 없으면 빈 값 유지.
    """
    from features.draft.template_metadata import get_template_path, parse_template_metadata
    from config import KNOWLEDGE_BASE_DIR

    templates_dir = KNOWLEDGE_BASE_DIR / "templates"
    enriched: list[DocTypeMeta] = []
    for it in items:
        path = get_template_path(it.id, templates_dir)
        if path is None:
            enriched.append(it)
            continue
        meta = parse_template_metadata(path)
        # var_metadata 는 list[dict] → list[VarMetadata] 변환은 Pydantic 이 자동 수행
        enriched.append(
            it.model_copy(
                update={
                    "usage_hint": meta.get("usage_hint", "") or it.usage_hint,
                    "dept_recommend": meta.get("dept_recommend", []) or it.dept_recommend,
                    "example_output": meta.get("example_output", "") or it.example_output,
                    "var_metadata": meta.get("var_metadata", []) or it.var_metadata,
                }
            )
        )
    return enriched


def _fallback_doc_types() -> DocTypeListResponse:
    """doc_type_config 미동작 시 — canonical Draft.jsx 16종 fallback."""
    items = [
        # 외부 6종
        DocTypeMeta(id="8d_report",   category="external", name_ko="8D Report",       name_en="8D Report",       required_fields=["title", "issue", "team"]),
        DocTypeMeta(id="ecn",         category="external", name_ko="ECN",             name_en="ECN",             required_fields=["title", "change_reason"]),
        DocTypeMeta(id="ppap",        category="external", name_ko="PPAP",            name_en="PPAP",            required_fields=["part_number", "level"]),
        DocTypeMeta(id="fmea",        category="external", name_ko="FMEA",            name_en="FMEA",            required_fields=["process", "risk"]),
        DocTypeMeta(id="msa",         category="external", name_ko="MSA",             name_en="MSA",             required_fields=["instrument", "study_type"]),
        DocTypeMeta(id="oem_email",   category="external", name_ko="OEM 영문 이메일", name_en="OEM Email",       required_fields=["recipient", "subject"]),
        # 내부 10종
        DocTypeMeta(id="internal_email",       category="internal", name_ko="사내 이메일",   name_en="Internal Email",         required_fields=["recipient", "subject"]),
        DocTypeMeta(id="meeting_min",          category="internal", name_ko="회의록",       name_en="Meeting Minutes",        required_fields=["date", "attendees"]),
        DocTypeMeta(id="weekly_report",        category="internal", name_ko="주간 보고",     name_en="Weekly Report",          required_fields=["week", "summary"]),
        DocTypeMeta(id="leave_request",        category="internal", name_ko="휴가 신청서",   name_en="Leave Request",          required_fields=["start_date", "reason"]),
        DocTypeMeta(id="business_trip_request",category="internal", name_ko="출장 신청서",   name_en="Business Trip Request",  required_fields=["destination", "purpose"]),
        DocTypeMeta(id="resignation_letter",   category="internal", name_ko="사직서",       name_en="Resignation Letter",      required_fields=["last_work_date", "reason"]),
        DocTypeMeta(id="personnel_notice",     category="internal", name_ko="인사발령 통지", name_en="Personnel Notice",       required_fields=["target_employee", "effective_date"]),
        DocTypeMeta(id="quote",                category="internal", name_ko="견적서",       name_en="Quote",                  required_fields=["customer", "items"]),
        DocTypeMeta(id="travel_report",        category="internal", name_ko="출장 보고서",   name_en="Travel Report",          required_fields=["destination", "purpose"]),
        DocTypeMeta(id="spc_report",           category="internal", name_ko="SPC Report",   name_en="SPC Report",             required_fields=["process", "period"]),
    ]
    items = _enrich_with_template_metadata(items)
    internal_count = sum(1 for i in items if i.category == "internal")
    external_count = sum(1 for i in items if i.category == "external")
    return DocTypeListResponse(items=items, internal_count=internal_count, external_count=external_count)


@router.post("/cc/recommend", response_model=CCRecResponse)
async def recommend_cc(req: CCRecRequest, user=Depends(get_current_user)):
    """CC 자동 추천 — 4-tier (필수/권장/자주 함께 보낸/선택). B5 v4.0 빈도 학습 포함.

    features/draft/cc_recommender.recommend_cc_with_history(doc_type, sender_dept, sender_div)
    → {"mandatory": [...], "recommended": [...], "frequent": [...], "optional": [...]}
    """
    try:
        from features.draft.cc_recommender import recommend_cc_with_history as _recommend

        # 사용자가 있으면 부서 우선, 없으면 요청 페이로드 부서
        sender_dept = req.sender_department or (getattr(user, "department", "") if user else "")
        sender_div = req.sender_division

        cc_data = _recommend(
            doc_type=req.doc_type,
            sender_department=sender_dept,
            sender_division=sender_div,
        )

        groups = [
            CCGroup(
                tier="required",
                label_ko="필수",
                label_en="REQUIRED",
                departments=cc_data.get("mandatory", []),
            ),
            CCGroup(
                tier="recommended",
                label_ko="권장",
                label_en="RECOMMENDED",
                departments=cc_data.get("recommended", []),
            ),
            CCGroup(
                tier="frequent",
                label_ko="자주 함께 보낸",
                label_en="FREQUENT",
                departments=cc_data.get("frequent", []),
            ),
            CCGroup(
                tier="optional",
                label_ko="선택",
                label_en="OPTIONAL",
                departments=cc_data.get("optional", []),
            ),
        ]

        return CCRecResponse(
            groups=groups,
            doc_type=req.doc_type,
            sender_department=sender_dept,
        )
    except Exception as e:
        logger.warning("[draft] cc-recommend 실패: %s", e)
        # fallback: 빈 그룹
        return CCRecResponse(
            groups=[
                CCGroup(tier="required", label_ko="필수", label_en="REQUIRED", departments=[]),
                CCGroup(tier="recommended", label_ko="권장", label_en="RECOMMENDED", departments=[]),
                CCGroup(tier="optional", label_ko="선택", label_en="OPTIONAL", departments=[]),
            ],
            doc_type=req.doc_type,
            sender_department=req.sender_department,
        )


@router.post("/quality/score", response_model=QualityResponse)
async def quality_score(req: QualityRequest, user=Depends(get_current_user)):
    """품질 평가 5기준 (구조/길이/전문성/완성도/톤) — 0~100점 + A/B+/B/C/D/F 등급.

    features/draft/doc_quality_scorer.evaluate_document(text, doc_type, reference_template)
    → QualityScore dataclass
    """
    try:
        from features.draft.doc_quality_scorer import evaluate_document

        score = evaluate_document(
            text=req.text,
            doc_type=req.doc_type,
            reference_template=req.reference_template,
        )

        # 등급 매핑: A>=90, B+>=85, B>=75, C>=60, D>=40, else F
        # (canonical Draft.jsx 의 grade 계산과 일치)
        total = float(score.total_score)
        if total >= 90:
            grade = "A"
        elif total >= 80:
            grade = "B+"
        elif total >= 70:
            grade = "B"
        elif total >= 50:
            grade = "C"
        elif total >= 30:
            grade = "D"
        else:
            grade = "F"

        # details 의 non-serializable 객체 정리
        clean_details: dict = {}
        for k, v in (score.details or {}).items():
            try:
                json.dumps(v)
                clean_details[k] = v
            except (TypeError, ValueError):
                clean_details[k] = str(v)

        return QualityResponse(
            total_score=total,
            grade=grade,
            scores=QualityScoresDetail(
                structure=float(score.structure_score),
                length=float(score.length_score),
                terminology=float(score.terminology_score),
                completeness=float(score.completeness_score),
                tone=float(score.tone_score),
            ),
            improvements=list(score.improvements or []),
            details=clean_details,
        )
    except Exception as e:
        logger.warning("[draft] quality-score 실패: %s — 기본값 반환", e)
        return QualityResponse(
            total_score=0.0,
            grade="C",
            scores=QualityScoresDetail(),
            improvements=[f"평가 실패: {e}"],
            details={},
        )


@router.post("/diff", response_model=DiffResponse)
async def doc_diff(req: DiffRequest, user=Depends(get_current_user)):
    """버전 diff — frontend lg-diff-line.{add/del/mod/ctx} 마크업 직접 매핑.

    features/draft/doc_diff.compute_diff(old, new, context_lines) → (html, stats)
    + compute_similarity_ratio(old, new) → float
    """
    try:
        from features.draft.doc_diff import compute_diff, compute_similarity_ratio

        diff_html, raw_stats = compute_diff(req.old, req.new, req.context_lines)
        similarity = compute_similarity_ratio(req.old, req.new)

        # difflib unified diff → DiffLine 변환 (Frontend lg-diff-line 매핑)
        import difflib
        old_lines = req.old.splitlines()
        new_lines = req.new.splitlines()
        unified = list(
            difflib.unified_diff(
                old_lines,
                new_lines,
                lineterm="",
                n=req.context_lines,
            )
        )

        lines: list[DiffLine] = []
        for raw in unified:
            if raw.startswith("+++") or raw.startswith("---"):
                continue  # 헤더 스킵
            if raw.startswith("@@"):
                lines.append(DiffLine(type="header", text=raw))
            elif raw.startswith("+"):
                lines.append(DiffLine(type="add", text=raw[1:]))
            elif raw.startswith("-"):
                lines.append(DiffLine(type="del", text=raw[1:]))
            else:
                lines.append(DiffLine(type="ctx", text=raw[1:] if raw.startswith(" ") else raw))

        return DiffResponse(
            lines=lines,
            stats=DiffStats(
                added=raw_stats.get("added", 0),
                removed=raw_stats.get("removed", 0),
                unchanged=raw_stats.get("unchanged", 0),
                similarity=float(similarity),
            ),
            diff_html=diff_html,
        )
    except Exception as e:
        logger.warning("[draft] diff 실패: %s", e)
        return DiffResponse(
            lines=[DiffLine(type="ctx", text=f"diff 계산 실패: {e}")],
            stats=DiffStats(),
            diff_html="",
        )


@router.get(
    "/diagnose",
    response_model=DiagnoseResponse,
    dependencies=[Depends(require_role_level(5))],
)
async def diagnose_draft(request: Request) -> DiagnoseResponse:
    """Module B 의 5개 의존성 진단 — UI 헬스 배너용.

    Plan v1.0 §1.1 — 모호한 "백엔드 연결 실패" 대신 구체 원인 표면화.
    """
    import os as _os
    from pathlib import Path as _Path

    import requests as _req

    from backend.routers.health import redact_diagnostic_url
    from config import KNOWLEDGE_BASE_DIR, OLLAMA_BASE_URL, ollama_headers

    safe_ollama_url = redact_diagnostic_url(OLLAMA_BASE_URL)

    # 1) Ollama (Plan A 변형: Caddy 경유 시 X-AJIN-Secret 부착)
    try:
        r = _req.get(f"{OLLAMA_BASE_URL}/api/tags", headers=ollama_headers(), timeout=2)
        if r.status_code == 200:
            models = [m.get("name", "") for m in r.json().get("models", [])][:20]
            ollama_chk = DiagnoseCheck(
                ok=True,
                detail=f"{len(models)}개 모델 설치됨",
                meta={"models": models, "base_url": safe_ollama_url},
            )
        else:
            ollama_chk = DiagnoseCheck(
                ok=False,
                detail=f"Ollama 응답 비정상 (HTTP {r.status_code})",
                meta={"base_url": safe_ollama_url},
            )
    except Exception as e:
        ollama_chk = DiagnoseCheck(
            ok=False,
            detail=f"Ollama 서버 연결 불가 — `ollama serve` 확인 ({type(e).__name__})",
            meta={"base_url": safe_ollama_url},
        )

    # 2) Gemini API key (.env)
    gemini_key = _os.environ.get("GEMINI_API_KEY", "").strip()
    gemini_chk = DiagnoseCheck(
        ok=bool(gemini_key),
        detail=("API 키 로드됨 (.env)" if gemini_key else "GEMINI_API_KEY 미설정 (.env)"),
        meta={"key_present": bool(gemini_key)},
    )

    # 3) DraftPipeline (ENABLE_FEATURE_B)
    try:
        pipeline_obj = getattr(request.app.state, "draft_pipeline", None)
        pipeline_chk = DiagnoseCheck(
            ok=pipeline_obj is not None,
            detail=(
                "DraftPipeline 부팅 완료"
                if pipeline_obj is not None
                else "DraftPipeline 미부팅 (ENABLE_FEATURE_B=false 또는 부팅 실패)"
            ),
            meta={"enabled": pipeline_obj is not None},
        )
    except Exception as e:
        pipeline_chk = DiagnoseCheck(
            ok=False, detail=f"파이프라인 상태 확인 실패: {e}", meta={}
        )

    # 4) 템플릿
    tdir = KNOWLEDGE_BASE_DIR / "templates"
    j2s: list[str] = []
    if tdir.exists():
        try:
            j2s = [p.name for p in tdir.glob("*.j2")][:30]
        except Exception:
            j2s = []
    templates_chk = DiagnoseCheck(
        ok=bool(j2s),
        detail=f"{len(j2s)}개 .j2 템플릿" if j2s else "템플릿 DB 누락 — data/knowledge_base/templates",
        meta={"count": len(j2s), "samples": j2s[:5]},
    )

    # 5) 프롬프트
    pdir = _Path("features/draft/prompts")
    prompts: list[str] = []
    if pdir.exists():
        try:
            prompts = [p.name for p in pdir.glob("*.txt")]
        except Exception:
            prompts = []
    prompts_chk = DiagnoseCheck(
        ok=bool(prompts),
        detail=f"{len(prompts)}개 프롬프트" if prompts else "프롬프트 누락 — features/draft/prompts",
        meta={"count": len(prompts), "samples": prompts[:5]},
    )

    summary_ok = all([ollama_chk.ok, pipeline_chk.ok, templates_chk.ok, prompts_chk.ok])
    # gemini 는 선택적 — summary_ok 에 포함하지 않음 (Feature B 에서는 차단되므로)

    return DiagnoseResponse(
        ollama=ollama_chk,
        gemini=gemini_chk,
        pipeline=pipeline_chk,
        templates=templates_chk,
        prompts=prompts_chk,
        summary_ok=summary_ok,
    )


# ═══════════════════════════════════════════════════════════
# Plan v1.0 — /stream-v2 (Jinja2 + LLMRouter 통합 SSE)
# ═══════════════════════════════════════════════════════════


@router.post("/stream-v2")
async def draft_stream_v2(req: DraftStreamV2Request, user=Depends(get_current_user)):
    """SSE v2 — Few-shot RAG + 모델 셀렉터(provider/model) + 단계 이벤트.

    Plan v1.0 §2.1 — 가이드라인 v1.7 정식 흐름을 SSE 로 결합.
    Feature B 보안: provider="gemini" 요청은 강제로 ollama 로 다운그레이드.
    """
    import asyncio as _asyncio
    import json as _json
    import os as _os

    from core.llm_router import LLMRouter
    from core.llm_types import LLMMode

    # ── Feature B 보안: Gemini 차단 (Plan v1.0 — 환경변수 토글) ──
    # FEATURE_B_BLOCK_GEMINI=true (기본): Gemini 요청 시 ollama 로 다운그레이드.
    # FEATURE_B_BLOCK_GEMINI=false (시연): Gemini 그대로 사용 (Cloud Run 시연 환경).
    #
    # v3.6 자동 안전망: Ollama 미가용 (OLLAMA_BASE_URL 비어있음) + block_gemini=true 라면
    # 모든 프로바이더가 차단되어 빈 응답이 발생하므로 자동으로 차단을 해제한다.
    # → "왜 출력이 안 나오는가" 의 근본 원인 (Cloud Run + 기본 env 조합).
    block_gemini = (
        os.environ.get("FEATURE_B_BLOCK_GEMINI", "true").strip().lower()
        in ("1", "true", "yes", "on")
    )
    ollama_url = os.environ.get("OLLAMA_BASE_URL", "").strip()
    if block_gemini and not ollama_url:
        # Ollama 없는 환경에서 Gemini 차단은 사실상 서비스 중단 → 자동 해제
        block_gemini = False
    blocked_gemini = block_gemini and req.provider == "gemini"
    effective_provider = "ollama" if blocked_gemini else req.provider
    effective_model = req.model

    # ── 프롬프트 합성 (doc_type_config.build_prompt 우선, fallback 보존) ──
    rag_context = ""
    try:
        from features.draft.fewshot_rag import retrieve_examples  # type: ignore

        examples = retrieve_examples(req.doc_type, req.meta, top_k=3)
        if examples:
            rag_context = "\n\n[유사 사례]\n" + "\n---\n".join(
                e.get("content", "") for e in examples[:3]
            )
    except Exception:
        pass

    prompt: str = ""
    try:
        from features.draft.doc_type_config import build_prompt

        prompt = build_prompt(req.context, req.doc_type, {**req.meta, "user_request": req.user_request})
    except Exception:
        meta_lines = "\n".join(f"- {k}: {v}" for k, v in (req.meta or {}).items() if v)
        prompt = (
            "당신은 아진산업의 업무 문서 작성 전문가입니다.\n"
            f"[문서유형] {req.doc_type}\n"
            f"[어조] {req.tone}\n"
            f"[맥락] {req.context}\n"
            f"[요청] {req.user_request}\n"
            f"[입력]\n{meta_lines}\n\n"
            "한국어로 제목·수신·발신·본문을 포함한 완결된 문서 초안을 작성하세요."
        )
    if rag_context:
        prompt += rag_context

    # v3.6 — 출력 양식 지시 (모든 경로 공통 후행 추가)
    # 일부 LLM(Qwen 3.5 등) 이 한 단락으로 출력하는 문제 보정. 프론트 포스트-프로세서와
    # 함께 작동해 사용자 복사·붙여넣기 시 자연스러운 양식으로 표시되도록 한다.
    prompt += (
        "\n\n[출력 양식 — 반드시 준수]\n"
        "1. 각 섹션(수신/발신/제목/본문/끝맺음)을 빈 줄(빈 라인 1개) 로 구분하세요.\n"
        "2. **헤더:** 형태의 라벨 뒤에는 줄바꿈 후 내용을 작성하세요.\n"
        "3. 문장이 끝나면 마침표 후 새 줄로 이동하세요. 한 단락에 너무 많은 문장을 합치지 마세요.\n"
        "4. 목록은 '- ' 로 시작하고 각 항목을 별도 줄에 작성하세요.\n"
        "5. 출력 끝에 추가 설명·요약·메타 코멘트를 넣지 마세요.\n"
    )

    # ── v3.6: 사용자 업로드 참조 양식 주입 ─────────────────────────────
    # POST /draft/upload-reference 로 추출된 텍스트가 있으면 강력한 지시문으로 prepend.
    # LLM 이 양식 구조(섹션·테이블·머리말 등)를 그대로 따르도록 유도.
    ref_text = (req.reference_template_text or "").strip()
    if ref_text:
        # 너무 큰 텍스트는 prompt 토큰 폭증 방지 — 8000자 (≈ 2k 토큰) 컷
        ref_clipped = ref_text[:8000]
        ref_truncated = len(ref_text) > 8000
        ref_block = (
            "\n\n=========== 참조 양식 (USER-UPLOADED TEMPLATE) ===========\n"
            f"파일명: {req.reference_template_name or 'uploaded.txt'}\n"
            "[중요] 아래 양식의 **구조·머리말·섹션 순서·테이블 형태**를 그대로 따라 작성하세요.\n"
            "내용은 사용자 요청에 맞게 새로 작성하되, 양식의 자리표시자/항목명은 동일하게 유지합니다.\n"
            "----------------------------------------------------------------\n"
            + ref_clipped
            + ("\n... (이하 생략 — 양식의 처음 8000자만 사용)\n" if ref_truncated else "\n")
            + "==============================================================\n"
        )
        # prepend (시스템 지시 다음·사용자 요청 앞)
        prompt = ref_block + "\n" + prompt

    # ── 감사 로깅 ──
    if user:
        try:
            from backend.auth_middleware import log_api_access

            log_api_access(
                endpoint="/api/draft/stream-v2",
                method="POST",
                detail=f"doc_type={req.doc_type}, provider={effective_provider}, model={effective_model}, blocked_gemini={blocked_gemini}",
                user=user,
            )
        except Exception:
            pass

    # ── SSE 생성기 ──
    router_inst: LLMRouter = LLMRouter()

    def _sse(event: str, data) -> str:
        # Plan v1.0 — SSE 사양 준수: data 안의 \n 은 각 줄마다 'data:' 접두사 필요.
        # 그렇지 않으면 빈 줄(\n\n)이 메시지 종료자로 인식되어 token 이 여러 메시지로 잘못 분할됨.
        body = data if isinstance(data, str) else _json.dumps(data, ensure_ascii=False)
        # \r\n / \r 도 정규화
        body = body.replace("\r\n", "\n").replace("\r", "\n")
        formatted = "\n".join(f"data: {line}" for line in body.split("\n"))
        return f"event: {event}\n{formatted}\n\n"

    async def _astream():
        # Plan v1.0 — buffering 우회 패딩
        # Firebase Hosting↔Cloud Run rewrite 구간은 GFE 가 chunked 응답을 ~16KB 까지 buffer 함.
        # 2KB 로는 부족 → 16KB SSE comment(`:` 로 시작 = 클라이언트 무시) 로 강제 flush.
        yield ":" + (" " * 16384) + "\n\n"

        # 1) classify 단계 (간이 — doc_type 이 명시된 경우 그대로 통과)
        yield _sse(
            "stage",
            {"name": "classify", "status": "ok", "meta": {"doc_type": req.doc_type, "context": req.context}},
        )
        # 2) RAG 단계
        yield _sse(
            "stage",
            {"name": "rag", "status": "ok", "meta": {"hits": rag_context.count("[유사 사례]")}},
        )
        # 3) LLM 단계 — security note
        if blocked_gemini:
            yield _sse(
                "stage",
                {
                    "name": "security",
                    "status": "warn",
                    "meta": {
                        "policy": "feature-b-blocks-gemini",
                        "message": "Feature B 보안 정책에 따라 Gemini 요청을 로컬 모델로 다운그레이드했습니다.",
                    },
                },
            )

        force = None
        if effective_provider and effective_model:
            force = (effective_provider, effective_model)

        yield _sse(
            "stage",
            {
                "name": "llm",
                "status": "running",
                "meta": {"provider": effective_provider or "auto", "model": effective_model or "auto"},
            },
        )

        # P1 — LLM stream 을 Queue 로 래핑해 5초 idle 마다 SSE comment heartbeat 발사.
        # 효과: (1) 프록시/GFE idle timeout 회피 (2) 모델 cold load 중에도 연결 유지.
        HEARTBEAT_SEC = 5.0
        queue: _asyncio.Queue = _asyncio.Queue(maxsize=64)

        async def _producer():
            """1차: 사용자 선택(force) 으로 시도.
            2차(P2): 1차에서 토큰 0개 + Gemini 미차단(시연 모드) 이면 폴백 체인으로 재시도.
            """
            primary_got_token = False
            try:
                async for ev in router_inst.stream(
                    prompt=prompt,
                    mode=LLMMode.DRAFT,
                    temperature=0.3,
                    response_format="text",
                    force_provider=force,
                ):
                    if ev.get("type") == "token" and (ev.get("content") or ""):
                        primary_got_token = True
                    await queue.put(("ev", ev))
            except Exception as exc:  # noqa: BLE001
                await queue.put(("exc", exc))
                await queue.put(("end", None))
                return

            # P2 — 1차 실패 + Gemini 차단 안 됨 → 자유 폴백 체인 재시도
            if (not primary_got_token) and (force is not None) and (not block_gemini):
                try:
                    await queue.put((
                        "ev",
                        {
                            "type": "metadata",
                            "content": None,
                            "metadata": {
                                "event": "primary_empty_retry_with_fallback_chain",
                                "primary_provider": effective_provider,
                                "primary_model": effective_model,
                            },
                        },
                    ))
                    async for ev in router_inst.stream(
                        prompt=prompt,
                        mode=LLMMode.DRAFT,
                        temperature=0.3,
                        response_format="text",
                        force_provider=None,  # ollama → ollama_alt → gemini → lm_studio
                    ):
                        await queue.put(("ev", ev))
                except Exception as exc:  # noqa: BLE001
                    await queue.put(("exc", exc))

            await queue.put(("end", None))

        producer = _asyncio.create_task(_producer())
        try:
            while True:
                try:
                    kind, payload = await _asyncio.wait_for(queue.get(), timeout=HEARTBEAT_SEC)
                except _asyncio.TimeoutError:
                    # 5초 동안 토큰 없음 → SSE comment heartbeat (클라이언트는 무시, 프록시 flush)
                    yield ": hb\n\n"
                    continue

                if kind == "end":
                    break
                if kind == "exc":
                    yield _sse(
                        "stage",
                        {"name": "llm", "status": "error", "meta": {"message": str(payload)}},
                    )
                    break

                ev = payload  # kind == "ev"
                etype = ev.get("type")
                if etype == "token":
                    chunk = ev.get("content") or ""
                    if not chunk:
                        continue
                    yield _sse("token", chunk)
                elif etype == "metadata":
                    yield _sse(
                        "stage",
                        {"name": "llm", "status": "running", "meta": ev.get("metadata") or {}},
                    )
                elif etype == "done":
                    yield _sse(
                        "stage",
                        {"name": "llm", "status": "ok", "meta": ev.get("metadata") or {}},
                    )
                elif etype == "error":
                    yield _sse(
                        "stage",
                        {"name": "llm", "status": "error", "meta": {"message": ev.get("content")}},
                    )
        finally:
            if not producer.done():
                producer.cancel()
                try:
                    await producer
                except (_asyncio.CancelledError, Exception):  # noqa: BLE001
                    pass

        yield _sse("stage", {"name": "render", "status": "ok", "meta": {"render_template": req.render_template}})
        yield _sse("done", {"ok": True})

    from fastapi.responses import StreamingResponse

    return StreamingResponse(
        _astream(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache, no-transform",
            "X-Accel-Buffering": "no",
            "Connection": "keep-alive",
        },
    )


@router.post("/stream")
async def draft_stream(req: DraftStreamRequest, user=Depends(get_current_user)):
    """SSE 스트리밍 초안 생성 — Few-shot RAG + LLM 라우터 (DRAFT 모드).

    canonical Draft.jsx 의 generate() 흐름 매핑:
    1. doc_type + tone + meta(title/recipient/content_request) 입력
    2. fewshot_rag retrieve (선택)
    3. doc_type_config.build_prompt 또는 fallback 프롬프트
    4. LLM SSE 스트리밍

    기존 /generate 와 분리 — meta 구조화 + draft 모드 전용.
    """
    from core.llm_client import auto_select_model, stream_generate

    model = req.model or auto_select_model("draft")

    # 1. Few-shot RAG (선택)
    rag_context = ""
    try:
        from features.draft.fewshot_rag import retrieve_examples  # type: ignore
        examples = retrieve_examples(req.doc_type, req.meta, top_k=3)
        if examples:
            rag_context = "\n\n[유사 사례]\n" + "\n---\n".join(
                e.get("content", "") for e in examples[:3]
            )
    except Exception:
        pass  # fewshot_rag 미동작 시 RAG 없이 진행

    # 2. doc_type_config 기반 프롬프트 합성
    prompt: str = ""
    try:
        from features.draft.doc_type_config import build_prompt
        prompt = build_prompt(req.context, req.doc_type, req.meta)
    except Exception:
        # fallback: 기본 프롬프트
        meta_lines = "\n".join(f"- {k}: {v}" for k, v in (req.meta or {}).items() if v)
        prompt = (
            f"당신은 아진산업의 업무 문서 작성 전문가입니다.\n"
            f"[문서유형] {req.doc_type}\n"
            f"[어조] {req.tone}\n"
            f"[맥락] {req.context}\n"
            f"[입력]\n{meta_lines}\n\n"
            f"한국어로 제목·수신·발신·본문을 포함한 완결된 문서 초안을 작성하세요."
        )

    if rag_context:
        prompt += rag_context

    # 3. 감사 로깅
    if user:
        try:
            from backend.auth_middleware import log_api_access
            log_api_access(
                endpoint="/api/draft/stream",
                method="POST",
                detail=f"doc_type={req.doc_type}, tone={req.tone}, ctx={req.context}",
                user=user,
            )
        except Exception:
            pass

    # 4. SSE 스트리밍
    return create_sse_response(
        sse_from_sync_generator(
            stream_generate,
            prompt=prompt,
            model=model,
            feature="draft",
        )
    )


# ─────────────────────────────────────────────────────────────────
# v3.6 — POST /draft/upload-reference
# 사용자가 자체 양식(DOCX/PDF/HWP/TXT/MD)을 업로드하면 텍스트로 추출하여
# /stream-v2 의 reference_template_text 필드에 채워 보낼 수 있게 한다.
# 외부용 신청서 등 양식 그대로 작성해야 하는 시나리오 대응.
# ─────────────────────────────────────────────────────────────────

UPLOAD_MAX_BYTES = int(os.environ.get("UPLOAD_MAX_BYTES", str(5 * 1024 * 1024)))  # 5 MB
UPLOAD_MAX_TEXT_CHARS = 30000  # 추출 텍스트 상한 (LLM 토큰 보호)


def _extract_docx(data: bytes) -> str:
    """DOCX 텍스트 추출 — python-docx 사용."""
    try:
        from docx import Document  # type: ignore
        from io import BytesIO
    except ImportError as e:
        raise HTTPException(status_code=503, detail=f"python-docx 미설치: {e}")
    doc = Document(BytesIO(data))
    parts: list[str] = []
    # 단락
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    # 표 (탭 구분)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append("\t".join(cells))
    return "\n".join(parts)


def _extract_pdf(data: bytes) -> str:
    """PDF 텍스트 추출 — pypdf 사용."""
    try:
        from pypdf import PdfReader  # type: ignore
        from io import BytesIO
    except ImportError as e:
        raise HTTPException(status_code=503, detail=f"pypdf 미설치: {e}")
    reader = PdfReader(BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        try:
            t = page.extract_text() or ""
            if t.strip():
                parts.append(t)
        except Exception:
            continue
    return "\n\n".join(parts)


def _extract_hwp(data: bytes) -> str:
    """HWP 5.0 (OLE compound) 텍스트 추출 — olefile 우선, 실패 시 빈 문자열."""
    try:
        import olefile  # type: ignore
        from io import BytesIO
    except ImportError:
        return ""
    try:
        ole = olefile.OleFileIO(BytesIO(data))
    except Exception:
        return ""
    parts: list[str] = []
    # PrvText 스트림 (한컴 5.0 표준 — 미리보기 텍스트)
    if ole.exists("PrvText"):
        try:
            with ole.openstream("PrvText") as stream:
                raw = stream.read()
                # UTF-16 LE
                text = raw.decode("utf-16-le", errors="ignore")
                if text.strip():
                    parts.append(text)
        except Exception:
            pass
    ole.close()
    return "\n".join(parts)


def _extract_hwpx(data: bytes) -> str:
    """HWPX (OWPML zip) 텍스트 추출 — Contents/section*.xml 텍스트만 긁어내기."""
    try:
        import zipfile
        import re as _re
        from io import BytesIO
    except ImportError:
        return ""
    parts: list[str] = []
    try:
        with zipfile.ZipFile(BytesIO(data)) as zf:
            for name in zf.namelist():
                if name.startswith("Contents/section") and name.endswith(".xml"):
                    raw = zf.read(name).decode("utf-8", errors="ignore")
                    # <hp:t>텍스트</hp:t> 만 추출
                    for m in _re.finditer(r"<hp:t[^>]*>([^<]*)</hp:t>", raw):
                        t = m.group(1).strip()
                        if t:
                            parts.append(t)
    except Exception:
        return ""
    return " ".join(parts)


@router.post("/upload-reference", response_model=UploadReferenceResponse)
async def upload_reference(
    file: UploadFile = File(...),
    user=Depends(get_current_user),
):
    """사용자 양식 업로드 — 텍스트 추출 후 반환.

    프론트는 응답의 `text` 를 그대로 `/stream-v2` 의 reference_template_text 필드로 전달.
    파일 자체는 서버에 영속 저장하지 않음 (메모리에서 추출 후 폐기).
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="파일명이 비어 있습니다.")

    # 크기 제한
    data = await file.read()
    if len(data) > UPLOAD_MAX_BYTES:
        raise HTTPException(
            status_code=413,
            detail=f"파일이 너무 큽니다 ({len(data)} > {UPLOAD_MAX_BYTES} bytes).",
        )
    if len(data) == 0:
        raise HTTPException(status_code=400, detail="빈 파일입니다.")

    # 확장자 기반 라우팅
    name = file.filename.lower()
    detected = "unsupported"
    text = ""
    warning = ""

    try:
        if name.endswith(".docx"):
            detected = "docx"
            text = _extract_docx(data)
        elif name.endswith(".pdf"):
            detected = "pdf"
            text = _extract_pdf(data)
            if not text.strip():
                warning = "PDF 에서 텍스트를 추출하지 못했습니다 (스캔 이미지 가능성)."
        elif name.endswith(".hwpx"):
            detected = "hwpx"
            text = _extract_hwpx(data)
            if not text.strip():
                warning = "HWPX 에서 텍스트를 추출하지 못했습니다 (양식이 비표준일 수 있습니다)."
        elif name.endswith(".hwp"):
            detected = "hwp"
            text = _extract_hwp(data)
            if not text.strip():
                warning = "HWP 본문 추출 제한 — 미리보기(PrvText) 만 사용 가능합니다."
        elif name.endswith((".txt", ".md", ".markdown")):
            detected = "txt" if name.endswith(".txt") else "md"
            text = data.decode("utf-8", errors="replace")
        else:
            raise HTTPException(
                status_code=415,
                detail=(
                    "지원하지 않는 파일 형식입니다. "
                    "지원: .docx, .pdf, .hwp, .hwpx, .txt, .md"
                ),
            )
    except HTTPException:
        raise
    except Exception as e:
        logger.warning("upload-reference 추출 실패 (%s): %s", name, e)
        raise HTTPException(
            status_code=422,
            detail=f"파일 추출 실패: {type(e).__name__}: {e}",
        )

    truncated = len(text) > UPLOAD_MAX_TEXT_CHARS
    if truncated:
        text = text[:UPLOAD_MAX_TEXT_CHARS]

    # 감사 로깅
    if user:
        try:
            from backend.auth_middleware import log_api_access
            log_api_access(
                endpoint="/api/draft/upload-reference",
                method="POST",
                detail=f"file={file.filename}, format={detected}, chars={len(text)}",
                user=user,
            )
        except Exception:
            pass

    return UploadReferenceResponse(
        ok=bool(text.strip()),
        filename=file.filename,
        extracted_chars=len(text),
        truncated=truncated,
        text=text,
        detected_format=detected,
        warning=warning,
    )


# ═══════════════════════════════════════════════════════════════
# B3 v4.0 — 대화형 부분 수정 (Partial Edit)
# 명시적 마커 파싱으로 본문을 섹션 단위로 분리하고, 사용자가 지정한 단일
# 섹션만 LLM 으로 재작성. 전체 재생성 대비 토큰·시간 ~50% 절감.
# ═══════════════════════════════════════════════════════════════


@router.post("/scan-sections", response_model=PartialEditScanResponse)
async def scan_sections(req: PartialEditScanRequest):
    """본문을 명시적 마커(##, D1., [Section: ...] 등)로 분리해 섹션 리스트 반환.

    LLM 호출 없음 — 순수 파싱. 프론트가 본문 위에 호버 영역을 그리는 데 사용.
    """
    from features.draft.partial_editor import find_sections

    sections = find_sections(req.body or "")
    items = [
        PartialEditSection(
            index=s.index,
            marker=s.marker,
            title=s.title,
            start=s.start,
            end=s.end,
            preview=s.content[:80].replace("\n", " "),
        )
        for s in sections
    ]
    return PartialEditScanResponse(sections=items, total=len(items))


@router.post("/partial-edit", response_model=PartialEditResponse)
async def partial_edit(req: PartialEditRequest, user=Depends(get_current_user)):
    """타겟 섹션 한 단락만 LLM 으로 재작성하고 본문 교체본을 반환.

    - 입력: 전체 본문 + 타겟 섹션 인덱스 + 사용자 지시 ("조금 더 정중하게")
    - 동작: 명시적 마커 파싱 → 직전/다음 섹션 컨텍스트와 함께 LLM 호출
    - 출력: 새 섹션 본문 + 본문 교체본 (전체)
    """
    from features.draft.partial_editor import (
        build_partial_edit_context,
        render_partial_edit_prompt,
        replace_section,
    )

    ctx = build_partial_edit_context(
        body=req.body,
        target_section_index=req.target_section_index,
        instruction=req.instruction,
        metadata={"doc_type": req.doc_type, "tone": req.tone},
    )
    if ctx is None:
        raise HTTPException(
            status_code=400,
            detail="섹션을 찾을 수 없거나 인덱스가 범위 밖입니다 — ##, D1., [Section: …] 같은 명시적 마커가 필요합니다.",
        )

    prompt = render_partial_edit_prompt(ctx)

    # LLM 호출 — 라우터 체인 사용 (gemini→ollama 폴백 자동)
    try:
        from core.llm_router import LLMRouter
        from core.llm_types import LLMTask

        router_instance = LLMRouter()
        result = router_instance.generate(
            task=LLMTask.DRAFT,
            prompt=prompt,
            stream=False,
        )
        new_section_content = (result.content or "").strip()
        model = result.model or ""
        provider = result.provider or ""
    except Exception as e:
        logger.warning("[draft.partial-edit] LLMRouter 실패 — 단순 stream_generate 폴백: %s", e)
        from core.llm_client import auto_select_model, generate

        model = auto_select_model("draft")
        provider = "ollama"
        new_section_content = generate(prompt=prompt, model=model, feature="draft").strip()

    if not new_section_content:
        raise HTTPException(status_code=502, detail="LLM 이 빈 응답을 반환했습니다.")

    new_body = replace_section(req.body, ctx.target_section, new_section_content)

    if user:
        try:
            from backend.auth_middleware import log_api_access
            log_api_access(
                endpoint="/api/draft/partial-edit",
                method="POST",
                detail=f"section={req.target_section_index} doc_type={req.doc_type}",
                user=user,
            )
        except Exception:
            pass

    return PartialEditResponse(
        new_body=new_body,
        new_section_content=new_section_content,
        target_section_index=req.target_section_index,
        model=model,
        provider=provider,
    )


# ═══════════════════════════════════════════════════════════════
# B4 v4.0 — 버전 관리 (사용자별 영속 + 단일 검토자 워크플로우)
# 상태: draft → under_review → approved/rejected
# ═══════════════════════════════════════════════════════════════


def _user_id_from(user) -> str:
    """UserContext or dict 에서 employee_id 추출. 없으면 빈 문자열."""
    if user is None:
        return ""
    eid = getattr(user, "employee_id", "") or ""
    if not eid and isinstance(user, dict):
        eid = user.get("employee_id", "") or ""
    return eid


def _user_dept_from(user) -> str:
    if user is None:
        return ""
    dept = getattr(user, "department", "") or ""
    if not dept and isinstance(user, dict):
        dept = user.get("department", "") or ""
    return dept


def _user_role_level_from(user) -> int:
    """Return the RBAC level for review-workflow decisions."""
    if user is None:
        return 0
    raw_level = getattr(user, "role_level", None)
    if raw_level is not None:
        try:
            return int(raw_level)
        except (TypeError, ValueError):
            pass
    role = getattr(user, "role", "") or ""
    if isinstance(user, dict):
        role = role or user.get("role", "")
    try:
        from core.auth.rbac import get_role_level
        return int(get_role_level(str(role)))
    except Exception:
        return 0


@router.post("/versions", response_model=VersionSaveResponse)
async def save_draft_version(
    req: VersionSaveRequest, user=Depends(get_current_user)
):
    """초안 버전 저장. document_id 가 없으면 새 문서로 생성, 있으면 기존 문서에 새 버전 append."""
    from features.draft.version_db import save_version

    author = _user_id_from(user) or "anonymous"
    dept = _user_dept_from(user)

    result = save_version(
        doc_type=req.doc_type,
        title=req.title or req.doc_type,
        author=author,
        department=dept,
        template_vars=req.template_vars,
        rendered_text=req.rendered_text,
        change_summary=req.change_summary,
        document_id=req.document_id,
    )
    return VersionSaveResponse(**result)


@router.get("/versions", response_model=VersionListResponse)
async def list_draft_versions(
    scope: Literal["mine", "review", "all"] = "mine",
    status: str = "",
    limit: int = 50,
    user=Depends(get_current_user),
):
    """사용자별 버전 목록.

    - scope=mine: 내가 작성한 버전 (created_by = me)
    - scope=review: 내가 검토자로 지정된 버전 (reviewer_id = me)
    - scope=all: 전체 (관리자 권장)
    """
    from features.draft.version_db import list_versions_for_user

    me = _user_id_from(user)
    author_filter = me if scope == "mine" and me else ""
    reviewer_filter = me if scope == "review" and me else ""

    rows = list_versions_for_user(
        author=author_filter,
        reviewer=reviewer_filter,
        status=status,
        limit=limit,
    )
    items = [VersionListItem(**r) for r in rows]
    return VersionListResponse(items=items, total=len(items))


@router.get("/versions/{version_id}", response_model=VersionDetailResponse)
async def get_draft_version(version_id: int, user=Depends(get_current_user)):
    """특정 버전의 본문 + 메타. 비교 / rollback 용."""
    from features.draft.version_db import get_version

    row = get_version(version_id)
    if not row:
        raise HTTPException(status_code=404, detail="버전을 찾을 수 없습니다.")
    # template_vars 가 dict 로 풀려있음 (version_db.get_version 이 처리)
    return VersionDetailResponse(
        version_id=row.get("id", version_id),
        document_id=row.get("document_id", 0),
        version_num=row.get("version_num", 1),
        rendered_text=row.get("rendered_text", ""),
        template_vars=row.get("template_vars", {}),
        change_summary=row.get("change_summary", ""),
        created_at=row.get("created_at", ""),
        created_by=row.get("created_by", ""),
        status=row.get("status", "draft"),
        reviewer_id=row.get("reviewer_id", ""),
        reviewed_at=row.get("reviewed_at", ""),
        review_note=row.get("review_note", ""),
        doc_type=row.get("doc_type", ""),
        title=row.get("title", ""),
        author=row.get("author", ""),
        department=row.get("department", ""),
    )


@router.post("/versions/{version_id}/review")
async def review_draft_version(
    version_id: int,
    req: VersionReviewRequest,
    user=Depends(get_current_user),
):
    """버전 검토 액션 (submit / approve / reject).

    - submit: 작성자가 검토자 지정 → status: draft → under_review
    - approve: 지정 검토자가 승인 → status: under_review → approved
    - reject: 지정 검토자가 반려 → status: under_review → rejected
    """
    from features.draft.version_db import (
        approve_version,
        get_version,
        reject_version,
        submit_for_review,
    )

    me = _user_id_from(user)

    result = None
    if req.action == "submit":
        result = submit_for_review(version_id, req.reviewer_id)
        if result is None:
            raise HTTPException(
                status_code=400,
                detail="제출 실패: 버전을 찾을 수 없거나 이미 검토 중/승인된 상태입니다.",
            )
    elif req.action == "approve":
        before = get_version(version_id)
        if before and (before.get("created_by") or before.get("author") or "").strip() == me:
            raise HTTPException(
                status_code=403,
                detail="승인 실패: 작성자는 본인 문서를 승인할 수 없습니다.",
            )
        reviewer = req.reviewer_id.strip() or me
        admin_override = bool(
            before
            and before.get("reviewer_id")
            and before.get("reviewer_id") != reviewer
            and _user_role_level_from(user) >= 5
        )
        note = req.note
        if admin_override:
            note = f"[L5 override by {me}] {note}".strip()
        result = approve_version(
            version_id,
            reviewer,
            note=note,
            allow_reviewer_override=admin_override,
        )
        if result is None:
            raise HTTPException(
                status_code=403,
                detail="승인 실패: 권한이 없거나 이미 처리된 버전입니다.",
            )
        if admin_override and user:
            try:
                from backend.auth_middleware import log_api_access
                log_api_access(
                    endpoint=f"/api/draft/versions/{version_id}/review",
                    method="POST",
                    status_code=200,
                    detail=(
                        "L5 reviewer override "
                        f"assigned={before.get('reviewer_id')} actual={reviewer}"
                    ),
                    user=user,
                    intent="draft_review_override",
                    result_count=1,
                )
            except Exception:
                pass
    elif req.action == "reject":
        reviewer = req.reviewer_id.strip() or me
        result = reject_version(version_id, reviewer, note=req.note)
        if result is None:
            raise HTTPException(
                status_code=403,
                detail="반려 실패: 권한이 없거나 검토 중 상태가 아닙니다.",
            )

    if result is None:
        raise HTTPException(status_code=400, detail="알 수 없는 review action 입니다.")

    return {
        "ok": True,
        "version_id": version_id,
        "new_status": result.get("status", ""),
        "reviewer_id": result.get("reviewer_id", ""),
    }


# ═══════════════════════════════════════════════════════════════
# B6 v4.0 — 메일 발송 인터페이스 + 첨부 추천
# ═══════════════════════════════════════════════════════════════


@router.get("/mail/attachment-recommendations", response_model=AttachmentRecommendResponse)
async def attachment_recommendations(doc_type: str):
    """doc_type 별 첨부 추천 + 필수 항목 라벨."""
    from features.draft.attachment_recommender import (
        recommend_attachments,
        required_attachment_labels,
    )

    suggestions = recommend_attachments(doc_type)
    return AttachmentRecommendResponse(
        doc_type=doc_type,
        items=[
            AttachmentSuggestionItem(
                label=s.label,
                description=s.description,
                required=s.required,
                file_hint=s.file_hint,
            )
            for s in suggestions
        ],
        required_labels=required_attachment_labels(doc_type),
    )


@router.post("/mail/send", response_model=MailSendResponse)
async def send_mail(req: MailSendRequest, user=Depends(get_current_user)):
    """메일 발송 — 어댑터(mock/SMTP/Graph) 위임.

    Feature B Sprint 1 P0 (plan §14.2):
      - 발송 전 MailSendGuard.validate() 로 7 규칙 검증
      - 가드 차단 시 412 + mail_audit_log 에 차단 사유 기록
      - 발송 후 mail_audit_log 에 결과 기록 (external_domains 포함)
      - watermark_id 는 export 단계에서 부여된 SHA1 8자 (감사 추적)
    """
    from features.draft.mail_sender import (
        MailMessage,
        MailRecipient,
        get_mail_adapter,
    )
    from features.draft.mail_guard import get_guard, ALLOW
    from features.draft import mail_audit

    sender_id = getattr(user, "employee_id", "") if user else ""

    # ── 1. Pre-send guard (7 규칙) ──────────────────────────────
    guard = get_guard()
    report = guard.validate(req, user)

    if not report.ok:
        # 차단 — 감사 로그에 기록 후 412.
        # `external_domains` 는 가드가 분류한 결과 사용.
        try:
            mail_audit.append(
                user_id=getattr(user, "user_id", None) if user else None,
                sender_email=sender_id,
                version_id=req.version_id or None,
                version_status=report.version_status,
                adapter="",
                ok=False,
                message_id="",
                to_recipients=[r.email for r in req.to],
                cc_recipients=[r.email for r in req.cc],
                bcc_recipients=[r.email for r in req.bcc],
                external_domains=report.classified.get("external", []),
                guard_decision=report.decision,
                watermark_id=req.watermark_id,
                detail=report.detail,
            )
        except Exception:
            pass
        raise HTTPException(
            status_code=412,
            detail={"decision": report.decision, "message": report.detail},
        )

    # ── 2. 어댑터 발송 ───────────────────────────────────────────
    adapter = get_mail_adapter()
    msg = MailMessage(
        subject=req.subject,
        body=req.body,
        body_format=req.body_format,
        to=[MailRecipient(email=r.email, name=r.name) for r in req.to],
        cc=[MailRecipient(email=r.email, name=r.name) for r in req.cc],
        bcc=[MailRecipient(email=r.email, name=r.name) for r in req.bcc],
        attachments=list(req.attachments),
        metadata={
            **req.metadata,
            "doc_type": req.doc_type,
            "sender": sender_id,
            "version_id": req.version_id,
            "watermark_id": req.watermark_id,
        },
    )
    result = adapter.send(msg)

    # ── 3. Audit log ─────────────────────────────────────────────
    try:
        mail_audit.append(
            user_id=getattr(user, "user_id", None) if user else None,
            sender_email=sender_id,
            version_id=req.version_id or None,
            version_status=report.version_status,
            adapter=result.adapter,
            ok=result.ok,
            message_id=result.message_id,
            to_recipients=[r.email for r in req.to],
            cc_recipients=[r.email for r in req.cc],
            bcc_recipients=[r.email for r in req.bcc],
            external_domains=report.classified.get("external", []),
            guard_decision=ALLOW,
            watermark_id=req.watermark_id,
            detail=result.detail,
        )
    except Exception:
        pass

    # ── 4. 기존 api_audit_log (compat) ───────────────────────────
    if user:
        try:
            from backend.auth_middleware import log_api_access
            log_api_access(
                endpoint="/api/draft/mail/send",
                method="POST",
                status_code=200,
                detail=(
                    f"adapter={result.adapter} ok={result.ok} "
                    f"to={len(req.to)} cc={len(req.cc)} att={len(req.attachments)} "
                    f"version={req.version_id}"
                ),
                user=user,
                intent="mail_send",
                result_count=len(req.to) + len(req.cc) + len(req.bcc),
            )
        except Exception:
            pass

    return MailSendResponse(
        ok=result.ok,
        message_id=result.message_id,
        sent_at=result.sent_at,
        adapter=result.adapter,
        detail=result.detail,
    )


# ═══════════════════════════════════════════════════════════════
# 부록 5 P3-1 — 사용자 prefs 서버 sync (즐겨찾기 다중 기기)
# ═══════════════════════════════════════════════════════════════


def _user_id_of(user) -> str:
    if user is None:
        return ""
    return str(getattr(user, "user_id", "") or getattr(user, "username", "") or "")


@router.get("/prefs/me")
async def get_my_draft_prefs(user=Depends(get_current_user)):
    """현재 사용자의 즐겨찾기 prefs 반환. 미인증 시 빈 결과."""
    from backend.services import draft_prefs
    return draft_prefs.get_prefs(_user_id_of(user))


@router.put("/prefs/me")
async def update_my_draft_prefs(payload: dict, user=Depends(get_current_user)):
    """현재 사용자의 즐겨찾기 prefs 저장.

    부록 6 P4-4 — 페이로드에 `base_version` 이 있으면 3-way merge (multi-device 충돌 해결).
    없으면 last-write-wins (이전 동작 호환).
    """
    from backend.services import draft_prefs
    uid = _user_id_of(user)
    if not uid:
        raise HTTPException(status_code=401, detail="인증 필요 — 미인증 사용자는 로컬 prefs 만 가능")
    favs = payload.get("favorited_doc_types") or []
    if not isinstance(favs, list):
        raise HTTPException(status_code=400, detail="favorited_doc_types 는 리스트")
    base_version = str(payload.get("base_version") or "")
    removed_ids = payload.get("removed_ids") or []
    if base_version:
        return draft_prefs.merge_prefs(uid, favs, base_version=base_version,
                                       removed_ids=list(removed_ids))
    return draft_prefs.upsert_prefs(uid, favs)


# ═══════════════════════════════════════════════════════════════
# 부록 5 P3-2 — TF-IDF 기반 템플릿 추천
# ═══════════════════════════════════════════════════════════════


@router.post("/recommend")
async def recommend_doc_types(payload: dict, user=Depends(get_current_user)):
    """자유 입력 query → cosine similarity → top-2 docType id 추천."""
    from backend.services import draft_recommender
    q = (payload.get("user_request") or "").strip()
    top_k = int(payload.get("top_k") or 2)
    return draft_recommender.recommend(q, top_k=top_k)


# ═══════════════════════════════════════════════════════════════
# 부록 6 P4-1 + P4-3 — 부서 / 개인 추천 (Firestore aggregation 결과 조회)
# ═══════════════════════════════════════════════════════════════


def _user_department_of(user) -> str:
    if user is None:
        return ""
    return str(getattr(user, "department", "") or getattr(user, "dept", "") or "")


@router.get("/recommend/by-dept")
async def recommend_by_dept(top_k: int = 3, user=Depends(get_current_user)):
    """현재 사용자의 부서 기준 자주 쓰는 docType top-K. 부서 미설정 시 빈 결과."""
    from backend.services.draft_dept_agg import get_dept_recommendations
    dept = _user_department_of(user)
    return {
        "department": dept,
        "recommendations": get_dept_recommendations(dept, top_k=top_k),
    }


@router.get("/recommend/personal")
async def recommend_personal(top_k: int = 5, user=Depends(get_current_user)):
    """사용자 개인 빈도 기반 top-K. min_count=3 이상만 (데이터 부족 시 빈 결과)."""
    from backend.services.draft_dept_agg import get_personal_picks
    uid = _user_id_of(user)
    return {
        "user_id": uid,
        "picks": get_personal_picks(uid, top_k=top_k, min_count=3),
    }


@router.post("/recommend/aggregate-run")
async def trigger_aggregate_run(user=Depends(get_current_user)):
    """수동 트리거 (admin/디버깅용) — daily 잡 이전에 즉시 실행."""
    from backend.services.draft_dept_agg import run_aggregation
    return run_aggregation()
