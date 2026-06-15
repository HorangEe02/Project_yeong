"""규정 준수 라우터.

v3.0: 3-Tier 권한 적용
- Tier 1 VIEW: 모든 인증 사용자 (scenarios, facilities, risk, timeline, network, tariff)
- Tier 2 ANALYZE: 관련 부서 EMPLOYEE+ (check, classify)
- Tier 3 OPERATE: TEAM_LEAD+ (crawl, changes/ack)
"""

import asyncio
import json
import logging
from pathlib import Path
from typing import Any

from fastapi import APIRouter, Depends, HTTPException
from fastapi.routing import APIRoute
from starlette.requests import Request

logger = logging.getLogger(__name__)

from backend.schemas.compliance import (
    AcknowledgeResponse,
    AffectedContractItem,
    AffectedContractsResponse,
    AffectedSupplierItem,
    AffectedSuppliersResponse,
    AlternativeSupplierItem,
    AlternativesResponse,
    CollabTicketItem,
    CollabTicketListResponse,
    CollabTicketOwner,
    CreateTicketResponse,
    DelegationDryRunMatch,
    DelegationDryRunResponse,
    DelegationRuleCreateRequest,
    DelegationRuleItem,
    DelegationRuleListResponse,
    DelegationRuleMutationResponse,
    DelegationRuleUpdateRequest,
    LearningMentorQueueItem,
    LearningMentorQueueResponse,
    LearningMyProgressResponse,
    LearningPathCreateRequest,
    LearningPathCreateResponse,
    LearningPathDetail,
    LearningProgressItem,
    LearningQuizPreviewResponse,
    LearningQuizRequest,
    LearningQuizResponse,
    LearningRequestReviewResponse,
    LearningReviewRequest,
    LearningReviewResponse,
    ApprovalChainCreateRequest,
    ApprovalChainCreateResponse,
    ApprovalChainDetailResponse,
    ApprovalDecideRequest,
    ApprovalDecideResponse,
    ApprovalMyPendingResponse,
    ApprovalStepItem,
    JiraHealthResponse,
    JiraWebhookAckResponse,
    JiraWebhookEvent,
    SupplierDiscoveryCandidate,
    SupplierDiscoveryListResponse,
    SupplierPromoteRequest,
    SupplierPromoteResponse,
    PnlImpactItem,
    SupplierCycleResponse,
    SupplierGraphNode,
    SupplierGraphResponse,
    WhatIfAccountingRequest,
    WhatIfAccountingResponse,
    WhatIfBaselineResponse,
    DeptHandlingTime,
    ExtendedTrendResponse,
    FeedbackLoopApplyRequest,
    FeedbackLoopApplyResponse,
    TicketTransitionRequest,
    TicketTransitionResponse,
    IndustryContextResponse,
    IndustryCorpFilings,
    IndustryFetchResponse,
    WhatIfRequest,
    WhatIfResponse,
    CaseLawIndexResponse,
    CaseLawItem,
    ChangeCorrectionRequest,
    ChangeCorrectionResponse,
    ChangeFeedItem,
    ChangeFeedResponse,
    ChangeItem,
    ChangeKpiResponse,
    ChangeListResponse,
    ChangeTransitionRequest,
    ChangeTransitionResponse,
    ContractIngestResponse,
    ContractListResponse,
    ContractMeta,
    CorrectionStatsResponse,
    CostSimulationRequest,
    CostSimulationResponse,
    CostSimulationSupplierBreakdown,
    SendAssessmentRequest,
    SendAssessmentResponse,
    SimilarCasesResponse,
    SupplierComponentItem,
    SupplierDetailResponse,
    SupplierImportRequest,
    SupplierImportResponse,
    SupplierListResponse,
    SupplierMeta,
    ClassifyRequest,
    ClassifyResponse,
    ComplianceCheckRequest,
    ComplianceCheckResponse,
    CrawlRunAllResponse,
    CrawlRunResponse,
    CrawlResultMeta,
    CrawlResultsListResponse,
    CrawlResultItem,
    CrawlResultDetailResponse,
    FacilityItem,
    PlotlyResponse,
    RiskScoreItem,
    RiskScoreResponse,
    ScenarioChangeVersion,
    ScenarioDetailResponse,
    ScenarioReference,
    ScenarioRegulationMeta,
    ScenarioSimRiskScore,
    ScenarioSimImpact,
    ScenarioSimEvidence,
    ScenarioSimulateRequest,
    ScenarioSimulateResponse,
    TariffSimulateRequest,
    TariffSimulateResponse,
)
from backend.dependencies import (
    get_current_user,
    require_permission,
    require_role_level,
    resolve_user_role_level,
)
from config import DATA_DIR


def _feature_d_group_for_path(path: str) -> str:
    """Compliance route path를 Feature D 하위 플래그로 분류한다.

    Args:
        path: APIRoute path_format 값. 예: "/changes/{change_id}/acknowledge".

    Returns:
        str: `FeatureDFlags` 속성명. 알 수 없는 신규 D endpoint는 기본 봉인된다.
    """
    normalized = path
    for prefix in ("/api/compliance", "/compliance"):
        if normalized.startswith(prefix + "/"):
            normalized = normalized[len(prefix):]
            break
    path = normalized

    if path in {
        "/facilities",
        "/changes/recent",
        "/changes/feed",
        "/changes/kpi",
        "/crawl/run-all",
        "/crawl/history",
        "/crawl/history/stats",
        "/crawl/results",
        "/alarms/recent",
        "/alarms/stream",
        "/scheduler/jobs",
        "/digest/run-now",
    }:
        return "d1_alerts"
    if path.startswith("/changes/") and (
        path.endswith("/acknowledge") or path.endswith("/transition")
    ):
        return "d1_alerts"
    if path.startswith("/crawl/run/"):
        return "d1_alerts"
    if path.startswith("/crawl/results/"):
        if "bulk-download" in path:
            return "d4_workflow"
        return "d1_alerts"
    if path.startswith("/alarms/") and path.endswith("/ack"):
        return "d1_alerts"
    if path.startswith("/scheduler/trigger/"):
        return "d1_alerts"

    if path.startswith((
        "/search",
        "/regulations",
        "/glossary",
        "/docs",
        "/case-law",
        "/contracts",
    )) or path in {"/check", "/classify"}:
        return "d2_rag"
    if path.startswith("/changes/") and (
        path.endswith("/similar-cases") or path.endswith("/affected-contracts")
    ):
        return "d2_rag"

    if path.startswith(("/tariff", "/whatif", "/scenarios", "/risk", "/timeline", "/network")):
        return "d3_whatif"
    if path == "/impact-network" or path.startswith("/changes/") and path.endswith("/cost-simulation"):
        return "d3_whatif"

    if path.startswith(("/suppliers", "/industry-trend", "/admin/suppliers")):
        return "d5_supply"
    if path.startswith("/changes/") and (
        path.endswith("/industry-context")
        or path.endswith("/affected-suppliers")
    ):
        return "d5_supply"

    if path.startswith((
        "/tickets",
        "/jira",
        "/approvals",
        "/delegation-rules",
        "/learning-path",
        "/sop",
        "/feedback-loop",
    )):
        return "d4_workflow"
    if path.startswith("/changes/"):
        return "d4_workflow"

    return "d4_workflow"


def _feature_d_enabled(group: str) -> bool:
    """Feature D 하위 플래그 활성 여부를 런타임 환경변수에서 확인한다.

    Args:
        group: `FeatureDFlags` 속성명.

    Returns:
        bool: 활성화되어 있으면 True.
    """
    from core.feature_flags import load_feature_d_flags

    return bool(getattr(load_feature_d_flags(), group, False))


class FeatureDGatedRoute(APIRoute):
    """Feature D 하위 기능을 런타임 404와 OpenAPI 숨김으로 봉인하는 route."""

    feature_d_group: str

    def __init__(self, *args: Any, **kwargs: Any) -> None:
        """APIRoute 생성 시 기본 OpenAPI 노출 여부를 결정한다.

        Args:
            *args: FastAPI APIRoute positional arguments.
            **kwargs: FastAPI APIRoute keyword arguments.
        """
        super().__init__(*args, **kwargs)
        self.feature_d_group = _feature_d_group_for_path(self.path_format)
        if not _feature_d_enabled(self.feature_d_group):
            self.include_in_schema = False

    def get_route_handler(self):
        """비활성 Feature D route를 런타임에 404로 숨기는 handler를 반환한다.

        Returns:
            Callable: FastAPI route handler.
        """
        original_handler = super().get_route_handler()

        async def gated_handler(request: Request):
            if not _feature_d_enabled(self.feature_d_group):
                raise HTTPException(status_code=404, detail="feature_disabled")
            return await original_handler(request)

        return gated_handler


router = APIRouter(prefix="/compliance", tags=["compliance"], route_class=FeatureDGatedRoute)


_VALID_CHANGE_STATUSES = {"pending", "reviewing", "planning", "announced", "done", "filtered"}


def _current_role_level(user) -> int:
    """Resolve a user's RBAC level using the same fallback as role dependencies.

    Args:
        user: Authenticated user context.

    Returns:
        int: RBAC level. Missing or invalid roles fail closed to L0.
    """
    return resolve_user_role_level(user)


def _transition_actor_id(user) -> str:
    """Return a stable actor id for compliance transition audit rows.

    Args:
        user: Authenticated user context.

    Returns:
        str: Employee id, user id, email, or username.
    """

    return str(
        getattr(user, "employee_id", "")
        or getattr(user, "user_id", "")
        or getattr(user, "email", "")
        or getattr(user, "username", "")
        or ""
    )


def _load_change_for_transition(change_id: int) -> dict[str, Any] | None:
    """Load a change row needed for legal transition guardrails.

    Args:
        change_id: Regulation change row id.

    Returns:
        dict[str, Any] | None: Change row when present.
    """

    import sqlite3
    from features.compliance.change_detector import CHANGE_DB_PATH, init_change_db

    init_change_db()
    conn = sqlite3.connect(CHANGE_DB_PATH)
    conn.row_factory = sqlite3.Row
    try:
        row = conn.execute(
            """SELECT id, grade, status, audit_trail, legal_class,
                      penalty_severity_krw_mn
               FROM regulation_changes WHERE id = ?""",
            (change_id,),
        ).fetchone()
        return dict(row) if row is not None else None
    finally:
        conn.close()


def _append_change_audit_event(change_id: int, event: dict[str, Any]) -> None:
    """Append a structured compliance audit event to a change row.

    Args:
        change_id: Regulation change row id.
        event: Secret-safe audit event.
    """

    import sqlite3
    from datetime import datetime
    from features.compliance.alerts.legal_guard import parse_json_list
    from features.compliance.change_detector import CHANGE_DB_PATH, init_change_db

    init_change_db()
    conn = sqlite3.connect(CHANGE_DB_PATH)
    conn.row_factory = sqlite3.Row
    try:
        row = conn.execute(
            "SELECT audit_trail FROM regulation_changes WHERE id = ?",
            (change_id,),
        ).fetchone()
        if row is None:
            return
        trail = parse_json_list(row["audit_trail"])
        payload = {"ts": datetime.now().isoformat(), **event}
        trail.append(payload)
        conn.execute(
            "UPDATE regulation_changes SET audit_trail = ? WHERE id = ?",
            (json.dumps(trail, ensure_ascii=False), change_id),
        )
        conn.commit()
    finally:
        conn.close()


def _enforce_change_transition_guard(
    change_id: int,
    req: ChangeTransitionRequest,
    user,
) -> tuple[bool, bool]:
    """Enforce RBAC and legal-review requirements for status transitions.

    Args:
        change_id: Regulation change row id.
        req: Transition request.
        user: Authenticated user context already checked at L3+.

    Returns:
        tuple[bool, bool]: ``review_required`` and ``override_used``.

    Raises:
        HTTPException: If the transition violates RBAC or review policy.
    """

    from features.compliance.alerts.legal_guard import (
        LEGAL_FINAL_STATUSES,
        has_independent_human_review,
        parse_json_list,
        requires_legal_review,
    )

    if req.new_status not in _VALID_CHANGE_STATUSES:
        raise HTTPException(
            status_code=400,
            detail=f"전환 실패 — invalid status '{req.new_status}' 또는 change_id={change_id} 미존재",
        )
    row = _load_change_for_transition(change_id)
    if row is None:
        raise HTTPException(
            status_code=400,
            detail=f"전환 실패 — invalid status '{req.new_status}' 또는 change_id={change_id} 미존재",
        )

    role_level = _current_role_level(user)
    actor_id = _transition_actor_id(user)
    is_final_status = req.new_status in LEGAL_FINAL_STATUSES
    if is_final_status and role_level < 4:
        raise HTTPException(status_code=403, detail="legal_final_status_requires_l4")

    review_required = is_final_status and requires_legal_review(row)
    override_used = False
    if review_required:
        trail = parse_json_list(row.get("audit_trail"))
        if has_independent_human_review(trail, actor_id):
            return review_required, override_used
        if role_level >= 5:
            if not req.override_reason.strip():
                raise HTTPException(status_code=409, detail="override_reason_required")
            override_used = True
        else:
            raise HTTPException(status_code=409, detail="legal_review_required")
    return review_required, override_used


# ═══════════════════════════════════════════════════════════════
# D-2-3  GET /scenarios
# ═══════════════════════════════════════════════════════════════

@router.get("/scenarios")
async def list_scenarios(user=Depends(get_current_user)):
    """로드된 시나리오 목록을 반환한다. (인증 필수)"""
    scenarios_dir = DATA_DIR / "scenarios"
    scenarios: list[dict] = []

    if scenarios_dir.exists():
        for f in sorted(scenarios_dir.glob("*.json")):
            try:
                data = json.loads(f.read_text(encoding="utf-8"))
                if isinstance(data, list):
                    scenarios.extend(d for d in data if isinstance(d, dict))
                elif isinstance(data, dict):
                    scenarios.append(data)
            except Exception:
                continue

    # scenario_id 가 비어있는 항목 보정
    cleaned = [s for s in scenarios if s.get("scenario_id") or s.get("id")]
    return {"scenarios": cleaned, "total": len(cleaned)}


# ═══════════════════════════════════════════════════════════════
# D-2-12  GET /facilities  ─ plants.json 기반 19개소
# ═══════════════════════════════════════════════════════════════

def _load_facilities() -> list[FacilityItem]:
    """data/facility_db/plants.json 의 plants + subsidiaries(국내/해외) 19개소를 통합."""
    path = DATA_DIR / "facility_db" / "plants.json"
    if not path.exists():
        return []

    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return []

    facilities: list[FacilityItem] = []

    def _coord(item: dict) -> dict:
        lat = item.get("lat")
        lng = item.get("lng")
        return {
            "lat": float(lat) if isinstance(lat, (int, float)) else None,
            "lng": float(lng) if isinstance(lng, (int, float)) else None,
        }

    # 자사 공장
    for p in data.get("plants", []):
        facilities.append(FacilityItem(
            plant_id=p.get("plant_id", ""),
            name=p.get("name", ""),
            location=p.get("location", ""),
            address=p.get("location", ""),
            certifications=p.get("certifications", []),
            processes=p.get("main_business", []) or p.get("main_processes", []),
            kind="plant",
            country="KR",
            **_coord(p),
        ))

    # 국내 계열사
    for s in data.get("subsidiaries_domestic", []):
        facilities.append(FacilityItem(
            plant_id=s.get("subsidiary_id") or s.get("plant_id", ""),
            name=s.get("name", ""),
            location=s.get("location", ""),
            address=s.get("location", ""),
            certifications=s.get("certifications", []),
            processes=s.get("main_business", []) or s.get("products", []),
            kind="subsidiary_domestic",
            country="KR",
            **_coord(s),
        ))

    # 해외 법인
    for s in data.get("subsidiaries_overseas", []):
        facilities.append(FacilityItem(
            plant_id=s.get("subsidiary_id") or s.get("id", ""),
            name=s.get("name", ""),
            location=s.get("location", "") or s.get("city", ""),
            address=s.get("location", "") or s.get("city", ""),
            certifications=s.get("certifications", []),
            processes=s.get("main_business", []) or s.get("products", []),
            kind="subsidiary_overseas",
            country=s.get("country", ""),
            **_coord(s),
        ))

    return facilities


@router.get("/facilities")
async def list_facilities(user=Depends(get_current_user)):
    """19개 사업장 (자사 + 국내 계열사 + 해외 법인) 통합 반환."""
    facilities = _load_facilities()
    return {
        "facilities": facilities,
        "total": len(facilities),
        "domestic": sum(1 for f in facilities if f.country == "KR"),
        "overseas": sum(1 for f in facilities if f.country and f.country != "KR"),
    }


# ═══════════════════════════════════════════════════════════════
# D-2-2.5  GET /impact-network  ─ D6 실 데이터 그래프
# ═══════════════════════════════════════════════════════════════

@router.get("/impact-network")
async def get_impact_network_graph(user=Depends(get_current_user)):
    """규제 ↔ 사업장 영향 네트워크 그래프 (정규화된 nodes/edges).

    plant_regulation_mapper.get_regulation_mapping_summary() 결과를
    `{nodes:[{id,type,label,meta}], edges:[{source,target}]}` 형태로 반환.
    프론트 SVG 가 슬라이스 없이 직접 렌더 가능한 형태.
    """
    from features.compliance.plant_regulation_mapper import (
        get_regulation_mapping_summary,
        _REGULATION_PLANT_RULES,
    )

    mapping = get_regulation_mapping_summary()
    nodes: list[dict[str, Any]] = []
    edges: list[dict[str, Any]] = []
    seen_plants: set[str] = set()

    for doc_type, plant_names in mapping.items():
        rule_meta = _REGULATION_PLANT_RULES.get(doc_type, {})
        nodes.append({
            "id": f"reg::{doc_type}",
            "type": "regulation",
            "label": doc_type,
            "meta": {
                "scope": rule_meta.get("scope", ""),
                "description": rule_meta.get("description", ""),
                "plant_count": len(plant_names),
            },
        })
        for plant_name in plant_names:
            plant_id = f"plant::{plant_name}"
            if plant_name not in seen_plants:
                seen_plants.add(plant_name)
                nodes.append({
                    "id": plant_id,
                    "type": "plant",
                    "label": plant_name,
                    "meta": {},
                })
            edges.append({
                "source": f"reg::{doc_type}",
                "target": plant_id,
            })

    return {
        "nodes": nodes,
        "edges": edges,
        "stats": {
            "regulation_count": sum(1 for n in nodes if n["type"] == "regulation"),
            "plant_count": sum(1 for n in nodes if n["type"] == "plant"),
            "edge_count": len(edges),
        },
    }


# ═══════════════════════════════════════════════════════════════
# D-2-2  GET /risk/scores  ─ 100점 스코어링
# ═══════════════════════════════════════════════════════════════

@router.get("/risk/scores", response_model=RiskScoreResponse)
async def list_risk_scores(user=Depends(get_current_user)):
    """모든 시나리오의 리스크 점수 (100점 + CRITICAL/HIGH/MEDIUM/LOW)."""
    from features.compliance.risk_scorer import score_all_scenarios, get_risk_summary

    try:
        scores = score_all_scenarios(str(DATA_DIR / "scenarios"))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"리스크 계산 실패: {e}")

    items = [
        RiskScoreItem(
            scenario_id=s.scenario_id,
            title=s.title,
            total_score=s.total_score,
            grade=s.grade,
            financial_impact=s.financial_impact,
            likelihood=s.likelihood,
            urgency=s.urgency,
            deadline=s.deadline,
            days_remaining=s.days_remaining,
            affected_plants=s.affected_plants,
            mitigation_status=s.mitigation_status,
        )
        for s in scores
    ]
    return RiskScoreResponse(
        total=len(items),
        summary=get_risk_summary(scores),
        scores=items,
    )


# ═══════════════════════════════════════════════════════════════
# D-2-5  POST /tariff/simulate
# ═══════════════════════════════════════════════════════════════

@router.post("/tariff/simulate", response_model=TariffSimulateResponse)
async def simulate_tariff_endpoint(
    req: TariffSimulateRequest,
    user=Depends(get_current_user),
):
    """관세 시뮬레이션 — 6품목(EWP/CCH/OBC/볼시트/도어/EV배터리) 기본 + 환율 적용."""
    from features.compliance.tariff_simulator import simulate_tariff

    out = simulate_tariff(tariff_rate=req.tariff_rate, exchange_rate=req.exchange_rate)
    return TariffSimulateResponse(
        tariff_rate=out["tariff_rate"],
        exchange_rate=out["exchange_rate"],
        total_annual_usd=out["total_annual_usd"],
        total_annual_krw=out["total_annual_krw"],
        total_annual_krw_billion=round(out["total_annual_krw"] / 1e8, 2),
        avg_cost_increase=out["avg_cost_increase"],
        results=[
            {
                "product": r.product,
                "tariff_rate": r.tariff_rate,
                "unit_tariff": r.unit_tariff,
                "annual_tariff": r.annual_tariff,
                "annual_tariff_krw": r.annual_tariff_krw,
                "cost_increase_pct": r.cost_increase_pct,
            }
            for r in out["results"]
        ],
    )


# ═══════════════════════════════════════════════════════════════
# D-2-4  GET /timeline  ─ Plotly Figure JSON
# ═══════════════════════════════════════════════════════════════

@router.get("/timeline", response_model=PlotlyResponse)
async def get_timeline(user=Depends(get_current_user)):
    """데드라인 간트 차트 (Plotly Figure JSON)."""
    from features.compliance.risk_scorer import score_all_scenarios
    from features.compliance.timeline_builder import build_deadline_timeline

    try:
        scores = score_all_scenarios(str(DATA_DIR / "scenarios"))
        fig = build_deadline_timeline(scores)
        return PlotlyResponse(figure=fig.to_plotly_json())
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"타임라인 생성 실패: {e}")


# ═══════════════════════════════════════════════════════════════
# D-2-7  GET /network/{scenario_id}  ─ Plotly Network
# ═══════════════════════════════════════════════════════════════

@router.get("/network/{scenario_id}", response_model=PlotlyResponse)
async def get_impact_network(scenario_id: str, user=Depends(get_current_user)):
    """규제 → 시설 → 부서 영향 네트워크 (Plotly Figure JSON)."""
    from features.compliance.impact_network import build_impact_network

    scenarios_dir = DATA_DIR / "scenarios"
    if not scenarios_dir.exists():
        raise HTTPException(status_code=404, detail="시나리오 디렉토리 없음")

    target: dict | None = None
    for f in scenarios_dir.glob("*.json"):
        try:
            d = json.loads(f.read_text(encoding="utf-8"))
            cands = d if isinstance(d, list) else [d]
            for c in cands:
                if isinstance(c, dict) and (
                    c.get("scenario_id") == scenario_id or c.get("id") == scenario_id
                ):
                    target = c
                    break
            if target:
                break
        except Exception:
            continue

    if not target:
        raise HTTPException(status_code=404, detail=f"시나리오 {scenario_id} 없음")

    try:
        fig = build_impact_network(target)
        return PlotlyResponse(figure=fig.to_plotly_json())
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"네트워크 생성 실패: {e}")


# ═══════════════════════════════════════════════════════════════
# D-2-6  GET /changes/recent  +  POST /changes/{id}/acknowledge
# ═══════════════════════════════════════════════════════════════

@router.get("/changes/recent", response_model=ChangeListResponse)
async def get_recent_changes_endpoint(
    limit: int = 20,
    unack_only: bool = False,
    regulation_type: str = "",
    item_id: str = "",
    user=Depends(get_current_user),
):
    """최근 규제 변경 이력 (compliance.db 기반).

    Issue 3: regulation_type + item_id 필터로 특정 규제의 변경 이력 + enrichment
    (before/after/diff/impact) 를 반환한다.
    """
    from features.compliance.change_detector import (
        init_change_db,
        get_recent_changes,
        get_change_stats,
    )

    init_change_db()
    raw = get_recent_changes(limit=limit, unacknowledged_only=unack_only)

    # regulation_type / item_id 필터 (선택)
    if regulation_type:
        raw = [c for c in raw if c.get("regulation_type") == regulation_type]
    if item_id:
        raw = [c for c in raw if c.get("item_id") == item_id]

    items = [
        ChangeItem(
            id=int(c.get("id", 0) or 0),
            regulation_type=c.get("regulation_type", ""),
            change_type=c.get("change_type", ""),
            item_id=c.get("item_id", ""),
            title=c.get("title", "") or c.get("description", ""),
            summary=c.get("summary", "") or c.get("description", ""),
            detected_at=c.get("detected_at", ""),
            acknowledged=bool(c.get("acknowledged", 0)),
            before_text=c.get("before_text", "") or "",
            after_text=c.get("after_text", "") or "",
            diff_html=c.get("diff_html", "") or "",
            impact_json=c.get("impact_json", "") or "",
        )
        for c in raw
    ]
    return ChangeListResponse(total=len(items), stats=get_change_stats(), changes=items)


@router.post("/changes/{change_id}/acknowledge", response_model=AcknowledgeResponse)
async def acknowledge_change_endpoint(
    change_id: int,
    user=Depends(require_role_level(3)),
):
    """변경 이력을 '확인 완료' 상태로 마킹한다. (role_level>=3)"""
    from features.compliance.change_detector import init_change_db, acknowledge_change

    init_change_db()
    try:
        acknowledge_change(change_id, user_id=getattr(user, "employee_id", "") or "")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"확인 처리 실패: {e}")
    return AcknowledgeResponse(ok=True, change_id=change_id)


# ═══════════════════════════════════════════════════════════════
# MVP — 변경 피드 + 워크플로우 + KPI (Stage 7 + 8c)
# ═══════════════════════════════════════════════════════════════


@router.get("/changes/feed", response_model=ChangeFeedResponse)
async def get_change_feed(
    grade: str | None = None,            # CRITICAL / HIGH / MEDIUM / LOW
    status: str | None = None,           # pending / reviewing / ... / filtered
    dept: str | None = None,             # 영향 부서 필터 (substring 매칭)
    legal_class: str | None = None,      # P1 D1 — criminal/administrative/civil/contract/standardization
    since: str | None = None,            # ISO datetime — 이후 변경만
    include_filtered: bool = False,      # 노이즈 archive 포함 여부
    limit: int = 50,
    offset: int = 0,
    user=Depends(get_current_user),
):
    """변경 피드 — 필터 + 페이지네이션. ChangeItem 보다 풍부한 enrich 필드 포함."""
    from features.compliance.change_detector import init_change_db, CHANGE_DB_PATH
    import sqlite3 as _sql

    init_change_db()
    conn = _sql.connect(CHANGE_DB_PATH)
    conn.row_factory = _sql.Row

    where: list[str] = []
    params: list[Any] = []

    if not include_filtered:
        where.append("status != 'filtered'")
    if grade:
        where.append("grade = ?")
        params.append(grade.upper())
    if status:
        where.append("status = ?")
        params.append(status)
    if dept:
        where.append("affected_departments LIKE ?")
        params.append(f"%{dept}%")
    if legal_class:
        # legal_class column 은 JSON list — substring 매칭으로 충분 (소량 enum)
        where.append("legal_class LIKE ?")
        params.append(f"%{legal_class}%")
    if since:
        where.append("detected_at >= ?")
        params.append(since)

    where_sql = (" WHERE " + " AND ".join(where)) if where else ""
    total = conn.execute(
        f"SELECT COUNT(*) as n FROM regulation_changes{where_sql}", params
    ).fetchone()["n"]

    rows = conn.execute(
        f"""SELECT * FROM regulation_changes{where_sql}
            ORDER BY detected_at DESC LIMIT ? OFFSET ?""",
        [*params, max(1, min(limit, 200)), max(0, offset)],
    ).fetchall()
    conn.close()

    items: list[ChangeFeedItem] = []
    for r in rows:
        def _json_list(key: str) -> list:
            try:
                v = json.loads(r[key] or "[]")
                return v if isinstance(v, list) else []
            except (json.JSONDecodeError, TypeError):
                return []

        items.append(ChangeFeedItem(
            id=int(r["id"] or 0),
            regulation_type=r["regulation_type"] or "",
            change_type=r["change_type"] or "",
            item_id=r["item_id"] or "",
            item_title=r["item_title"] or "",
            summary_ko=r["summary_ko"] or "",
            grade=r["grade"] or "MEDIUM",
            severity=r["severity"] or "info",
            status=r["status"] or "pending",
            affected_departments=_json_list("affected_departments"),
            affected_plants=_json_list("affected_plants"),
            assigned_to=r["assigned_to"] or "",
            detected_at=r["detected_at"] or "",
            acknowledged=bool(r["acknowledged"] or 0),
            audit_trail=_json_list("audit_trail"),
            old_value=(r["old_value"] or "")[:1000],
            new_value=(r["new_value"] or "")[:1000],
            # P1 D1 — 법무 분류·벌칙
            legal_class=_json_list("legal_class"),
            penalty_extract=r["penalty_extract"] or "",
            penalty_severity_krw_mn=int(r["penalty_severity_krw_mn"] or 0),
        ))
    return ChangeFeedResponse(
        total=int(total),
        items=items,
        has_more=(offset + len(items)) < int(total),
    )


@router.post("/changes/{change_id}/transition", response_model=ChangeTransitionResponse)
async def transition_change_status(
    change_id: int,
    req: ChangeTransitionRequest,
    user=Depends(require_role_level(3)),
):
    """워크플로우 상태 전환 (pending → reviewing → planning → announced → done).

    audit_trail 에 자동 append (ts, user, action, from, to).
    P4 D14 — transition 후 위임 룰 평가 (사용자 transition 우선, 룰은 보조 — notify/ticket 만 발화).
    """
    from features.compliance.change_detector import update_change_status

    review_required, override_used = _enforce_change_transition_guard(change_id, req, user)
    user_id = _transition_actor_id(user)
    ok = update_change_status(change_id, req.new_status, user_id)
    if not ok:
        raise HTTPException(
            status_code=400,
            detail=f"전환 실패 — invalid status '{req.new_status}' 또는 change_id={change_id} 미존재",
        )

    # P4 D14 — 사용자 transition 후 룰 평가 (transition_to 액션은 자동 skip; notify/ticket 만 발화)
    try:
        from features.compliance.delegation_rules import evaluate as _eval_rules
        import sqlite3 as _sql
        from features.compliance.change_detector import CHANGE_DB_PATH
        c = _sql.connect(CHANGE_DB_PATH)
        c.row_factory = _sql.Row
        row = c.execute(
            """SELECT id, item_title, summary_ko, grade, status, affected_departments,
                      legal_class, penalty_severity_krw_mn
               FROM regulation_changes WHERE id = ?""",
            (change_id,),
        ).fetchone()
        c.close()
        if row is not None:
            _eval_rules(dict(row), dry_run=False, trigger="post_user_transition")
    except Exception as _e:
        logger.debug("delegation post-transition hook 실패 change_id=%s: %s", change_id, _e)

    # P5 §6 — Jira 양방향 sync: 우리 transition → 연결된 Jira issue transition
    try:
        from features.compliance import jira_sync
        if jira_sync.jira_enabled():
            from datetime import datetime as _dt
            import sqlite3 as _sql2
            from features.compliance.change_detector import CHANGE_DB_PATH as _CHDB
            c2 = _sql2.connect(_CHDB)
            c2.row_factory = _sql2.Row
            t_row = c2.execute(
                """SELECT id, external_id FROM collab_tickets
                   WHERE change_id = ? AND external_id != '' LIMIT 1""",
                (change_id,),
            ).fetchone()
            if t_row is not None:
                issue_key = t_row["external_id"]
                target_jira = jira_sync.STATUS_MAP.get(req.new_status)
                if target_jira:
                    out = jira_sync.transition_issue(issue_key, target_jira)
                    if out.get("ok"):
                        c2.execute(
                            "UPDATE collab_tickets SET jira_last_sync_at = ? WHERE id = ?",
                            (_dt.now().isoformat(), t_row["id"]),
                        )
                        c2.commit()
                    else:
                        logger.warning(
                            "Jira transition 실패 change_id=%s issue=%s reason=%s",
                            change_id, issue_key, out.get("error"),
                        )
            c2.close()
    except Exception as _e:
        logger.warning("Jira transition hook 예외 change_id=%s: %s",
                       change_id, type(_e).__name__)

    if req.review_note.strip():
        _append_change_audit_event(
            change_id,
            {
                "user": user_id,
                "action": "human_review",
                "note": req.review_note.strip()[:500],
            },
        )
    if override_used:
        _append_change_audit_event(
            change_id,
            {
                "user": user_id,
                "action": "legal_admin_override",
                "reason": req.override_reason.strip()[:500],
                "to": req.new_status,
            },
        )

    return ChangeTransitionResponse(
        ok=True,
        change_id=change_id,
        new_status=req.new_status,
        review_required=review_required,
        override_used=override_used,
    )


@router.get("/changes/kpi", response_model=ChangeKpiResponse)
async def get_change_kpi_endpoint(
    window_days: int = 30,
    user=Depends(get_current_user),
):
    """임원 KPI 1페이지 — 등급별 카운트, 미해결, 평균 처리시간, 트렌드.

    RBAC: 모든 인증 사용자 read 허용. 임원만 보이도록 하는 건 프런트 측에서 처리.
    """
    from features.compliance.change_detector import get_change_kpi

    kpi = get_change_kpi(window_days=max(1, min(window_days, 365)))
    return ChangeKpiResponse(**kpi)


# P3 D13 — 6개월 트렌드 KPI 확장
@router.get("/changes/extended-trend", response_model=ExtendedTrendResponse)
async def get_extended_trend_endpoint(
    window_days: int = 180,
    user=Depends(get_current_user),
):
    """확장 트렌드 — 90/180일 윈도우. 월별 등급 + 부서별 처리시간 + 법무 분포."""
    from features.compliance.change_detector import get_extended_trend

    out = get_extended_trend(window_days=max(30, min(window_days, 730)))
    return ExtendedTrendResponse(
        window_days=out["window_days"],
        monthly_grade_trend=out["monthly_grade_trend"],
        by_dept_handling_hours=[
            DeptHandlingTime(**d) for d in out["by_dept_handling_hours"]
        ],
        by_legal_class=out["by_legal_class"],
        total=out["total"],
    )


# P3 D12 — Feedback Loop 수동 적용 (cron 도 동일 함수 호출)


@router.post("/feedback-loop/apply", response_model=FeedbackLoopApplyResponse)
async def apply_feedback_loop(
    req: FeedbackLoopApplyRequest,
    user=Depends(get_current_user),
):
    """누적 수정 데이터 → 룰 dict / fewshot 자동 보강 (관리자, dry_run 기본)."""
    from features.compliance.feedback_loop import apply_aggregated_rules

    out = apply_aggregated_rules(
        window_days=max(7, min(req.window_days, 365)),
        min_occurrences=max(1, min(req.min_occurrences, 100)),
        dry_run=req.dry_run,
    )
    return FeedbackLoopApplyResponse(**out)


# P3 D9 — 협업 티켓 (다중 부서 매핑)


@router.post("/changes/{change_id}/tickets", response_model=CreateTicketResponse)
async def create_collab_ticket(
    change_id: int,
    user=Depends(get_current_user),
):
    """다중 부서 영향 변경 → 협업 티켓 자동 생성 (Slack DM + 옵션 webhook)."""
    from features.compliance.change_detector import init_change_db, CHANGE_DB_PATH
    from features.compliance.collab_ticket import create_ticket
    import sqlite3 as _sql

    init_change_db()
    conn = _sql.connect(CHANGE_DB_PATH)
    conn.row_factory = _sql.Row
    row = conn.execute(
        """SELECT id, item_title, summary_ko, grade, affected_departments
           FROM regulation_changes WHERE id = ?""",
        (change_id,),
    ).fetchone()
    conn.close()
    if row is None:
        raise HTTPException(status_code=404, detail=f"change_id={change_id} 미존재")

    change = dict(row)
    try:
        change["affected_departments"] = json.loads(change["affected_departments"] or "[]")
    except (json.JSONDecodeError, TypeError):
        change["affected_departments"] = []

    out = create_ticket(change, change_id=change_id)
    if not out.get("ok"):
        return CreateTicketResponse(
            ok=False, ticket_id=0,
            error=out.get("error", "티켓 생성 실패"),
        )

    return CreateTicketResponse(
        ok=True,
        ticket_id=out["ticket_id"],
        departments=out["departments"],
        owners=[CollabTicketOwner(**o) for o in out["owners"]],
        slack_sent=out["slack_sent"],
        external_url=out.get("external_url", ""),
    )


@router.get("/tickets", response_model=CollabTicketListResponse)
async def get_tickets(
    status: str | None = None,
    department: str | None = None,
    limit: int = 50,
    user=Depends(get_current_user),
):
    """협업 티켓 목록 (필터: status, department)."""
    from features.compliance.collab_ticket import list_tickets

    from features.compliance.collab_ticket import ensure_kanban_columns
    ensure_kanban_columns()
    rows = list_tickets(status=status, department=department, limit=limit)
    items = [
        CollabTicketItem(
            id=int(r["id"]),
            change_id=int(r["change_id"] or 0),
            title=r.get("title") or "",
            departments=r.get("departments") or [],
            owners=r.get("owners") or [],
            status=r.get("status") or "created",
            external_id=r.get("external_id") or "",
            external_url=r.get("external_url") or "",
            created_at=r.get("created_at") or "",
            resolved_at=r.get("resolved_at") or "",
            assignee=r.get("assignee") or "",
            deadline=r.get("deadline") or "",
            progress_pct=int(r.get("progress_pct") or 0),
        )
        for r in rows
    ]
    return CollabTicketListResponse(items=items, total=len(items))


@router.post("/tickets/{ticket_id}/transition", response_model=TicketTransitionResponse)
async def transition_ticket(
    ticket_id: int,
    req: TicketTransitionRequest,
    user=Depends(get_current_user),
):
    """티켓 상태 전환 (created → acknowledged → in_progress → resolved)."""
    from features.compliance.collab_ticket import update_ticket_status

    user_id = getattr(user, "employee_id", "") or getattr(user, "email", "") or ""
    ok = update_ticket_status(ticket_id, req.new_status, user_id)
    if not ok:
        raise HTTPException(
            status_code=400,
            detail=f"전환 실패 — invalid status '{req.new_status}' 또는 ticket_id={ticket_id} 미존재",
        )
    return TicketTransitionResponse(ok=True, ticket_id=ticket_id, new_status=req.new_status)


# P3 D10 — What-if 시뮬레이션


@router.post("/whatif/simulate", response_model=WhatIfResponse)
async def whatif_simulate(
    req: WhatIfRequest,
    user=Depends(get_current_user),
):
    """5 시나리오 또는 자연어 입력 → 시뮬레이션 결과.

    natural_language 시나리오는 정규식 + LLM (Ollama) 으로 자동 라우팅.
    """
    from features.compliance.whatif_engine import simulate, VALID_SCENARIOS

    if req.scenario_type not in VALID_SCENARIOS:
        raise HTTPException(
            status_code=400,
            detail=f"scenario_type must be one of {sorted(VALID_SCENARIOS)}",
        )

    out = simulate(req.scenario_type, req.params, change_id=req.change_id)
    return WhatIfResponse(**out)


# P4 D17 — What-if 정밀화 (baseline + 회계 매핑)


@router.get("/whatif/baseline", response_model=WhatIfBaselineResponse)
async def get_whatif_baseline(
    corp_code: str | None = None,
    user=Depends(get_current_user),
):
    """현재 baseline (실 데이터 우선, 부재 시 hardcoded confidence=0.30)."""
    from features.compliance.financial_baseline import get_baseline

    baseline = get_baseline(corp_code=corp_code)
    return WhatIfBaselineResponse(**baseline.to_dict())


@router.post("/whatif/simulate/accounting", response_model=WhatIfAccountingResponse)
async def whatif_to_accounting(
    req: WhatIfAccountingRequest,
    user=Depends(require_role_level(3)),
):
    """시뮬 결과 → P&L line item 매핑 (재무팀 검토용, role_level≥3)."""
    from features.compliance.accounting_trace import map_to_pnl

    items = map_to_pnl(req.whatif_result)
    return WhatIfAccountingResponse(items=[PnlImpactItem(**i) for i in items])


# P5 §6 — Jira 양방향 sync


@router.get("/jira/health", response_model=JiraHealthResponse)
async def get_jira_health(user=Depends(require_role_level(3))):
    """Jira 자격 + 토큰 유효성 + default project 존재 확인 (role_level≥3)."""
    from features.compliance import jira_sync
    return JiraHealthResponse(**jira_sync.health())


@router.post("/jira/webhook", response_model=JiraWebhookAckResponse)
async def receive_jira_webhook(req: JiraWebhookEvent):
    """Jira → 우리 시스템 inbound webhook.

    payload: {webhookEvent: 'jira:issue_updated', issue: {key, fields: {status: {name}}}}
    JIRA_WEBHOOK_SECRET 가 설정돼 있으면 X-Atlassian-Webhook-Identifier 헤더 HMAC 검증
    (현재는 secret 없으면 통과 — 운영 시 secret 필수 설정 권장).

    NOTE: webhook endpoint 는 인증 없이 외부에서 호출됨 — secret 검증으로 보호.
    """
    from features.compliance import jira_sync
    from features.compliance.collab_ticket import update_ticket_status
    import sqlite3 as _sql
    from features.compliance.change_detector import (
        CHANGE_DB_PATH, init_change_db, update_change_status,
    )
    from datetime import datetime as _dt

    # 1) issue_key + 우리 status 추출
    issue_key, our_status = jira_sync.parse_webhook_status(
        req.model_dump() if hasattr(req, "model_dump") else {},
    )
    if not issue_key:
        return JiraWebhookAckResponse(
            ok=True, note=f"unsupported_event:{req.webhookEvent}",
        )

    # 2) collab_ticket 조회 (external_id 매칭)
    init_change_db()
    conn = _sql.connect(CHANGE_DB_PATH)
    conn.row_factory = _sql.Row
    t_row = conn.execute(
        "SELECT id, change_id, status FROM collab_tickets WHERE external_id = ? LIMIT 1",
        (issue_key,),
    ).fetchone()
    if t_row is None:
        conn.close()
        return JiraWebhookAckResponse(
            ok=True, issue_key=issue_key,
            note="no_matching_ticket",
        )

    if not our_status:
        conn.close()
        status_name = ((req.issue or {}).get("fields") or {}).get("status", {}).get("name", "")
        return JiraWebhookAckResponse(
            ok=True, issue_key=issue_key,
            synced_ticket_id=int(t_row["id"]),
            note=f"unmapped_status:{status_name}",
        )

    # 3) 우리 ticket / change 상태 동기화 (ticket → change 둘 다)
    ticket_id = int(t_row["id"])
    change_id = int(t_row["change_id"] or 0)
    # collab_tickets 의 valid status 와 우리 our_status 매핑 — pending/reviewing/done →
    # ticket status 는 created/acknowledged/in_progress/resolved 라 별도 매핑 필요.
    ticket_status_map = {
        "pending": "created",
        "reviewing": "in_progress",
        "done": "resolved",
    }
    ticket_status = ticket_status_map.get(our_status)
    if ticket_status:
        update_ticket_status(ticket_id, ticket_status, user_id=f"jira_webhook:{issue_key}")
    if change_id:
        try:
            update_change_status(change_id, our_status, user_id=f"jira_webhook:{issue_key}")
        except Exception as _e:
            logger.warning("Jira webhook change update 실패 change=%s: %s",
                           change_id, type(_e).__name__)

    conn.execute(
        "UPDATE collab_tickets SET jira_last_sync_at = ? WHERE id = ?",
        (_dt.now().isoformat(), ticket_id),
    )
    conn.commit()
    conn.close()

    return JiraWebhookAckResponse(
        ok=True,
        synced_ticket_id=ticket_id,
        issue_key=issue_key,
        new_status=our_status,
        note="synced",
    )


# P5 §10 — 2차 협력사 자동 발굴 (DART 공시 데이터 기반 후보 제안)


@router.get("/suppliers/discovery/candidates", response_model=SupplierDiscoveryListResponse)
async def list_supplier_discovery_candidates(
    limit: int = 50,
    min_filings: int = 1,
    user=Depends(require_role_level(3)),
):
    """DART 공시 동종업체 ⊖ 우리 suppliers DB → 미등록 후보 list (role_level≥3)."""
    from features.compliance.supplier_discovery import discover_candidates
    cands = discover_candidates(
        limit=max(1, min(int(limit), 200)),
        min_filings=max(1, int(min_filings)),
    )
    note = ""
    if not cands:
        note = "industry_filings 데이터 없음 — POST /industry-trend/fetch 먼저 실행"
    return SupplierDiscoveryListResponse(
        candidates=[SupplierDiscoveryCandidate(**c) for c in cands],
        total=len(cands),
        note=note,
    )


@router.post(
    "/suppliers/discovery/{corp_code}/promote",
    response_model=SupplierPromoteResponse,
)
async def promote_discovery_candidate(
    corp_code: str,
    req: SupplierPromoteRequest,
    user=Depends(require_role_level(4)),
):
    """후보 corp_code 를 실 supplier 로 등록 (role_level≥4 admin only)."""
    from features.compliance.supplier_discovery import promote_to_supplier
    out = promote_to_supplier(
        corp_code,
        name_override=req.name_override,
        tier=max(1, min(int(req.tier), 5)),
        relation_type=req.relation_type,
        parent_supplier_id=req.parent_supplier_id,
    )
    if not out.get("ok"):
        return SupplierPromoteResponse(ok=False, error=out.get("error", ""))
    return SupplierPromoteResponse(
        ok=True,
        supplier_id=out["supplier_id"],
        name=out["name"],
        tier_depth=int(out["tier_depth"]),
    )


# P5 §7 — 자체 결재 워크플로


@router.post("/approvals", response_model=ApprovalChainCreateResponse)
async def start_approval_chain(
    req: ApprovalChainCreateRequest,
    user=Depends(get_current_user),
):
    """변경 1건에 대한 다단 결재 chain 시작."""
    from features.compliance.approval_workflow import start_chain
    requested_by = _user_id(user)
    if not requested_by:
        raise HTTPException(status_code=400, detail="employee_id 또는 email 필요")
    if not req.steps:
        raise HTTPException(status_code=400, detail="steps 비어있음")
    try:
        cid = start_chain(
            change_id=req.change_id,
            name=req.name[:200],
            requested_by=requested_by,
            steps=[
                {"approver_id": s.approver_id, "role_label": s.role_label}
                for s in req.steps
            ],
        )
    except ValueError as e:
        return ApprovalChainCreateResponse(ok=False, error=str(e))
    return ApprovalChainCreateResponse(ok=True, chain_id=cid)


@router.get("/approvals/my", response_model=ApprovalMyPendingResponse)
async def get_my_pending_approvals(user=Depends(get_current_user)):
    """본인이 결재할 차례인 step list."""
    from features.compliance.approval_workflow import my_pending_steps
    me = _user_id(user)
    if not me:
        return ApprovalMyPendingResponse(items=[], total=0)
    rows = my_pending_steps(me)
    return ApprovalMyPendingResponse(items=rows, total=len(rows))


@router.get("/approvals/{chain_id}", response_model=ApprovalChainDetailResponse)
async def get_approval_chain(
    chain_id: int,
    user=Depends(get_current_user),
):
    """결재 chain 상세 (요청자 / step 별 결정 history)."""
    from features.compliance.approval_workflow import get_chain_detail
    detail = get_chain_detail(chain_id)
    if detail is None:
        raise HTTPException(status_code=404, detail=f"chain_id={chain_id} 미존재")
    return ApprovalChainDetailResponse(
        id=int(detail["id"]),
        change_id=int(detail["change_id"] or 0),
        name=detail["name"],
        requested_by=detail.get("requested_by", ""),
        status=detail.get("status", "pending"),
        current_step=int(detail.get("current_step") or 1),
        created_at=detail.get("created_at", ""),
        completed_at=detail.get("completed_at", ""),
        steps=[ApprovalStepItem(**s) for s in detail["steps"]],
    )


@router.post(
    "/approvals/steps/{step_id}/decide", response_model=ApprovalDecideResponse,
)
async def decide_approval_step(
    step_id: int,
    req: ApprovalDecideRequest,
    user=Depends(get_current_user),
):
    """결재 step 에 결정 — 본인 (approver_id 일치) 만 가능."""
    from features.compliance.approval_workflow import act_on_step
    actor = _user_id(user)
    out = act_on_step(step_id, req.decision, actor=actor, comment=req.comment)
    if not out.get("ok"):
        return ApprovalDecideResponse(
            ok=False,
            chain_id=int(out.get("chain_id") or 0),
            step_order=int(out.get("step_order") or 0),
            error=out.get("error", ""),
        )
    return ApprovalDecideResponse(
        ok=True,
        chain_id=int(out["chain_id"]),
        step_order=int(out["step_order"]),
        new_chain_status=out["new_chain_status"],
        next_step_order=int(out["next_step_order"]),
    )


# P3 D11 — 산업 트렌드


@router.post("/industry-trend/fetch", response_model=IndustryFetchResponse)
async def industry_trend_fetch(
    days: int = 90,
    user=Depends(get_current_user),
):
    """동종업계 5개사 DART 분기보고서 fetch + 인덱싱 (cron, admin).

    DART_API_KEY 미설정 시 빈 응답 (graceful skip).
    """
    from features.compliance.industry_trend import fetch_dart_filings, index_filings
    filings = fetch_dart_filings(days=max(7, min(days, 365)))
    indexed = index_filings(filings)
    return IndustryFetchResponse(
        fetched=len(filings),
        indexed=indexed,
        note=(
            "DART_API_KEY 미설정"
            if not filings else
            f"{indexed}건 신규 적재 (총 {len(filings)} fetch, 중복 제외)"
        ),
    )


@router.get("/changes/{change_id}/industry-context", response_model=IndustryContextResponse)
async def get_industry_context(
    change_id: int,
    user=Depends(get_current_user),
):
    """변경 1건 → 산업 평균 대비 비교 + 동종업계 사례."""
    from features.compliance.change_detector import init_change_db, CHANGE_DB_PATH
    from features.compliance.industry_trend import compare_change_to_industry
    import sqlite3 as _sql

    init_change_db()
    conn = _sql.connect(CHANGE_DB_PATH)
    conn.row_factory = _sql.Row
    row = conn.execute(
        """SELECT id, item_title, summary_ko, new_value
           FROM regulation_changes WHERE id = ?""",
        (change_id,),
    ).fetchone()
    conn.close()
    if row is None:
        raise HTTPException(status_code=404, detail=f"change_id={change_id} 미존재")

    out = compare_change_to_industry(dict(row))
    return IndustryContextResponse(
        change_keywords=out["change_keywords"],
        matching_filings_count=out["matching_filings_count"],
        by_corp=[IndustryCorpFilings(**c) for c in out["by_corp"]],
        industry_average_filings=out["industry_average_filings"],
        verdict=out["verdict"],
        available=out["available"],
    )


# P4 D14 — 권한 위임 룰 엔진


@router.get("/delegation-rules", response_model=DelegationRuleListResponse)
async def list_delegation_rules(
    enabled_only: bool = False,
    user=Depends(require_role_level(3)),
):
    """위임 룰 목록 (role_level≥3)."""
    from features.compliance.delegation_rules import list_rules

    rows = list_rules(enabled_only=enabled_only)
    items = [DelegationRuleItem(**r) for r in rows]
    return DelegationRuleListResponse(items=items, total=len(items))


@router.post("/delegation-rules", response_model=DelegationRuleMutationResponse)
async def create_delegation_rule(
    req: DelegationRuleCreateRequest,
    user=Depends(require_role_level(4)),
):
    """위임 룰 생성 (role_level≥4)."""
    from features.compliance.delegation_rules import create_rule

    owner = getattr(user, "employee_id", "") or getattr(user, "email", "") or "admin"
    rid = create_rule(
        name=req.name,
        owner=owner,
        conditions=req.conditions or {},
        actions=req.actions or {},
        enabled=req.enabled,
        priority=req.priority,
    )
    return DelegationRuleMutationResponse(ok=True, rule_id=rid)


@router.patch("/delegation-rules/{rule_id}", response_model=DelegationRuleMutationResponse)
async def update_delegation_rule(
    rule_id: int,
    req: DelegationRuleUpdateRequest,
    user=Depends(require_role_level(4)),
):
    """위임 룰 수정 (role_level≥4)."""
    from features.compliance.delegation_rules import update_rule

    ok = update_rule(
        rule_id,
        name=req.name,
        enabled=req.enabled,
        conditions=req.conditions,
        actions=req.actions,
        priority=req.priority,
    )
    if not ok:
        raise HTTPException(status_code=404, detail=f"rule_id={rule_id} 미존재")
    return DelegationRuleMutationResponse(ok=True, rule_id=rule_id)


@router.delete("/delegation-rules/{rule_id}", response_model=DelegationRuleMutationResponse)
async def delete_delegation_rule(
    rule_id: int,
    user=Depends(require_role_level(4)),
):
    """위임 룰 삭제 (role_level≥4)."""
    from features.compliance.delegation_rules import delete_rule

    ok = delete_rule(rule_id)
    if not ok:
        raise HTTPException(status_code=404, detail=f"rule_id={rule_id} 미존재")
    return DelegationRuleMutationResponse(ok=True, rule_id=rule_id)


@router.post("/delegation-rules/dry-run", response_model=DelegationDryRunResponse)
async def dry_run_delegation_rules(
    limit: int = 50,
    user=Depends(require_role_level(3)),
):
    """최근 N개 변경에 dry-run 평가 (role_level≥3, 사이드이펙트 없음)."""
    from features.compliance.delegation_rules import dry_run_recent

    out = dry_run_recent(limit=max(1, min(limit, 200)))
    return DelegationDryRunResponse(
        scanned=out["scanned"],
        matched=out["matched"],
        by_rule={str(k): v for k, v in (out.get("by_rule") or {}).items()},
        matches=[DelegationDryRunMatch(**m) for m in out["matches"]],
    )


# P4 D15 — 신입 학습경로 큐레이션


def _user_id(user) -> str:
    return getattr(user, "employee_id", "") or getattr(user, "email", "") or ""


@router.post("/learning-path", response_model=LearningPathCreateResponse)
async def create_learning_path(
    req: LearningPathCreateRequest,
    user=Depends(get_current_user),
):
    """사수가 신입에게 학습경로 큐레이션."""
    from features.compliance.learning_path import curate_path

    owner = _user_id(user)
    if not owner:
        raise HTTPException(status_code=400, detail="사수 employee_id 필요")
    try:
        pid = curate_path(
            name=req.name,
            owner_employee_id=owner,
            assignee_employee_id=req.assignee_employee_id,
            week_split=req.week_split,
            min_quiz_score=max(0, min(req.min_quiz_score, 100)),
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    return LearningPathCreateResponse(ok=True, path_id=pid)


@router.get("/learning-path/my", response_model=LearningMyProgressResponse)
async def get_my_learning_progress(user=Depends(get_current_user)):
    """신입 본인 학습 진도."""
    from features.compliance.learning_path import my_progress

    me = _user_id(user)
    if not me:
        return LearningMyProgressResponse(paths=[])
    out = my_progress(me)
    return LearningMyProgressResponse(
        paths=[
            LearningPathDetail(
                id=p["id"],
                name=p["name"],
                owner_employee_id=p["owner_employee_id"],
                week_count=p["week_count"],
                curriculum=p["curriculum"],
                progress=[LearningProgressItem(**pr) for pr in p["progress"]],
                status=p["status"],
                created_at=p["created_at"],
            )
            for p in out["paths"]
        ]
    )


@router.get("/learning-path/mentor-queue", response_model=LearningMentorQueueResponse)
async def get_mentor_queue(user=Depends(get_current_user)):
    """사수 평가 대기 큐."""
    from features.compliance.learning_path import mentor_queue

    me = _user_id(user)
    if not me:
        return LearningMentorQueueResponse(items=[], total=0)
    rows = mentor_queue(me)
    items = [
        LearningMentorQueueItem(
            id=r["id"], path_id=r["path_id"],
            path_name=r.get("path_name") or "",
            assignee_employee_id=r.get("assignee_employee_id") or "",
            week=r["week"], change_id=r["change_id"],
            quiz_score=r.get("quiz_score") or -1,
            quiz_attempts=r.get("quiz_attempts") or 0,
            mentor_review=r.get("mentor_review") or "pending",
            updated_at=r.get("updated_at") or "",
        )
        for r in rows
    ]
    return LearningMentorQueueResponse(items=items, total=len(items))


@router.get("/learning-path/{progress_id}/quiz", response_model=LearningQuizPreviewResponse)
async def get_learning_quiz_preview(
    progress_id: int,
    user=Depends(get_current_user),
):
    """P4.1 §6 — 응시 전 question/choices 미리보기. answer_index 미반환, sideeffect 없음."""
    from features.compliance.learning_path import (
        get_path_owner, get_progress_path_id, get_quiz_preview, _conn,
    )
    me = _user_id(user)
    pid = get_progress_path_id(progress_id)
    if pid == 0:
        raise HTTPException(status_code=404, detail="progress 미존재")

    c = _conn()
    row = c.execute(
        "SELECT assignee_employee_id, owner_employee_id FROM learning_paths WHERE id = ?",
        (pid,),
    ).fetchone()
    c.close()
    if row is None or not me or (
        me != row["assignee_employee_id"] and me != row["owner_employee_id"]
    ):
        raise HTTPException(status_code=403, detail="본인 학습 경로의 퀴즈만 미리보기 가능")

    out = get_quiz_preview(progress_id)
    if not out.get("ok"):
        return LearningQuizPreviewResponse(ok=False, error=out.get("error", ""))
    return LearningQuizPreviewResponse(
        ok=True,
        progress_id=out["progress_id"],
        change_id=out["change_id"],
        question=out["question"],
        choices=out["choices"],
        generated_by=out["generated_by"],
    )


@router.post("/learning-path/{progress_id}/quiz", response_model=LearningQuizResponse)
async def take_learning_quiz(
    progress_id: int,
    req: LearningQuizRequest,
    user=Depends(get_current_user),
):
    """퀴즈 채점 — 신입 본인만 응시 가능."""
    from features.compliance.learning_path import (
        get_progress_path_id, get_path_owner, take_quiz,
    )
    me = _user_id(user)
    pid = get_progress_path_id(progress_id)
    if pid == 0:
        raise HTTPException(status_code=404, detail="progress 미존재")
    # path 의 assignee 가 본인이어야 (owner 도 시뮬 가능)
    from features.compliance.learning_path import _conn
    c = _conn()
    row = c.execute(
        "SELECT assignee_employee_id, owner_employee_id FROM learning_paths WHERE id = ?",
        (pid,),
    ).fetchone()
    c.close()
    # P4 D15 fix #3 — `not me` 가드: 인증 user 의 employee_id/email 이 비어있을 때 RBAC bypass 방지
    if row is None or not me or (me != row["assignee_employee_id"] and me != row["owner_employee_id"]):
        raise HTTPException(status_code=403, detail="본인 학습 경로의 퀴즈만 응시 가능")

    # P4.1 §4 — owner 응시는 mentor_dryrun_score 만 갱신 (assignee 점수 보호)
    is_owner = bool(me) and me == row["owner_employee_id"]
    out = take_quiz(progress_id, req.choice_index, as_owner=is_owner)
    if not out.get("ok"):
        return LearningQuizResponse(ok=False, error=out.get("error", ""))
    return LearningQuizResponse(
        ok=True,
        correct=bool(out["correct"]),
        score=int(out["score"]),
        answer_index=int(out["answer_index"]),
        question=out["question"],
    )


@router.post(
    "/learning-path/{progress_id}/request-review",
    response_model=LearningRequestReviewResponse,
)
async def request_learning_review(
    progress_id: int,
    user=Depends(get_current_user),
):
    """사수에게 평가 요청 — Slack DM."""
    from features.compliance.learning_path import request_mentor_review
    out = request_mentor_review(progress_id)
    if not out.get("ok"):
        return LearningRequestReviewResponse(ok=False, error=out.get("error", ""))
    return LearningRequestReviewResponse(ok=True, slack_sent=bool(out.get("slack_sent")))


@router.get("/learning-path/{path_id}/export.scorm.zip")
async def export_learning_path_scorm(
    path_id: int,
    user=Depends(get_current_user),
):
    """P5 §5 — 학습경로 SCORM 1.2 ZIP export (외부 LMS 이식용).

    path owner / assignee / 또는 admin (role_level≥4) 만 다운로드 가능.
    """
    from fastapi.responses import Response
    from features.compliance.learning_path import (
        _conn, get_path_owner,
    )
    from features.compliance.lms_export import export_scorm_package

    me = _user_id(user)
    owner = get_path_owner(path_id)
    if not owner:
        raise HTTPException(status_code=404, detail=f"path_id={path_id} 미존재")

    c = _conn()
    row = c.execute(
        "SELECT assignee_employee_id FROM learning_paths WHERE id = ?", (path_id,),
    ).fetchone()
    c.close()
    assignee = row["assignee_employee_id"] if row else ""
    role_level = int(getattr(user, "role_level", 0) or 0)
    if not me or (me != owner and me != assignee and role_level < 4):
        raise HTTPException(status_code=403, detail="본인 학습 경로 또는 admin 만 다운로드 가능")

    data = export_scorm_package(path_id)
    if data is None:
        raise HTTPException(status_code=404, detail="export 실패")
    return Response(
        content=data,
        media_type="application/zip",
        headers={
            "Content-Disposition": f'attachment; filename="learning_path_{path_id}.zip"',
        },
    )


@router.get("/learning-path/{path_id}/export.xapi.json")
async def export_learning_path_xapi(
    path_id: int,
    user=Depends(get_current_user),
):
    """P5 §5 — 학습경로 xAPI (Tin Can) statements export (LRS POST 가능 형식)."""
    from features.compliance.learning_path import _conn, get_path_owner
    from features.compliance.lms_export import export_xapi_statements

    me = _user_id(user)
    owner = get_path_owner(path_id)
    if not owner:
        raise HTTPException(status_code=404, detail=f"path_id={path_id} 미존재")
    c = _conn()
    row = c.execute(
        "SELECT assignee_employee_id FROM learning_paths WHERE id = ?", (path_id,),
    ).fetchone()
    c.close()
    assignee = row["assignee_employee_id"] if row else ""
    role_level = int(getattr(user, "role_level", 0) or 0)
    if not me or (me != owner and me != assignee and role_level < 4):
        raise HTTPException(status_code=403, detail="본인 학습 경로 또는 admin 만 export 가능")

    statements = export_xapi_statements(path_id)
    if statements is None:
        raise HTTPException(status_code=404, detail="export 실패")
    return {"statements": statements, "count": len(statements)}


@router.post("/learning-path/{progress_id}/review", response_model=LearningReviewResponse)
async def submit_learning_review(
    progress_id: int,
    req: LearningReviewRequest,
    user=Depends(get_current_user),
):
    """사수 평가 — pass / redo / pass_with_comment. path owner 만 가능."""
    from features.compliance.learning_path import (
        get_progress_path_id, get_path_owner, submit_review,
    )
    me = _user_id(user)
    pid = get_progress_path_id(progress_id)
    if pid == 0:
        raise HTTPException(status_code=404, detail="progress 미존재")
    owner = get_path_owner(pid)
    # P4 D15 fix #3 — `not me` 가드: me/owner 빈 값일 때 RBAC bypass 방지
    if not me or not owner or me != owner:
        raise HTTPException(status_code=403, detail="본인이 큐레이션한 학습 경로의 평가만 가능")

    out = submit_review(progress_id, req.verdict, comment=req.comment, reviewer=me)
    if not out.get("ok"):
        return LearningReviewResponse(ok=False, verdict="", error=out.get("error", ""))
    return LearningReviewResponse(ok=True, verdict=out["verdict"], reviewer=me)


# P1 D2 — 임원 보고서 자동 생성 (markdown / docx)
from fastapi.responses import Response as _FastAPIResponse


@router.get("/changes/exec-report")
async def get_exec_report(
    format: str = "markdown",       # markdown | docx_boon_bujang
    since: str | None = None,        # ISO date (default = 이번달 1일)
    until: str | None = None,        # ISO date (default = 오늘)
    user=Depends(get_current_user),
):
    """기간 내 변경 보고서 다운로드. markdown 또는 docx 양식 선택.

    빈 기간 → 안내 본문만 (HTTP 200, 파일은 정상 생성).
    """
    from features.compliance.exec_report import generate_report, VALID_FORMATS
    from datetime import date as _date, datetime as _dt

    if format not in VALID_FORMATS:
        raise HTTPException(
            status_code=400,
            detail=f"format 은 {sorted(VALID_FORMATS)} 중 하나",
        )

    try:
        since_d = _dt.strptime(since, "%Y-%m-%d").date() if since else None
    except ValueError:
        raise HTTPException(status_code=400, detail="since 가 ISO 날짜 형식 (YYYY-MM-DD) 이 아님")
    try:
        until_d = _dt.strptime(until, "%Y-%m-%d").date() if until else None
    except ValueError:
        raise HTTPException(status_code=400, detail="until 이 ISO 날짜 형식 (YYYY-MM-DD) 이 아님")

    body, content_type, filename = generate_report(
        format=format, since=since_d, until=until_d
    )
    return _FastAPIResponse(
        content=body,
        media_type=content_type,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# P2 D5 — Feedback Loop


@router.post("/changes/{change_id}/correct", response_model=ChangeCorrectionResponse)
async def correct_change(
    change_id: int,
    req: ChangeCorrectionRequest,
    user=Depends(get_current_user),
):
    """사용자 수정 이벤트 적재 + change row 의 해당 field 즉시 갱신.

    수정 이력은 `change_corrections` 테이블 + 기존 `audit_trail` 양쪽에 기록.
    """
    from features.compliance.feedback_loop import record_correction

    user_id = getattr(user, "employee_id", "") or getattr(user, "email", "") or ""
    user_role = (
        getattr(user, "role", "")
        or getattr(user, "position", "")
        or ""
    )

    out = record_correction(
        change_id=change_id,
        field=req.field,
        new_value=req.new_value,
        user_id=user_id,
        user_role=user_role,
        note=req.note,
    )
    if not out.get("ok"):
        raise HTTPException(status_code=400, detail=out.get("error") or "correction 실패")
    return ChangeCorrectionResponse(
        ok=True,
        correction_id=int(out.get("correction_id") or 0),
    )


@router.get("/changes/correction-stats", response_model=CorrectionStatsResponse)
async def get_correction_stats(
    window_days: int = 30,
    user=Depends(get_current_user),
):
    """누적 수정 통계 — 정확도 metric + 부서 매핑 보강 후보."""
    from features.compliance.feedback_loop import aggregate_corrections

    stats = aggregate_corrections(window_days=max(1, min(window_days, 365)))
    # tuple → list 변환 (Pydantic 직렬화)
    stats["frequent_dept_changes"] = [list(t) for t in stats.get("frequent_dept_changes", [])]
    return CorrectionStatsResponse(**stats)


# P2 D8 — 외부 판례 RAG


@router.post("/case-law/index", response_model=CaseLawIndexResponse)
async def index_case_law(
    keywords: list[str] | None = None,
    display_per_keyword: int = 10,
    user=Depends(get_current_user),
):
    """외부 판례 corpus 인덱싱 (관리자, cron). LAW_GO_KR_OC 미설정 시 빈 응답."""
    from features.compliance.case_law_indexer import (
        DEFAULT_KEYWORDS,
        fetch_cases_for_keyword,
        index_cases,
    )

    target_keywords = keywords if keywords else list(DEFAULT_KEYWORDS)
    by_keyword: dict[str, int] = {}
    for kw in target_keywords:
        cases = fetch_cases_for_keyword(kw, display=max(1, min(display_per_keyword, 50)))
        by_keyword[kw] = index_cases(cases)
    return CaseLawIndexResponse(
        by_keyword=by_keyword,
        total=sum(by_keyword.values()),
    )


@router.get("/changes/{change_id}/similar-cases", response_model=SimilarCasesResponse)
async def get_similar_cases(
    change_id: int,
    top_k: int = 3,
    user=Depends(get_current_user),
):
    """변경 1건 → 유사 판례 top-k. 외부 자료 미가용 시 안전 응답."""
    from features.compliance.change_detector import init_change_db, CHANGE_DB_PATH
    from features.compliance.case_law_indexer import find_similar_for_change, collection_stats
    import sqlite3 as _sql

    init_change_db()
    conn = _sql.connect(CHANGE_DB_PATH)
    conn.row_factory = _sql.Row
    row = conn.execute(
        """SELECT id, item_title, summary_ko, new_value, legal_class
           FROM regulation_changes WHERE id = ?""",
        (change_id,),
    ).fetchone()
    conn.close()
    if row is None:
        raise HTTPException(status_code=404, detail=f"change_id={change_id} 미존재")

    # ChromaDB 가용성 체크 — 미가용 시 안전 응답
    stats = collection_stats()
    if not stats.get("available") or stats.get("count", 0) == 0:
        return SimilarCasesResponse(
            items=[],
            available=False,
            note="외부 판례 corpus 가 인덱싱되지 않았습니다. 관리자가 /case-law/index 를 실행해야 합니다.",
        )

    change = dict(row)
    hits = find_similar_for_change(change, top_k=max(1, min(top_k, 10)))
    return SimilarCasesResponse(
        items=[CaseLawItem(**h) for h in hits],
        available=True,
        note="" if hits else "유사도 임계값(0.7) 이상 매칭된 판례 없음",
    )


# P2 D7 — 계약 영향 분석


from fastapi import UploadFile, File, Form
import tempfile as _tempfile
import shutil as _shutil


@router.post("/contracts/upload", response_model=ContractIngestResponse)
async def upload_contract(
    file: UploadFile = File(...),
    contract_id: str = Form(...),
    counterparty: str = Form(""),
    contract_type: str = Form("OEM"),
    effective_date: str = Form(""),
    expiry_date: str = Form(""),
    annual_value_krw_mn: int = Form(0),
    user=Depends(get_current_user),
):
    """계약 PDF/docx/txt 업로드 → 조항 split → DB + ChromaDB 적재.

    파일은 임시 디렉토리에 저장 후 처리. 영구 보관은 storage rule 별도 설정 필요.
    """
    from features.compliance.contract_indexer import ingest_contract

    suffix = Path(file.filename or "contract").suffix or ".txt"
    with _tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        _shutil.copyfileobj(file.file, tmp)
        tmp_path = tmp.name

    out = ingest_contract(
        file_path=tmp_path,
        contract_id=contract_id,
        counterparty=counterparty,
        contract_type=contract_type,
        effective_date=effective_date,
        expiry_date=expiry_date,
        annual_value_krw_mn=annual_value_krw_mn,
    )
    if not out.get("ok"):
        return ContractIngestResponse(
            ok=False, contract_id=contract_id, clause_count=0,
            error=out.get("error", "ingest 실패"),
        )
    return ContractIngestResponse(
        ok=True,
        contract_id=out["contract_id"],
        clause_count=out["clause_count"],
    )


@router.get("/contracts", response_model=ContractListResponse)
async def list_contracts_endpoint(
    search: str = "",
    limit: int = 50,
    user=Depends(get_current_user),
):
    """계약 목록 (검색·필터)."""
    from features.compliance.contract_indexer import list_contracts

    rows = list_contracts(limit=max(1, min(limit, 500)), search=search)
    items = [
        ContractMeta(
            contract_id=str(r.get("contract_id") or ""),
            counterparty=str(r.get("counterparty") or ""),
            type=str(r.get("type") or ""),
            effective_date=str(r.get("effective_date") or ""),
            expiry_date=str(r.get("expiry_date") or ""),
            annual_value_krw_mn=int(r.get("annual_value_krw_mn") or 0),
            status=str(r.get("status") or "active"),
        )
        for r in rows
    ]
    return ContractListResponse(items=items, total=len(items))


@router.get("/changes/{change_id}/affected-contracts", response_model=AffectedContractsResponse)
async def get_affected_contracts(
    change_id: int,
    top_k: int = 5,
    user=Depends(get_current_user),
):
    """변경 1건 → 영향 계약·조항 매핑."""
    from features.compliance.change_detector import init_change_db, CHANGE_DB_PATH
    from features.compliance.contract_indexer import match_contracts
    import sqlite3 as _sql

    init_change_db()
    conn = _sql.connect(CHANGE_DB_PATH)
    conn.row_factory = _sql.Row
    row = conn.execute(
        """SELECT id, item_title, summary_ko, new_value, legal_class
           FROM regulation_changes WHERE id = ?""",
        (change_id,),
    ).fetchone()
    conn.close()
    if row is None:
        raise HTTPException(status_code=404, detail=f"change_id={change_id} 미존재")

    hits = match_contracts(dict(row), top_k=max(1, min(top_k, 20)))
    items = [
        AffectedContractItem(
            contract_id=h.get("contract_id", ""),
            counterparty=h.get("counterparty", ""),
            type=h.get("type", ""),
            clause_no=h.get("clause_no", ""),
            title=h.get("title", ""),
            body_excerpt=h.get("body_excerpt", ""),
            similarity=h.get("similarity"),
            match_keywords=h.get("match_keywords") or [],
            source=h.get("source", ""),
        )
        for h in hits
    ]
    return AffectedContractsResponse(items=items)


# P2 D6 — 공급망 (8 endpoints)


@router.post("/admin/suppliers/import", response_model=SupplierImportResponse)
async def admin_import_suppliers(
    req: SupplierImportRequest,
    user=Depends(get_current_user),
):
    """CSV 텍스트 → suppliers 또는 supplier_components 일괄 적재 (관리자)."""
    from features.compliance.supplier_compliance import (
        import_suppliers_csv,
        import_components_csv,
    )

    if req.target == "components":
        out = import_components_csv(req.csv_text)
    else:
        out = import_suppliers_csv(req.csv_text)
    return SupplierImportResponse(
        imported=out["imported"],
        skipped=out["skipped"],
        errors=out["errors"][:20],
    )


@router.get("/suppliers", response_model=SupplierListResponse)
async def get_suppliers(
    tier: int | None = None,
    country: str | None = None,
    min_score: int = 0,
    limit: int = 100,
    user=Depends(get_current_user),
):
    """협력사 목록 (필터: tier / country / min_score)."""
    from features.compliance.supplier_compliance import list_suppliers

    rows = list_suppliers(tier=tier, country=country, min_score=min_score, limit=limit)
    items = [SupplierMeta(**{k: v for k, v in r.items() if k in SupplierMeta.model_fields}) for r in rows]
    return SupplierListResponse(items=items, total=len(items))


@router.get("/suppliers/{supplier_id}", response_model=SupplierDetailResponse)
async def get_supplier_detail(
    supplier_id: str,
    user=Depends(get_current_user),
):
    """협력사 1건 상세 + 부품 list."""
    from features.compliance.supplier_compliance import get_supplier

    s = get_supplier(supplier_id)
    if s is None:
        raise HTTPException(status_code=404, detail=f"supplier_id={supplier_id} 미존재")
    meta = SupplierMeta(**{k: v for k, v in s.items() if k in SupplierMeta.model_fields})
    components = [
        SupplierComponentItem(
            component_code=str(c.get("component_code") or ""),
            hs_code=str(c.get("hs_code") or ""),
            unit_price_krw=int(c.get("unit_price_krw") or 0),
            qty_per_year=int(c.get("qty_per_year") or 0),
        )
        for c in s.get("components", [])
    ]
    return SupplierDetailResponse(meta=meta, components=components)


@router.get("/changes/{change_id}/affected-suppliers", response_model=AffectedSuppliersResponse)
async def get_affected_suppliers(
    change_id: int,
    top_k: int = 20,
    max_depth: int = 1,
    relation_type: str | None = None,
    user=Depends(get_current_user),
):
    """변경 1건 → 영향 협력사 자동 매핑 (HS / 키워드 / 국가).

    P4 D16 — max_depth>1 이면 supplier_graph cascading 으로 2차/3차 협력사 포함.
    P4.1 §11 — `relation_type` 필터 (direct / sub_assembly / raw_material / logistics).
    """
    from features.compliance.change_detector import init_change_db, CHANGE_DB_PATH
    import sqlite3 as _sql

    init_change_db()
    conn = _sql.connect(CHANGE_DB_PATH)
    conn.row_factory = _sql.Row
    row = conn.execute(
        """SELECT id, item_title, summary_ko, new_value
           FROM regulation_changes WHERE id = ?""",
        (change_id,),
    ).fetchone()
    conn.close()
    if row is None:
        raise HTTPException(status_code=404, detail=f"change_id={change_id} 미존재")

    md = max(1, min(max_depth, 5))
    rt_filter = (relation_type or "").strip().lower() or None

    def _passes_rt(h: dict) -> bool:
        if rt_filter is None:
            return True
        return (h.get("relation_type") or "direct").lower() == rt_filter

    if md == 1:
        from features.compliance.supplier_compliance import match_suppliers
        hits = match_suppliers(dict(row), top_k=max(1, min(top_k, 50)))
        items = [
            AffectedSupplierItem(
                supplier_id=h["supplier_id"],
                name=h.get("name", ""),
                tier=int(h.get("tier") or 1),
                country=h.get("country", ""),
                contact_email=h.get("contact_email", ""),
                annual_volume_krw_mn=int(h.get("annual_volume_krw_mn") or 0),
                compliance_score=int(h.get("compliance_score") or 0),
                match_reasons=h.get("match_reasons") or [],
                relation_type=h.get("relation_type") or "direct",
            )
            for h in hits if _passes_rt(h)
        ]
        return AffectedSuppliersResponse(items=items)

    from features.compliance.supplier_graph import affected_suppliers_multi_tier
    hits = affected_suppliers_multi_tier(
        dict(row), max_depth=md, top_k=max(1, min(top_k, 50)),
    )
    items = [
        AffectedSupplierItem(
            supplier_id=h["supplier_id"],
            name=h.get("name", ""),
            tier=int(h.get("tier") or 1),
            country=h.get("country", ""),
            contact_email=h.get("contact_email", ""),
            annual_volume_krw_mn=int(h.get("annual_volume_krw_mn") or 0),
            compliance_score=int(h.get("compliance_score") or 0),
            match_reasons=h.get("match_reasons") or [],
            tier_depth=int(h.get("tier_depth") or 1),
            cascade_path=h.get("cascade_path") or [],
            depth_from_match=int(h.get("depth_from_match") or 0),
            relation_type=h.get("relation_type") or "direct",
            match_method=h.get("match_method") or "primary",
        )
        for h in hits if _passes_rt(h)
    ]
    return AffectedSuppliersResponse(items=items)


@router.get("/suppliers/{supplier_id}/graph", response_model=SupplierGraphResponse)
async def get_supplier_graph(
    supplier_id: str,
    direction: str = "down",
    max_depth: int = 3,
    user=Depends(get_current_user),
):
    """공급망 트리 트래버설 (down=자식 / up=부모 체인). P4 D16."""
    from features.compliance.supplier_graph import traverse, detect_cycles
    if direction not in ("down", "up"):
        raise HTTPException(status_code=400, detail="direction 은 'down' 또는 'up'")
    nodes = traverse(supplier_id, max_depth=max_depth, direction=direction)
    cycles = detect_cycles()
    return SupplierGraphResponse(
        origin=supplier_id,
        direction=direction,
        max_depth=max(1, min(max_depth, 5)),
        nodes=[SupplierGraphNode(**n) for n in nodes],
        cycles_detected=cycles,
    )


@router.get("/suppliers/cycles", response_model=SupplierCycleResponse)
async def get_supplier_cycles(user=Depends(get_current_user)):
    """전체 공급망 그래프 cycle 감지 (CSV import 검증용). P4 D16."""
    from features.compliance.supplier_graph import detect_cycles
    return SupplierCycleResponse(cycles=detect_cycles())


@router.post("/suppliers/{supplier_id}/send-assessment", response_model=SendAssessmentResponse)
async def send_supplier_assessment(
    supplier_id: str,
    req: SendAssessmentRequest,
    user=Depends(get_current_user),
):
    """협력사 자가진단 폼 자동 생성 + SMTP 발송 (또는 queued)."""
    from features.compliance.change_detector import init_change_db, CHANGE_DB_PATH
    from features.compliance.supplier_compliance import (
        get_supplier,
        generate_self_assessment_form,
        send_self_assessment_email,
    )
    import sqlite3 as _sql

    s = get_supplier(supplier_id)
    if s is None:
        raise HTTPException(status_code=404, detail=f"supplier_id={supplier_id} 미존재")

    init_change_db()
    conn = _sql.connect(CHANGE_DB_PATH)
    conn.row_factory = _sql.Row
    row = conn.execute(
        """SELECT id, item_title, summary_ko, legal_class
           FROM regulation_changes WHERE id = ?""",
        (req.change_id,),
    ).fetchone()
    conn.close()
    if row is None:
        raise HTTPException(status_code=404, detail=f"change_id={req.change_id} 미존재")

    body = generate_self_assessment_form(dict(row), s)
    out = send_self_assessment_email(req.change_id, supplier_id, body)
    if not out.get("ok"):
        return SendAssessmentResponse(
            ok=False, status="error", error=out.get("error", "발송 실패"),
        )
    return SendAssessmentResponse(
        ok=True,
        assessment_id=int(out.get("assessment_id") or 0),
        status=out.get("status", ""),
        sent_via_smtp=bool(out.get("sent_via_smtp", False)),
    )


@router.get("/changes/{change_id}/cost-simulation", response_model=CostSimulationResponse)
async def get_cost_simulation(
    change_id: int,
    scenario_rate_pct: float = 25.0,
    user=Depends(get_current_user),
):
    """변경 1건 → 관세 시뮬레이션 + 화학물질 대체 추정."""
    from features.compliance.change_detector import init_change_db, CHANGE_DB_PATH
    from features.compliance.cost_simulator import (
        simulate_tariff_impact,
        simulate_chemical_substitution,
    )
    import sqlite3 as _sql

    init_change_db()
    conn = _sql.connect(CHANGE_DB_PATH)
    conn.row_factory = _sql.Row
    row = conn.execute(
        """SELECT id, item_title, summary_ko, new_value
           FROM regulation_changes WHERE id = ?""",
        (change_id,),
    ).fetchone()
    conn.close()
    if row is None:
        raise HTTPException(status_code=404, detail=f"change_id={change_id} 미존재")

    change = dict(row)
    tariff = simulate_tariff_impact(change, scenario_rate_pct=scenario_rate_pct)
    chemical = simulate_chemical_substitution(change)

    by_supplier = [
        CostSimulationSupplierBreakdown(
            supplier_id=s["supplier_id"],
            name=s.get("name", ""),
            country=s.get("country", ""),
            baseline_krw_mn=s["baseline_krw_mn"],
            additional_tariff_krw_mn=s["additional_tariff_krw_mn"],
        )
        for s in tariff.get("by_supplier", [])
    ]

    return CostSimulationResponse(
        baseline_cost_krw_mn=tariff["baseline_cost_krw_mn"],
        new_cost_krw_mn=tariff["new_cost_krw_mn"],
        delta_krw_mn=tariff["delta_krw_mn"],
        delta_pct=tariff["delta_pct"],
        by_supplier=by_supplier,
        applicable_hs=tariff["applicable_hs"],
        scenario_rate_pct=scenario_rate_pct,
        chemical_substitution=chemical,
    )


@router.get("/suppliers/{supplier_id}/alternatives", response_model=AlternativesResponse)
async def get_supplier_alternatives(
    supplier_id: str,
    top_k: int = 3,
    user=Depends(get_current_user),
):
    """대체 협력사 추천 (다차원 점수)."""
    from features.compliance.supplier_recommender import recommend_alternatives

    items = recommend_alternatives(supplier_id, top_k=max(1, min(top_k, 10)))
    return AlternativesResponse(
        items=[AlternativeSupplierItem(**a) for a in items],
    )


# ═══════════════════════════════════════════════════════════════
# D-2-8  POST /classify  ─ TF-IDF + RF
# ═══════════════════════════════════════════════════════════════

@router.post("/classify", response_model=ClassifyResponse)
async def classify_regulation(req: ClassifyRequest, user=Depends(get_current_user)):
    """규제 텍스트의 리스크 레벨을 분류한다."""
    from features.compliance.alerts.legal_guard import COMPLIANCE_AI_DISCLAIMER
    from features.compliance.regulation_classifier import get_regulation_classifier

    if not req.text or not req.text.strip():
        raise HTTPException(status_code=400, detail="text 가 비어있습니다.")

    try:
        clf = get_regulation_classifier()
        out = clf.classify(req.text)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"분류 실패: {e}")

    return ClassifyResponse(
        severity=str(out.severity),
        confidence=float(out.confidence),
        all_scores={str(k): float(v) for k, v in out.all_scores.items()},
        related_departments=list(out.related_departments),
        affected_plants=list(out.affected_plants),
        risk_score=int(out.risk_score),
        recommended_actions=list(out.recommended_actions),
        response_deadline=out.response_deadline or "",
        disclaimer=COMPLIANCE_AI_DISCLAIMER,
    )


# ═══════════════════════════════════════════════════════════════
# D-2-1 / D-2-12  Crawler control
# ═══════════════════════════════════════════════════════════════

# 단일 크롤러 인스턴스 매핑 (지연 로딩)
_CRAWLER_KEYS = {
    "iso": ("features.compliance.iso_crawler", "ISOCrawler"),
    "apqp": ("features.compliance.apqp_crawler", "APQPCrawler"),
    "msds": ("features.compliance.msds_crawler", "MSDSCrawler"),
    "domestic_law": ("features.compliance.domestic_law_crawler", "DomesticLawCrawler"),
    "eu_regulation": ("features.compliance.eu_regulation_crawler", "EURegulationCrawler"),
    "oem_quality": ("features.compliance.oem_quality_crawler", "OEMQualityCrawler"),
    "carbon_esg": ("features.compliance.carbon_esg_crawler", "CarbonESGCrawler"),
    "ev_battery": ("features.compliance.ev_battery_crawler", "EVBatteryCrawler"),
    "global_trade": ("features.compliance.global_trade_crawler", "GlobalTradeCrawler"),
}


def _summarize_crawl(name: str, result: Any) -> CrawlRunResponse:
    """크롤러 결과(dataclass 또는 dict)를 공통 응답 스키마로 정규화.

    Phase 1 BaseCrawler 는 dict 를 반환하고, 기존 6개 크롤러는 dataclass 를
    반환한다 — 양쪽 모두 동일 인터페이스로 처리.
    """
    def _g(key: str, default: Any = "") -> Any:
        if isinstance(result, dict):
            return result.get(key, default)
        return getattr(result, key, default)

    return CrawlRunResponse(
        name=name,
        crawled_at=str(_g("crawled_at", "")),
        source=str(_g("source", "")),
        total_count=int(
            _g("total_count", 0)
            or _g("total_records", 0)
            or _g("total_phases", 0)
        ),
        updates_found=int(
            _g("updates_found", 0)
            or _g("updates_needed", 0)
            or _g("total_updates", 0)
            or _g("action_needed", 0)
        ),
        source_type=str(_g("source_type", "curated")),
        errors=list(_g("errors", []) or []),
    )


_CRAWLER_TIMEOUT_SEC = 25.0


def _resolve_crawler_classes() -> dict[str, type]:
    """모듈 import 시점에 한 번만 9 크롤러 클래스를 사전 로드.

    Python 의 module import lock 은 thread 동시 import 시 _DeadlockError 를
    유발하므로, run_in_executor 로 병렬 실행하기 전에 모든 모듈을 사전 로드한다.
    """
    out: dict[str, type] = {}
    for name, (mod_path, cls_name) in _CRAWLER_KEYS.items():
        try:
            mod = __import__(mod_path, fromlist=[cls_name])
            out[name] = getattr(mod, cls_name)
        except Exception as e:
            logger.warning("crawler 사전 로드 실패 — name=%s err=%s", name, e)
    return out


_CRAWLER_CLASSES: dict[str, type] = _resolve_crawler_classes()


async def _run_one_crawler(
    name: str,
    trigger_source: str = "api",
    user_id: str = "",
) -> CrawlRunResponse:
    """단일 크롤러를 thread pool 에서 실행 + per-crawler 25s timeout.

    F5 구조화 로깅 + F12 crawl_runs 감사 기록을 모두 포함한다.
    페르소나: 시니어/신입 모두 incident 회고 + 학습 자료로 활용.
    """
    from backend.services import crawl_audit

    cls = _CRAWLER_CLASSES.get(name)
    started_iso = __import__("datetime").datetime.now().isoformat(timespec="seconds")

    if cls is None:
        logger.error("crawler.run", extra={"crawler": name, "ok": False,
                                            "elapsed_ms": 0,
                                            "error": "preload_failed"})
        crawl_audit.record_run(
            crawler_name=name, started_at=started_iso,
            elapsed_ms=0, ok=False,
            errors="preload_failed", trigger_source=trigger_source,
            user_id=user_id,
        )
        return CrawlRunResponse(
            name=name, errors=[f"크롤러 사전 로드 실패: {name}"]
        )

    loop = asyncio.get_running_loop()
    started = loop.time()

    def _do() -> Any:
        inst = cls(DATA_DIR / "crawled")
        return inst.crawl()

    try:
        res = await asyncio.wait_for(
            loop.run_in_executor(None, _do),
            timeout=_CRAWLER_TIMEOUT_SEC,
        )
        elapsed_ms = int((loop.time() - started) * 1000)
        summary = _summarize_crawl(name, res)
        ok = not summary.errors
        logger.info(
            "crawler.run",
            extra={
                "crawler": name,
                "ok": ok,
                "elapsed_ms": elapsed_ms,
                "updates_found": summary.updates_found or 0,
                "errors_count": len(summary.errors or []),
            },
        )
        # F8/F9 호환: 크롤러가 결과에 http_meta dict 를 첨부하면 audit 에 그대로 적재.
        http_meta = getattr(res, "http_meta", None) or (
            res.get("http_meta") if isinstance(res, dict) else None
        ) or {}
        crawl_audit.record_run(
            crawler_name=name, started_at=started_iso,
            elapsed_ms=elapsed_ms, ok=ok,
            updates_found=summary.updates_found or 0,
            errors=list(summary.errors or []),
            trigger_source=trigger_source,
            http_status=http_meta.get("status"),
            http_etag=http_meta.get("etag", ""),
            http_last_modified=http_meta.get("last_modified", ""),
            user_id=user_id,
        )
        return summary
    except asyncio.TimeoutError:
        elapsed_ms = int((loop.time() - started) * 1000)
        logger.warning(
            "crawler.run",
            extra={"crawler": name, "ok": False,
                   "elapsed_ms": elapsed_ms, "error": "timeout"},
        )
        crawl_audit.record_run(
            crawler_name=name, started_at=started_iso,
            elapsed_ms=elapsed_ms, ok=False,
            errors="timeout", trigger_source=trigger_source,
            user_id=user_id,
        )
        return CrawlRunResponse(
            name=name,
            errors=[f"timeout: {_CRAWLER_TIMEOUT_SEC:.0f}s"],
        )
    except Exception as e:
        elapsed_ms = int((loop.time() - started) * 1000)
        logger.error(
            "crawler.run",
            extra={"crawler": name, "ok": False,
                   "elapsed_ms": elapsed_ms,
                   "error": f"{type(e).__name__}: {e}"},
        )
        crawl_audit.record_run(
            crawler_name=name, started_at=started_iso,
            elapsed_ms=elapsed_ms, ok=False,
            errors=f"{type(e).__name__}: {e}",
            trigger_source=trigger_source,
            user_id=user_id,
        )
        return CrawlRunResponse(name=name, errors=[f"{type(e).__name__}: {e}"])


def _user_id_str(user) -> str:
    return str(getattr(user, "user_id", "") or getattr(user, "username", "") or "")


@router.post("/crawl/run/{name}", response_model=CrawlRunResponse)
async def run_single_crawler(
    name: str,
    user=Depends(require_role_level(3)),
):
    """개별 크롤러 실행 — async wrapping + 25s timeout."""
    if name not in _CRAWLER_KEYS:
        raise HTTPException(
            status_code=400,
            detail=f"지원되지 않는 크롤러: {name}. 가능: {list(_CRAWLER_KEYS.keys())}",
        )
    return await _run_one_crawler(name, trigger_source="api_retry",
                                  user_id=_user_id_str(user))


@router.post("/crawl/run-all", response_model=CrawlRunAllResponse)
async def run_all_crawlers(user=Depends(require_role_level(3))):
    """9개 크롤러 병렬 실행 — asyncio.gather + per-crawler 25s timeout."""
    uid = _user_id_str(user)
    tasks = [_run_one_crawler(name, trigger_source="api_run_all", user_id=uid)
             for name in _CRAWLER_KEYS.keys()]
    results = await asyncio.gather(*tasks)
    out = {r.name: r for r in results}
    total_changes = sum((r.updates_found or 0) for r in results)
    return CrawlRunAllResponse(crawlers=out, total_changes=total_changes)


# F12 — 크롤 실행 이력 조회 (시니어 incident / 신입 학습자료)
@router.get("/crawl/history")
async def list_crawl_runs(
    limit: int = 50,
    crawler_name: str | None = None,
    only_failed: bool = False,
    since_iso: str | None = None,
    user=Depends(get_current_user),
):
    """`crawl_runs` 감사 테이블 조회. 필터: crawler_name, only_failed, since_iso(ISO 8601)."""
    from backend.services import crawl_audit
    runs = crawl_audit.list_runs(
        limit=limit, crawler_name=crawler_name,
        only_failed=only_failed, since_iso=since_iso,
    )
    return {"runs": runs, "total": len(runs)}


@router.get("/crawl/history/stats")
async def crawl_history_stats(user=Depends(get_current_user)):
    """지난 24시간 크롤 통계 — 일일 다이제스트 첨부용."""
    from backend.services import crawl_audit
    return crawl_audit.stats_24h()


# ═══════════════════════════════════════════════════════════════
# 기존: POST /check  (키워드 매핑)
# ═══════════════════════════════════════════════════════════════

@router.post("/check", response_model=ComplianceCheckResponse)
async def check_compliance(
    req: ComplianceCheckRequest,
    user=Depends(get_current_user),
    _perm=Depends(require_permission("compliance.run_analysis")),
):
    """규정 준수 검사를 수행한다. (관련 부서 EMPLOYEE+ 필요)"""
    from features.compliance.alerts.legal_guard import COMPLIANCE_AI_DISCLAIMER

    try:
        query = req.query.lower()
        standards: list[str] = []
        status = "확인 필요"

        keyword_map = {
            "iatf": ["IATF 16949"],
            "iso 14001": ["ISO 14001"],
            "iso 45001": ["ISO 45001"],
            "reach": ["EU REACH"],
            "rohs": ["EU RoHS"],
            "ppap": ["IATF 16949 - PPAP"],
            "fmea": ["IATF 16949 - FMEA"],
            "spc": ["IATF 16949 - SPC"],
            "msds": ["화학물질관리법", "산업안전보건법"],
        }

        for kw, stds in keyword_map.items():
            if kw in query:
                standards.extend(stds)

        if standards:
            status = "관련 규정 발견"

        return ComplianceCheckResponse(
            answer=f"'{req.query}'에 대한 규정 검사 결과입니다.",
            relevant_standards=list(set(standards)),
            compliance_status=status,
            source="rules",
            disclaimer=COMPLIANCE_AI_DISCLAIMER,
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ═══════════════════════════════════════════════════════════════
# v3.6 Phase 3 — GET /scenarios/{scenario_id}/detail
# 시나리오 원문 상세: 법규명·조항·시행일·변경 전/후 텍스트·체크리스트·근거 링크
# (시뮬레이션과 분리 — 분석 vs 레퍼런스)
# ═══════════════════════════════════════════════════════════════


def _load_scenario_raw(scenario_id: str) -> dict[str, Any] | None:
    """data/scenarios/*.json 에서 매칭 시나리오 raw dict 반환."""
    scenarios_dir = DATA_DIR / "scenarios"
    if not scenarios_dir.exists():
        return None
    for f in scenarios_dir.glob("*.json"):
        try:
            with open(f, encoding="utf-8") as fp:
                data = json.load(fp)
        except Exception:
            continue
        sid = data.get("scenario_id") or data.get("id") or f.stem
        if sid == scenario_id:
            return data
        # us_trade_regulations.json 처럼 여러 시나리오를 배열로 묶은 경우
        if isinstance(data.get("scenarios"), list):
            for s in data["scenarios"]:
                if isinstance(s, dict) and (s.get("scenario_id") == scenario_id or s.get("id") == scenario_id):
                    return s
    return None


def _days_until(date_str: str) -> int:
    """ISO 날짜 문자열 → 오늘까지의 남은 일수 (음수 = 지난)"""
    if not date_str:
        return 0
    try:
        from datetime import datetime
        d = datetime.fromisoformat(date_str.split("T")[0])
        delta = (d - datetime.now()).days
        return max(delta, 0)
    except Exception:
        return 0


@router.get(
    "/scenarios/{scenario_id}/detail",
    response_model=ScenarioDetailResponse,
)
async def get_scenario_detail(scenario_id: str, user=Depends(get_current_user)):
    """선택한 시나리오의 원문 상세 정보를 반환한다.

    시뮬레이션 라우트(`POST /scenarios/{id}/simulate`)와 별개:
    - 시뮬레이션 = 위험도 점수 + 영향 시설/부서 + 비용 (분석)
    - 상세    = 법규 원문 + 변경 전/후 + 체크리스트 + 근거 (레퍼런스)
    """
    raw = _load_scenario_raw(scenario_id)
    if raw is None:
        raise HTTPException(
            status_code=404,
            detail=f"시나리오 '{scenario_id}' 를 찾을 수 없습니다.",
        )

    # 법규 메타 (없으면 빈값)
    reg_raw = raw.get("regulation") or {}
    regulation = ScenarioRegulationMeta(
        name=str(reg_raw.get("name", "")),
        article=str(reg_raw.get("article", "")),
        authority=str(reg_raw.get("authority", reg_raw.get("issuer", ""))),
        category=str(reg_raw.get("category", "")),
    )

    # 변경 전/후
    change_detail = raw.get("change_detail") or {}
    before = None
    after = None
    if isinstance(change_detail.get("before"), dict):
        b = change_detail["before"]
        before = ScenarioChangeVersion(
            text=str(b.get("text", "")),
            effective_date=str(b.get("effective_date", "")),
            version=str(b.get("version", "")),
        )
    if isinstance(change_detail.get("after"), dict):
        a = change_detail["after"]
        after = ScenarioChangeVersion(
            text=str(a.get("text", "")),
            effective_date=str(a.get("effective_date", "")),
            version=str(a.get("version", "")),
        )

    # 참고 자료 — reference_url + 추가 references 배열
    refs: list[ScenarioReference] = []
    if raw.get("reference_url"):
        refs.append(
            ScenarioReference(
                title=f"{regulation.authority or '관련 기관'} 공식 자료",
                url=str(raw["reference_url"]),
            )
        )
    extra_refs = raw.get("references") or []
    for r in extra_refs:
        if isinstance(r, str):
            refs.append(ScenarioReference(title=r, url=r if r.startswith("http") else ""))
        elif isinstance(r, dict):
            refs.append(
                ScenarioReference(
                    title=str(r.get("title", r.get("name", "참고 자료"))),
                    url=str(r.get("url", "")),
                )
            )

    deadline = str(raw.get("deadline", ""))
    days_remaining = _days_until(deadline)

    return ScenarioDetailResponse(
        scenario_id=scenario_id,
        title=str(raw.get("title", scenario_id)),
        description=str(raw.get("description", "")),
        regulation=regulation,
        change_before=before,
        change_after=after,
        severity=str(raw.get("severity", "medium")).lower(),
        impact_areas=list(raw.get("impact_areas", [])),
        applicable_plants=list(raw.get("applicable_plants", [])),
        affected_facility_ids=list(raw.get("affected_facility_ids", [])),
        affected_process_types=list(raw.get("affected_process_types", [])),
        deadline=deadline,
        days_remaining=days_remaining,
        required_actions=list(raw.get("required_actions", [])),
        estimated_cost=str(raw.get("estimated_cost", "")),
        references=refs[:10],
        raw={k: v for k, v in raw.items() if k not in ("change_detail",)},  # change_detail 은 텍스트 큼
    )


# ═══════════════════════════════════════════════════════════════
# v3.6 Phase 2 — POST /scenarios/{scenario_id}/simulate
# 시나리오 통합 시뮬레이션: 위험도 + 영향 시설/부서 + 비용 추정 + 권장 액션
# ═══════════════════════════════════════════════════════════════


def _grade_to_category(grade: str) -> str:
    """위험도 등급 문자열 → UI 카테고리 매핑."""
    g = (grade or "").upper()
    if g in ("CRITICAL", "C", "S"):
        return "CRITICAL"
    if g in ("HIGH", "H", "A"):
        return "HIGH"
    if g in ("MEDIUM", "M", "B"):
        return "MEDIUM"
    return "LOW"


def _default_recommended_actions(category: str, deadline_days: int) -> list[str]:
    """카테고리별 기본 권장 액션 (시나리오 JSON 에 명시 없을 때 폴백)."""
    base = []
    if category == "CRITICAL":
        base = [
            f"D-{deadline_days} 내 영향 시설·부서 비상 점검",
            "관련 부서장 긴급 회의 소집 (24시간 내)",
            "법무팀·외부 전문가 자문 의뢰",
            "공급망·고객사 사전 통지 검토",
            "변경 이력 추적 및 증빙 문서 보관",
        ]
    elif category == "HIGH":
        base = [
            f"D-{deadline_days} 내 시정 계획 수립",
            "관련 부서 협업 회의 (1주일 내)",
            "변경된 기준에 따른 절차 개정",
            "교육 자료 업데이트",
        ]
    elif category == "MEDIUM":
        base = [
            "월간 점검 일정에 추가 확인",
            "관련 부서 공유 및 인지 확보",
            "필요시 절차 검토",
        ]
    else:  # LOW
        base = [
            "정기 점검 사항으로 반영",
            "관련 부서 공유",
        ]
    return base


@router.post(
    "/scenarios/{scenario_id}/simulate",
    response_model=ScenarioSimulateResponse,
)
async def simulate_scenario(
    scenario_id: str,
    req: ScenarioSimulateRequest = ScenarioSimulateRequest(),
    user=Depends(get_current_user),
):
    """선택한 시나리오의 통합 시뮬레이션 결과를 반환한다.

    조합:
      - 시나리오 메타 (제목·설명·deadline) — data/scenarios/*.json
      - 위험도 점수 — features.compliance.compliance_db
      - 영향 시설/부서 — features.compliance.impact_network
      - 비용 추정 — features.compliance.tariff_simulator (관세 카테고리만)
      - 권장 액션 — 시나리오 JSON 의 recommended_actions 또는 카테고리 기본값
      - 근거 링크 — 시나리오 JSON 의 references (있으면)
    """
    # ── 1) 시나리오 로드 ──
    # detail 라우트와 동일 helper 재활용 — us_trade_regulations.json 처럼
    # `{"scenarios": [...]}` 배열 형식까지 정상 처리한다 (이전 in-place 로직은
    # 단일 객체만 가정하여 US-TRADE-* 시나리오가 모두 404 였음).
    scenario = _load_scenario_raw(scenario_id)
    if scenario is None:
        raise HTTPException(
            status_code=404,
            detail=f"시나리오 '{scenario_id}' 를 찾을 수 없습니다.",
        )

    title = scenario.get("title") or scenario.get("name") or scenario_id
    description = (
        scenario.get("description")
        or (scenario.get("regulation") or {}).get("article", "")
        or ""
    )
    deadline_days = int(scenario.get("days_remaining", scenario.get("deadline_days", 0)) or 0)

    # ── 2) 위험도 점수 ──
    # 이전: 존재하지 않는 `compute_risk_scores` 를 import 해 ImportError → 무음 폴백 →
    # 모든 시나리오가 0/100 + grade=MEDIUM 으로 표시되던 버그.
    # 카드 리스트(/risk-scores)와 동일 스코어러를 재사용해 동일 점수 보장.
    risk = ScenarioSimRiskScore()
    try:
        from features.compliance.risk_scorer import calculate_risk_score  # type: ignore
        rs = calculate_risk_score(scenario)
        risk = ScenarioSimRiskScore(
            total=int(round(rs.total_score)),
            fin=int(round(rs.financial_impact)),
            pos=int(round(rs.likelihood)),
            urg=int(round(rs.urgency)),
        )
        grade = rs.grade or scenario.get("grade", "MEDIUM")
    except Exception:
        logger.exception("calculate_risk_score 실패: scenario_id=%s", scenario_id)
        grade = scenario.get("grade", "MEDIUM")

    category = _grade_to_category(grade)

    # ── 3) 영향 시설/부서 ──
    # us_trade_regulations.json 은 `applicable_plants` 키를 사용 — 누락 시 시뮬레이션
    # 모달의 "AFFECTED PLANTS" 가 빈 칸으로 보였던 것을 키 alias 로 보강.
    plants = list(
        scenario.get("affected_plants")
        or scenario.get("applicable_plants")
        or scenario.get("affected_facility_ids")
        or scenario.get("sites")
        or []
    )
    departments: list[str] = []
    try:
        from features.compliance.impact_network import REGULATION_DEPT_MAP  # type: ignore
        # 규제 키워드 → 부서 매핑
        text_pool = " ".join(
            [
                title,
                description,
                str(scenario.get("regulation", {})),
                " ".join(scenario.get("keywords", [])),
            ]
        )
        for kw, deps in REGULATION_DEPT_MAP.items():
            if kw in text_pool:
                departments.extend(deps)
        departments = list(dict.fromkeys(departments))  # 중복 제거 (순서 유지)
    except Exception:
        departments = scenario.get("affected_departments") or []

    # ── 4) 비용 추정 (관세 시나리오만) ──
    cost_estimate = 0.0
    cost_breakdown: list[dict[str, Any]] = []
    if "관세" in title or "tariff" in scenario_id.lower() or "trade" in (scenario.get("category", "")).lower():
        try:
            from features.compliance.tariff_simulator import simulate_tariff  # type: ignore
            tariff_rate = req.tariff_rate if req.tariff_rate is not None else 25.0
            exchange_rate = req.exchange_rate if req.exchange_rate is not None else 1380.0
            sim = simulate_tariff(tariff_rate=tariff_rate, exchange_rate=exchange_rate)
            cost_breakdown = sim.get("items", [])
            total_krw = sum(item.get("annual_tariff_krw", 0) for item in cost_breakdown)
            cost_estimate = round(total_krw / 1_000_000_000, 2)  # 10억 단위
        except Exception:
            pass

    impact = ScenarioSimImpact(
        plants=plants,
        departments=departments,
        cost_estimate_krw_bn=cost_estimate,
        cost_breakdown=cost_breakdown[:20],  # 최대 20개
    )

    # ── 5) 권장 액션 ──
    recommended = scenario.get("recommended_actions") or scenario.get("actions") or []
    if not recommended:
        recommended = _default_recommended_actions(category, deadline_days)

    # ── 6) 근거 링크 ──
    evidence: list[ScenarioSimEvidence] = []
    refs = scenario.get("references") or scenario.get("evidence_links") or []
    for ref in refs[:5]:
        if isinstance(ref, str):
            evidence.append(ScenarioSimEvidence(title=ref, url=ref if ref.startswith("http") else ""))
        elif isinstance(ref, dict):
            evidence.append(
                ScenarioSimEvidence(
                    title=ref.get("title", ref.get("name", "참고 자료")),
                    url=ref.get("url", ""),
                )
            )

    return ScenarioSimulateResponse(
        scenario_id=scenario_id,
        title=title,
        category=category,
        deadline_days=deadline_days,
        description=description,
        risk_score=risk,
        impact=impact,
        recommended_actions=recommended[:10],
        evidence_links=evidence,
    )


# ═══════════════════════════════════════════════════════════════
# v3.6 Phase 2 — GET /crawl/results, GET /crawl/results/{name}
# 크롤러 실행 결과 조회 (data/crawled/*.json)
# ═══════════════════════════════════════════════════════════════


# 크롤러 이름 → 실제 파일명 매핑 (이름이 다를 수 있음)
_CRAWLER_FILE_MAP = {
    "iso": "iso_standards.json",
    "apqp": "apqp_process.json",
    "msds": "msds_data.json",
    "domestic_law": "domestic_laws.json",
    "eu_regulation": "eu_regulations.json",
    "oem_quality": "oem_quality.json",
    "carbon_esg": "carbon_esg.json",
    "ev_battery": "ev_battery.json",
    "global_trade": "global_trade.json",
}

# 결과 JSON 안에서 항목 배열을 담는 필드 이름 후보 (크롤러마다 다름)
_ITEMS_FIELD_CANDIDATES = [
    "standards", "laws", "regulations", "items", "products",
    "data", "results", "records", "phases",
]


def _extract_items(payload: dict[str, Any]) -> list[dict[str, Any]]:
    """크롤러 JSON 의 메인 배열을 자동 탐지."""
    for field in _ITEMS_FIELD_CANDIDATES:
        if isinstance(payload.get(field), list):
            return payload[field]
    # 마지막 수단: dict 의 list 값 중 가장 큰 것
    largest: list[dict[str, Any]] = []
    for v in payload.values():
        if isinstance(v, list) and len(v) > len(largest):
            largest = v
    return largest


@router.get("/crawl/results", response_model=CrawlResultsListResponse)
async def list_crawl_results(user=Depends(get_current_user)):
    """모든 크롤러의 결과 메타데이터를 반환한다.

    실제 데이터는 data/crawled/{filename}.json 에 저장됨. 본 엔드포인트는
    각 파일을 열어 메타 (crawled_at, source, total_count 등) 만 추출.
    """
    crawled_dir = DATA_DIR / "crawled"
    out: list[CrawlResultMeta] = []

    for name, filename in _CRAWLER_FILE_MAP.items():
        path = crawled_dir / filename
        if not path.exists():
            continue
        try:
            with open(path, encoding="utf-8") as f:
                data = json.load(f)
            items = _extract_items(data)
            size_bytes = path.stat().st_size
            out.append(
                CrawlResultMeta(
                    name=name,
                    filename=filename,
                    crawled_at=str(data.get("crawled_at", "")),
                    source=str(data.get("source", "")),
                    total_count=int(data.get("total_count", len(items))),
                    updates_found=int(data.get("updates_found", 0)),
                    errors=list(data.get("errors", []))[:5],
                    size_bytes=size_bytes,
                )
            )
        except Exception as e:
            out.append(
                CrawlResultMeta(
                    name=name,
                    filename=filename,
                    crawled_at="",
                    source="",
                    total_count=0,
                    errors=[f"파일 읽기 실패: {e}"],
                )
            )

    return CrawlResultsListResponse(crawlers=out, total=len(out))


# ═══════════════════════════════════════════════════════════════
# v3.6 Phase 3 Item 2 — GET /crawl/results/{name}/download?format=...
# 5 포맷 다운로드: JSON · CSV · XLSX · DOCX · PDF
# ═══════════════════════════════════════════════════════════════


def _build_csv(items: list[dict[str, Any]]) -> bytes:
    """크롤러 항목 배열 → CSV. 모든 항목의 키 합집합을 컬럼으로 사용."""
    import csv
    from io import StringIO

    if not items:
        return b"title,url,summary\n"
    columns = ["title", "url", "summary"]
    extra_keys: list[str] = []
    for it in items:
        for k in it.keys():
            if k not in columns and k not in extra_keys and k != "extra":
                extra_keys.append(k)
    columns = columns + extra_keys[:20]  # 추가 컬럼 최대 20개

    out = StringIO()
    writer = csv.DictWriter(out, fieldnames=columns, extrasaction="ignore")
    writer.writeheader()
    for it in items:
        row = {col: str(it.get(col, "")).replace("\n", " ") for col in columns}
        writer.writerow(row)
    return ("﻿" + out.getvalue()).encode("utf-8")  # BOM (Excel 한글 호환)


def _build_xlsx(name: str, meta: dict[str, Any], items: list[dict[str, Any]]) -> bytes:
    """openpyxl 사용 — 헤더 굵게 + 자동 너비."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment
        from io import BytesIO
    except ImportError as e:
        raise HTTPException(status_code=503, detail=f"openpyxl 미설치: {e}")

    wb = Workbook()
    ws = wb.active
    ws.title = name[:30] or "crawl"

    # 메타 정보 (1-3행)
    ws.cell(row=1, column=1, value=f"크롤러: {name}")
    ws.cell(row=1, column=1).font = Font(bold=True, size=14)
    ws.cell(row=2, column=1, value=f"실행 시각: {meta.get('crawled_at', '')}")
    ws.cell(row=3, column=1, value=f"출처: {meta.get('source', '')}")

    # 헤더 (5행)
    columns = ["title", "url", "summary"]
    extra_keys: list[str] = []
    for it in items:
        for k in it.keys():
            if k not in columns and k not in extra_keys and k != "extra":
                extra_keys.append(k)
    columns = columns + extra_keys[:15]

    header_fill = PatternFill("solid", fgColor="D89400")
    for col_idx, col_name in enumerate(columns, start=1):
        c = ws.cell(row=5, column=col_idx, value=col_name)
        c.font = Font(bold=True, color="FFFFFF")
        c.fill = header_fill
        c.alignment = Alignment(horizontal="center")

    # 데이터 (6행~)
    for row_idx, it in enumerate(items, start=6):
        for col_idx, col_name in enumerate(columns, start=1):
            v = it.get(col_name, "")
            ws.cell(row=row_idx, column=col_idx, value=str(v)[:32000])  # XLSX 셀 한계

    # 자동 너비 (간단 추정)
    for col_idx, col_name in enumerate(columns, start=1):
        col_letter = ws.cell(row=5, column=col_idx).column_letter
        max_len = max([len(str(it.get(col_name, ""))) for it in items[:50]] + [len(col_name)])
        ws.column_dimensions[col_letter].width = min(max_len + 2, 50)

    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _build_docx(name: str, meta: dict[str, Any], items: list[dict[str, Any]]) -> bytes:
    """python-docx — 회사 양식 보고서."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from io import BytesIO
    except ImportError as e:
        raise HTTPException(status_code=503, detail=f"python-docx 미설치: {e}")

    doc = Document()
    title = doc.add_heading(f"크롤링 결과 보고서 — {name}", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0xD8, 0x94, 0x00)

    p = doc.add_paragraph()
    p.add_run(f"실행 시각: {meta.get('crawled_at', '')}\n").font.size = Pt(10)
    p.add_run(f"출처: {meta.get('source', '')}\n").font.size = Pt(10)
    p.add_run(f"항목 수: {len(items)}건").font.size = Pt(10)

    doc.add_paragraph()  # 빈 줄

    for i, it in enumerate(items, 1):
        h = doc.add_heading(f"{i}. {it.get('title', '(제목 없음)')}", level=2)
        for run in h.runs:
            run.font.size = Pt(13)
        if it.get("url"):
            url_p = doc.add_paragraph()
            url_run = url_p.add_run(f"🔗 {it.get('url')}")
            url_run.font.size = Pt(9)
            url_run.font.color.rgb = RGBColor(0x29, 0x80, 0xB9)
        if it.get("summary"):
            doc.add_paragraph(it.get("summary"))

    # 푸터
    doc.add_paragraph()
    foot = doc.add_paragraph()
    fr = foot.add_run("아진산업(주) | 본 보고서는 컴플라이언스 모니터링 시스템에서 자동 생성됨")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_pdf(name: str, meta: dict[str, Any], items: list[dict[str, Any]]) -> bytes:
    """fpdf2 — PDF 보고서 (NanumGothic 한글 폰트)."""
    try:
        from fpdf import FPDF
        from io import BytesIO
        import os.path
    except ImportError as e:
        raise HTTPException(status_code=503, detail=f"fpdf2 미설치: {e}")

    pdf = FPDF(format="A4")
    pdf.add_page()

    # 한글 폰트 — Dockerfile 의 fonts-nanum 패키지가 /usr/share/fonts/truetype/nanum/ 에 설치됨
    font_loaded = False
    for font_path in [
        "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
        "/usr/share/fonts/truetype/nanum/NanumGothicBold.ttf",
        "/Library/Fonts/AppleSDGothicNeo.ttc",  # macOS 개발 환경
    ]:
        if os.path.exists(font_path):
            try:
                pdf.add_font("Korean", style="", fname=font_path)
                pdf.add_font("Korean", style="B", fname=font_path)
                pdf.set_font("Korean", size=10)
                font_loaded = True
                break
            except Exception:
                continue
    if not font_loaded:
        pdf.set_font("helvetica", size=10)

    pdf.set_font_size(18)
    pdf.cell(0, 12, f"크롤링 결과 보고서 - {name}", ln=True)
    pdf.set_font_size(9)
    pdf.cell(0, 6, f"실행 시각: {meta.get('crawled_at', '')}", ln=True)
    pdf.cell(0, 6, f"출처: {meta.get('source', '')[:80]}", ln=True)
    pdf.cell(0, 6, f"항목 수: {len(items)}건", ln=True)
    pdf.ln(4)

    for i, it in enumerate(items, 1):
        if pdf.get_y() > 270:
            pdf.add_page()
        pdf.set_font_size(11)
        title = str(it.get("title", "(제목 없음)"))[:120]
        pdf.multi_cell(0, 6, f"{i}. {title}")
        pdf.set_font_size(8)
        if it.get("url"):
            pdf.multi_cell(0, 5, f"URL: {str(it['url'])[:200]}")
        if it.get("summary"):
            summary = str(it["summary"])[:600]
            pdf.multi_cell(0, 5, summary)
        pdf.ln(2)

    return bytes(pdf.output())


@router.get("/crawl/results/{name}/download")
async def download_crawl_result(
    name: str,
    format: str = "json",
    enriched: bool = False,
    user=Depends(get_current_user),
):
    """크롤러 결과를 5 포맷 중 하나로 다운로드.

    format: json | csv | xlsx | docx | pdf | report

    enriched=true (Issue 3): docx/pdf/report 포맷에서 regulation_changes
    테이블의 과거 vs 현재 diff + 아진산업 영향 분석(사업장/공정/인원/위험도)
    섹션을 항목별로 추가한다. JSON/CSV/XLSX 는 영향 없음.
    """
    from fastapi.responses import Response

    if name not in _CRAWLER_FILE_MAP:
        raise HTTPException(status_code=404, detail=f"크롤러 '{name}' 미등록.")

    filename = _CRAWLER_FILE_MAP[name]
    path = DATA_DIR / "crawled" / filename
    if not path.exists():
        raise HTTPException(
            status_code=404,
            detail=f"'{name}' 결과 파일이 없습니다. 'RUN ALL' 으로 먼저 실행하세요.",
        )

    try:
        with open(path, encoding="utf-8") as f:
            data = json.load(f)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"결과 파일 파싱 실패: {e}")

    raw_items = _extract_items(data)
    # 정규화 — title/url/summary 추출 + 나머지 보존
    items: list[dict[str, Any]] = []
    for raw in raw_items:
        if not isinstance(raw, dict):
            items.append({"title": str(raw), "url": "", "summary": ""})
            continue
        title = str(
            raw.get("title")
            or raw.get("name")
            or raw.get("standard")
            or raw.get("law_name")
            or raw.get("id")
            or "(제목 없음)"
        )
        url = str(raw.get("url", raw.get("link", raw.get("source_url", ""))))
        summary = str(
            raw.get("summary") or raw.get("description") or raw.get("content", "")
        )[:1000]
        # 추가 필드는 그대로 보존
        merged = {"title": title, "url": url, "summary": summary}
        for k, v in raw.items():
            if k not in merged:
                merged[k] = v
        items.append(merged)

    fmt = format.lower().strip()
    from datetime import datetime
    today = datetime.now().strftime("%Y%m%d")
    base_name = f"{name}_{today}"

    meta = {
        "crawled_at": data.get("crawled_at", ""),
        "source": data.get("source", ""),
    }

    if fmt == "json":
        return Response(
            content=json.dumps(data, ensure_ascii=False, indent=2),
            media_type="application/json; charset=utf-8",
            headers={"Content-Disposition": f"attachment; filename={base_name}.json"},
        )
    if fmt == "csv":
        return Response(
            content=_build_csv(items),
            media_type="text/csv; charset=utf-8",
            headers={"Content-Disposition": f"attachment; filename={base_name}.csv"},
        )
    if fmt == "xlsx":
        return Response(
            content=_build_xlsx(name, meta, items),
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f"attachment; filename={base_name}.xlsx"},
        )
    if fmt == "docx":
        if enriched:
            from features.compliance.learning.regulation_exporter import (
                export_regulations_docx,
            )
            blob = export_regulations_docx(
                items=items,
                doc_type=name,
                display_name=name.upper(),
                crawled_at=meta.get("crawled_at", ""),
                include_enrichment=True,
                regulation_type=name,
            )
            if blob is not None:
                return Response(
                    content=blob,
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={"Content-Disposition": f"attachment; filename={base_name}_enriched.docx"},
                )
            # enrichment 실패 시 기본 docx 로 폴백하지 않고 명확히 실패 보고
            raise HTTPException(status_code=500, detail="enriched DOCX 생성 실패")
        return Response(
            content=_build_docx(name, meta, items),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={base_name}.docx"},
        )
    if fmt == "pdf":
        return Response(
            content=_build_pdf(name, meta, items),
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={base_name}.pdf"},
        )
    # v3.6 Phase 4 — 6번째 포맷: 회사 양식 보고서 (DOCX, 표지+요약+본문+부록)
    if fmt == "report":
        author = (
            f"{user.get('username', '')} {user.get('position', '')}".strip()
            if isinstance(user, dict)
            else ""
        )
        return Response(
            content=_build_report_docx(name, meta, items, author=author),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={base_name}_report.docx"},
        )

    raise HTTPException(
        status_code=400,
        detail=f"지원하지 않는 포맷: '{fmt}'. (json|csv|xlsx|docx|pdf|report)",
    )


def _build_report_docx(
    name: str,
    meta: dict[str, Any],
    items: list[dict[str, Any]],
    author: str = "",
) -> bytes:
    """단일 크롤러 회사 양식 보고서 (DOCX, 표지+요약+본문+부록).

    Phase 4 6번째 포맷. 일반 _build_docx 보다 구조화된 양식.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from io import BytesIO
        from datetime import datetime
    except ImportError as e:
        raise HTTPException(status_code=503, detail=f"python-docx 미설치: {e}")

    doc = Document()
    today = datetime.now()
    today_str = today.strftime("%Y년 %m월 %d일")
    period_str = today.strftime("%Y년 %m월")

    # ── 표지 ──
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("아진산업(주)")
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0xD8, 0x94, 0x00)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AJIN INDUSTRIAL CO., LTD.")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x8A, 0x82, 0x76)

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("법규 모니터링 보고서")
    r.font.size = Pt(22)
    r.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"— {name.upper()} —")
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    for _ in range(8):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"작성일 · {today_str}")
    r.font.size = Pt(11)
    if author:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"작성자 · {author}")
        r.font.size = Pt(11)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"기간 · {period_str} 월간 보고")
    r.font.size = Pt(11)

    doc.add_page_break()

    # ── Executive Summary ──
    doc.add_heading("1. 개요 (Executive Summary)", level=1)
    p = doc.add_paragraph()
    p.add_run(f"본 보고서는 {name.upper()} 크롤러로 수집한 ").font.size = Pt(11)
    r = p.add_run(f"{len(items)}건")
    r.font.bold = True
    r.font.color.rgb = RGBColor(0xD8, 0x94, 0x00)
    p.add_run(f"의 규제 항목을 정리한 것입니다.").font.size = Pt(11)

    p = doc.add_paragraph()
    p.add_run(f"  • 출처: {meta.get('source', '—')}").font.size = Pt(10)
    p = doc.add_paragraph()
    p.add_run(f"  • 마지막 갱신: {meta.get('crawled_at', '—')}").font.size = Pt(10)
    p = doc.add_paragraph()
    p.add_run(f"  • 항목 수: {len(items)}건").font.size = Pt(10)

    doc.add_paragraph()

    # ── 본문 ──
    doc.add_heading("2. 수집 항목 상세", level=1)
    for i, it in enumerate(items, 1):
        h = doc.add_heading(f"2.{i} {it.get('title', '(제목 없음)')}", level=2)
        if it.get("url"):
            p = doc.add_paragraph()
            r = p.add_run(f"🔗 출처: {it.get('url')}")
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x29, 0x80, 0xB9)
        if it.get("summary"):
            p = doc.add_paragraph(it.get("summary"))
            p.paragraph_format.space_after = Pt(8)

    # ── 부록 ──
    doc.add_page_break()
    doc.add_heading("부록 A. 전체 항목 색인", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "#"
    hdr[1].text = "제목"
    hdr[2].text = "URL"
    for i, it in enumerate(items, 1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = str(it.get("title", ""))[:80]
        row[2].text = str(it.get("url", ""))[:100]

    # ── 푸터 ──
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "아진산업(주) · 본 보고서는 컴플라이언스 모니터링 시스템에서 자동 생성됨"
    )
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════
# v3.6 Phase 4 — GET /crawl/results/bulk-download?format=...
# 9개 크롤러 결과를 단일 파일로 묶음 (JSON / XLSX / DOCX / PDF / ZIP / report)
# ═══════════════════════════════════════════════════════════════


def _gather_all_crawler_results() -> list[dict[str, Any]]:
    """9개 크롤러 결과 + 메타 + 항목 정규화 리스트."""
    out: list[dict[str, Any]] = []
    for crawler_name, filename in _CRAWLER_FILE_MAP.items():
        path = DATA_DIR / "crawled" / filename
        if not path.exists():
            out.append(
                {
                    "name": crawler_name,
                    "filename": filename,
                    "crawled_at": "",
                    "source": "",
                    "items": [],
                    "raw": {},
                    "errors": ["파일 없음 — RUN ALL 필요"],
                }
            )
            continue
        try:
            with open(path, encoding="utf-8") as f:
                data = json.load(f)
            raw_items = _extract_items(data)
            items: list[dict[str, Any]] = []
            for raw in raw_items:
                if not isinstance(raw, dict):
                    items.append({"title": str(raw), "url": "", "summary": ""})
                    continue
                title = str(
                    raw.get("title")
                    or raw.get("name")
                    or raw.get("standard")
                    or raw.get("law_name")
                    or raw.get("id")
                    or "(제목 없음)"
                )
                url = str(raw.get("url", raw.get("link", raw.get("source_url", ""))))
                summary = str(
                    raw.get("summary") or raw.get("description") or raw.get("content", "")
                )[:1000]
                items.append({"title": title, "url": url, "summary": summary, **raw})
            out.append(
                {
                    "name": crawler_name,
                    "filename": filename,
                    "crawled_at": str(data.get("crawled_at", "")),
                    "source": str(data.get("source", "")),
                    "items": items,
                    "raw": data,
                    "errors": list(data.get("errors", [])),
                }
            )
        except Exception as e:
            out.append(
                {
                    "name": crawler_name,
                    "filename": filename,
                    "errors": [f"파싱 실패: {e}"],
                    "items": [],
                    "raw": {},
                }
            )
    return out


def _build_bulk_xlsx(crawlers: list[dict[str, Any]]) -> bytes:
    """9개 시트 + Summary 시트."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment
        from io import BytesIO
        from datetime import datetime
    except ImportError as e:
        raise HTTPException(status_code=503, detail=f"openpyxl 미설치: {e}")

    wb = Workbook()
    # Summary 시트 (첫 번째)
    summary = wb.active
    summary.title = "Summary"
    summary["A1"] = "법규 모니터링 통합 보고서"
    summary["A1"].font = Font(bold=True, size=16, color="D89400")
    summary["A2"] = f"생성: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    summary["A3"] = f"크롤러 수: {len(crawlers)}"

    headers = ["#", "크롤러", "파일명", "마지막 갱신", "항목 수", "출처"]
    for i, h in enumerate(headers, 1):
        c = summary.cell(row=5, column=i, value=h)
        c.font = Font(bold=True, color="FFFFFF")
        c.fill = PatternFill("solid", fgColor="D89400")
        c.alignment = Alignment(horizontal="center")

    for ridx, c in enumerate(crawlers, start=6):
        summary.cell(row=ridx, column=1, value=ridx - 5)
        summary.cell(row=ridx, column=2, value=c["name"])
        summary.cell(row=ridx, column=3, value=c["filename"])
        summary.cell(row=ridx, column=4, value=c["crawled_at"])
        summary.cell(row=ridx, column=5, value=len(c["items"]))
        summary.cell(row=ridx, column=6, value=c["source"][:80])

    # 각 크롤러별 시트
    for c in crawlers:
        ws = wb.create_sheet(title=c["name"][:30])
        ws.cell(row=1, column=1, value=f"크롤러: {c['name']}").font = Font(bold=True, size=14)
        ws.cell(row=2, column=1, value=f"갱신: {c['crawled_at']}")
        ws.cell(row=3, column=1, value=f"출처: {c['source']}")
        # 헤더
        cols = ["#", "title", "url", "summary"]
        for i, h in enumerate(cols, 1):
            cell = ws.cell(row=5, column=i, value=h)
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor="D89400")
        for ridx, it in enumerate(c["items"], start=6):
            ws.cell(row=ridx, column=1, value=ridx - 5)
            ws.cell(row=ridx, column=2, value=str(it.get("title", ""))[:200])
            ws.cell(row=ridx, column=3, value=str(it.get("url", ""))[:300])
            ws.cell(row=ridx, column=4, value=str(it.get("summary", ""))[:1000])
        # 자동 너비
        ws.column_dimensions["A"].width = 5
        ws.column_dimensions["B"].width = 40
        ws.column_dimensions["C"].width = 50
        ws.column_dimensions["D"].width = 60

    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _build_bulk_docx(
    crawlers: list[dict[str, Any]],
    author: str = "",
) -> bytes:
    """9개 통합 회사 양식 보고서 (표지+목차+요약+9섹션+부록)."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from io import BytesIO
        from datetime import datetime
    except ImportError as e:
        raise HTTPException(status_code=503, detail=f"python-docx 미설치: {e}")

    doc = Document()
    today = datetime.now()
    today_str = today.strftime("%Y년 %m월 %d일")
    period_str = today.strftime("%Y년 %m월")
    total_items = sum(len(c["items"]) for c in crawlers)
    successful = sum(1 for c in crawlers if not c.get("errors"))

    # 표지
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("아진산업(주)")
    r.font.size = Pt(30)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0xD8, 0x94, 0x00)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AJIN INDUSTRIAL CO., LTD.")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x8A, 0x82, 0x76)

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("법규 모니터링 통합 보고서")
    r.font.size = Pt(24)
    r.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"— {len(crawlers)}개 크롤러 통합 —")
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    for _ in range(8):
        doc.add_paragraph()
    for line in [
        f"작성일 · {today_str}",
        f"작성자 · {author}" if author else "",
        f"기간 · {period_str} 월간 보고",
    ]:
        if not line:
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.size = Pt(11)

    doc.add_page_break()

    # 목차
    doc.add_heading("목차", level=1)
    doc.add_paragraph("1. 개요 (Executive Summary)")
    doc.add_paragraph("2. 크롤러별 상세")
    for i, c in enumerate(crawlers, 1):
        doc.add_paragraph(f"   2.{i} {c['name']} — {len(c['items'])}건")
    doc.add_paragraph("3. 부록 — 전체 항목 색인")
    doc.add_page_break()

    # Executive Summary
    doc.add_heading("1. 개요 (Executive Summary)", level=1)
    p = doc.add_paragraph()
    p.add_run("본 보고서는 ").font.size = Pt(11)
    r = p.add_run(f"{successful}/{len(crawlers)}개")
    r.font.bold = True
    r.font.color.rgb = RGBColor(0xD8, 0x94, 0x00)
    p.add_run(f" 크롤러에서 수집한 ").font.size = Pt(11)
    r = p.add_run(f"총 {total_items}건")
    r.font.bold = True
    r.font.color.rgb = RGBColor(0xD8, 0x94, 0x00)
    p.add_run(f"의 규제 항목을 통합 정리한 것입니다.").font.size = Pt(11)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "크롤러"
    hdr[1].text = "항목 수"
    hdr[2].text = "마지막 갱신"
    hdr[3].text = "상태"
    for c in crawlers:
        row = table.add_row().cells
        row[0].text = c["name"]
        row[1].text = str(len(c["items"]))
        row[2].text = c["crawled_at"][:10] if c["crawled_at"] else "—"
        row[3].text = "정상" if not c.get("errors") else "에러"

    doc.add_page_break()

    # 본문 - 각 크롤러 섹션
    doc.add_heading("2. 크롤러별 상세", level=1)
    for i, c in enumerate(crawlers, 1):
        doc.add_heading(f"2.{i} {c['name']}", level=2)
        p = doc.add_paragraph()
        p.add_run(f"  • 출처: {c['source'] or '—'}\n").font.size = Pt(10)
        p.add_run(f"  • 갱신: {c['crawled_at'] or '—'}\n").font.size = Pt(10)
        p.add_run(f"  • 항목 수: {len(c['items'])}건").font.size = Pt(10)
        if c.get("errors"):
            p = doc.add_paragraph()
            r = p.add_run(f"⚠ 에러: {'; '.join(c['errors'])}")
            r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
            r.font.size = Pt(10)
        # 상위 5개 항목만 본문에 표시 (큰 보고서 방지)
        for j, it in enumerate(c["items"][:5], 1):
            p = doc.add_paragraph()
            r = p.add_run(f"    {j}. {it.get('title', '(제목 없음)')[:100]}")
            r.font.size = Pt(10)
            r.font.bold = True
            if it.get("url"):
                p = doc.add_paragraph()
                r = p.add_run(f"       🔗 {it.get('url')[:120]}")
                r.font.size = Pt(8)
                r.font.color.rgb = RGBColor(0x29, 0x80, 0xB9)
        if len(c["items"]) > 5:
            p = doc.add_paragraph()
            r = p.add_run(
                f"    ... 외 {len(c['items']) - 5}건 (전체 항목은 부록 참조)"
            )
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # 부록
    doc.add_page_break()
    doc.add_heading("3. 부록 — 전체 항목 색인", level=1)
    for i, c in enumerate(crawlers, 1):
        doc.add_heading(f"A.{i} {c['name']}", level=2)
        if not c["items"]:
            doc.add_paragraph("(항목 없음)")
            continue
        for j, it in enumerate(c["items"], 1):
            p = doc.add_paragraph()
            r = p.add_run(f"{j}. {it.get('title', '')[:120]}")
            r.font.size = Pt(9)
            if it.get("url"):
                p = doc.add_paragraph()
                r = p.add_run(f"   {it.get('url')[:150]}")
                r.font.size = Pt(8)
                r.font.color.rgb = RGBColor(0x29, 0x80, 0xB9)

    # 푸터
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "아진산업(주) · 본 보고서는 컴플라이언스 모니터링 시스템에서 자동 생성됨"
    )
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_bulk_zip(crawlers: list[dict[str, Any]]) -> bytes:
    """9개 개별 JSON + index.txt → ZIP."""
    import zipfile
    from io import BytesIO

    buf = BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        index_lines = ["크롤러,파일명,갱신시각,항목수,출처"]
        for c in crawlers:
            zf.writestr(
                f"{c['name']}.json",
                json.dumps(c["raw"], ensure_ascii=False, indent=2),
            )
            index_lines.append(
                f"{c['name']},{c['filename']},{c['crawled_at']},{len(c['items'])},{c['source'][:80]}"
            )
        zf.writestr("index.csv", "\n".join(index_lines))
    return buf.getvalue()


@router.get("/crawl/results/bulk-download")
async def download_bulk_crawl_results(
    format: str = "report",
    user=Depends(get_current_user),
):
    """9개 크롤러 결과 통합 다운로드.

    format: report (회사 양식 DOCX, 기본) | docx | xlsx | pdf | json | zip
    """
    from fastapi.responses import Response
    from datetime import datetime

    crawlers = _gather_all_crawler_results()
    today = datetime.now().strftime("%Y%m%d")
    base = f"compliance_bulk_{today}"
    fmt = format.lower().strip()

    author = ""
    if isinstance(user, dict):
        username = user.get("username", "")
        position = user.get("position", "")
        author = f"{username} {position}".strip()

    if fmt == "json":
        payload = {
            "generated_at": datetime.now().isoformat(),
            "crawlers": {c["name"]: c["raw"] for c in crawlers},
            "total_items": sum(len(c["items"]) for c in crawlers),
        }
        return Response(
            content=json.dumps(payload, ensure_ascii=False, indent=2),
            media_type="application/json; charset=utf-8",
            headers={"Content-Disposition": f"attachment; filename={base}.json"},
        )
    if fmt == "xlsx":
        return Response(
            content=_build_bulk_xlsx(crawlers),
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f"attachment; filename={base}.xlsx"},
        )
    if fmt in ("docx", "report"):
        return Response(
            content=_build_bulk_docx(crawlers, author=author),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={base}_report.docx"},
        )
    if fmt == "pdf":
        # PDF는 단일 크롤러용 _build_pdf 를 9번 페이지로 합칠 수 있지만
        # 현재 fpdf2 기반으로 통합 보고서가 무거워서 첫 5개만 포함.
        # 실용적으로는 docx 가 더 좋은 선택. 단순 폴백.
        meta_combined = {
            "crawled_at": datetime.now().isoformat(),
            "source": f"{len(crawlers)}개 크롤러 통합",
        }
        all_items: list[dict[str, Any]] = []
        for c in crawlers:
            for it in c["items"][:10]:  # 각 크롤러 상위 10개
                all_items.append(
                    {
                        "title": f"[{c['name']}] {it.get('title', '')}",
                        "url": it.get("url", ""),
                        "summary": it.get("summary", ""),
                    }
                )
        return Response(
            content=_build_pdf("BULK", meta_combined, all_items),
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={base}.pdf"},
        )
    if fmt == "zip":
        return Response(
            content=_build_bulk_zip(crawlers),
            media_type="application/zip",
            headers={"Content-Disposition": f"attachment; filename={base}.zip"},
        )

    raise HTTPException(
        status_code=400,
        detail=f"지원하지 않는 포맷: '{fmt}'. (report|docx|xlsx|pdf|json|zip)",
    )


@router.get(
    "/crawl/results/{name}",
    response_model=CrawlResultDetailResponse,
)
async def get_crawl_result_detail(
    name: str,
    limit: int = 50,
    offset: int = 0,
    user=Depends(get_current_user),
):
    """특정 크롤러의 항목 리스트를 반환 (페이지네이션)."""
    if name not in _CRAWLER_FILE_MAP:
        raise HTTPException(status_code=404, detail=f"크롤러 '{name}' 가 등록되지 않았습니다.")

    filename = _CRAWLER_FILE_MAP[name]
    path = DATA_DIR / "crawled" / filename

    if not path.exists():
        # 실행 전 — 빈 결과 반환
        return CrawlResultDetailResponse(
            name=name,
            filename=filename,
            total=0,
            items=[],
            has_more=False,
        )

    try:
        with open(path, encoding="utf-8") as f:
            data = json.load(f)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"파일 파싱 실패: {e}")

    raw_items = _extract_items(data)
    lo = max(0, offset)
    hi = lo + max(1, min(limit, 200))
    page = raw_items[lo:hi]

    items: list[CrawlResultItem] = []
    for raw in page:
        if not isinstance(raw, dict):
            items.append(CrawlResultItem(title=str(raw), summary="", url=""))
            continue
        # Title — 한국어 변형 우선, MSDS 의 substance_name(_ko) 도 인식.
        # 9개 크롤러 JSON 전수 조사 결과: name_ko 가 6개, title_ko 1개,
        # substance_name_ko 1개에 있음. msds_data.json 은 chain 미스로 "(제목 없음)"
        # 폴백 발생하던 이슈 해결.
        title = str(
            raw.get("title_ko")
            or raw.get("name_ko")
            or raw.get("substance_name_ko")
            or raw.get("title")
            or raw.get("name")
            or raw.get("substance_name")
            or raw.get("standard")
            or raw.get("law_name")
            or raw.get("id")
            or "(제목 없음)"
        )
        url = str(raw.get("url", raw.get("link", raw.get("source_url", ""))))
        # Summary — ajin_relevance 가 가장 사용자 친화적 (사내 영향 분석, 6개 크롤러
        # 공통 한국어 필드). amendment_summary(국내법규)·changes_summary_ko(OEM/ISO) 차순위.
        summary = str(
            raw.get("summary_ko")
            or raw.get("ajin_relevance")
            or raw.get("amendment_summary")
            or raw.get("changes_summary_ko")
            or raw.get("summary")
            or raw.get("description")
            or raw.get("changes_summary")
            or raw.get("content", "")
        )[:500]  # 최대 500자
        # extra: 위에서 추출한 표준 필드 외 모든 필드 (중복 노출 방지)
        _extracted_keys = {
            "title", "title_ko", "name", "name_ko", "substance_name", "substance_name_ko",
            "standard", "law_name", "id",
            "url", "link", "source_url",
            "summary", "summary_ko", "description", "ajin_relevance",
            "amendment_summary", "changes_summary", "changes_summary_ko", "content",
        }
        extra = {k: v for k, v in raw.items() if k not in _extracted_keys}
        items.append(CrawlResultItem(title=title, url=url, summary=summary, extra=extra))

    return CrawlResultDetailResponse(
        name=name,
        filename=filename,
        crawled_at=str(data.get("crawled_at", "")),
        source=str(data.get("source", "")),
        total=len(raw_items),
        items=items,
        has_more=hi < len(raw_items),
    )


# ═══════════════════════════════════════════════════════════════
# D1 — 법규 전문 검색 (FTS5 + ChromaDB 하이브리드)
# ═══════════════════════════════════════════════════════════════


@router.get("/search")
async def compliance_search(
    q: str,
    limit: int = 20,
    offset: int = 0,
    doc_type: str | None = None,
    user=Depends(get_current_user),
):
    """하이브리드 검색 — FTS5 BM25 + ChromaDB 의미. 1,784 row 대상."""
    from backend.services.search.hybrid import search

    return search(q, limit=limit, offset=offset, doc_type=doc_type)


@router.get("/search/health")
async def compliance_search_health(user=Depends(get_current_user)):
    """검색 인덱스 상태 — fts5 row 수 + chromadb 컬렉션 상태."""
    from backend.services.search.hybrid import health
    from backend.services.search import fts_index

    info = health()
    info["counts"]["glossary"] = _glossary_count()
    info["fts_count"] = fts_index.index_count()
    return info


@router.get("/regulations/{reg_id}")
async def get_regulation_detail(reg_id: str, user=Depends(get_current_user)):
    """단일 규제 상세 (regulations 테이블 + content_json 평탄화)."""
    from backend.services.search import fts_index

    doc = fts_index.fetch_regulation(reg_id)
    if not doc:
        raise HTTPException(status_code=404, detail=f"규제 {reg_id} 없음")
    return doc


def _glossary_path() -> Path:
    return DATA_DIR / "glossary" / "terms.json"


def _glossary_count() -> int:
    p = _glossary_path()
    if not p.exists():
        return 0
    try:
        d = json.loads(p.read_text(encoding="utf-8"))
        terms = d.get("terms", []) if isinstance(d, dict) else d
        return len(terms) if isinstance(terms, list) else 0
    except Exception:
        return 0


# ═══════════════════════════════════════════════════════════════
# D7 — 법규 문서 라이브러리 (compliance_documents CRUD)
# ═══════════════════════════════════════════════════════════════


@router.get("/docs")
async def list_compliance_docs(
    limit: int = 100, offset: int = 0,
    doc_type: str | None = None,
    user=Depends(get_current_user),
):
    """법규 문서 라이브러리 — DB-backed 동적 목록."""
    from backend.services import docs_library
    docs_library.seed_if_empty()
    docs = docs_library.list_docs(limit=limit, offset=offset, doc_type=doc_type)
    return {"docs": docs, "total": len(docs)}


@router.get("/docs/{doc_id}")
async def get_compliance_doc(doc_id: int, user=Depends(get_current_user)):
    from backend.services import docs_library
    d = docs_library.get_doc(doc_id)
    if not d:
        raise HTTPException(status_code=404, detail=f"문서 {doc_id} 없음")
    return d


@router.post("/docs")
async def create_compliance_doc(payload: dict, user=Depends(get_current_user)):
    from backend.services import docs_library
    if not payload.get("title") or not payload.get("doc_type"):
        raise HTTPException(status_code=400, detail="title, doc_type 필수")
    new_id = docs_library.create_doc(
        payload, uploaded_by=str(getattr(user, "user_id", "") or getattr(user, "username", ""))
    )
    return {"id": new_id, "ok": True}


@router.delete("/docs/{doc_id}")
async def delete_compliance_doc(doc_id: int, user=Depends(get_current_user)):
    from backend.services import docs_library
    ok = docs_library.delete_doc(doc_id)
    if not ok:
        raise HTTPException(status_code=404, detail=f"문서 {doc_id} 없음")
    return {"id": doc_id, "deleted": True}


# ═══════════════════════════════════════════════════════════════
# D3 — Celery 스케줄러 트리거 (동기 fallback 포함)
# ═══════════════════════════════════════════════════════════════


# ═══════════════════════════════════════════════════════════════
# D5 — Kanban: collab_tickets PATCH (assignee/deadline/progress_pct)
# ═══════════════════════════════════════════════════════════════


# ═══════════════════════════════════════════════════════════════
# D4 — SOP diff (사내 SOP ↔ 규제 변경 차이 분석)
# ═══════════════════════════════════════════════════════════════


@router.get("/sop")
async def list_sops_endpoint(limit: int = 100, user=Depends(get_current_user)):
    from features.compliance.sop_diff import list_sops, ensure_table
    ensure_table()
    return {"sops": list_sops(limit=limit), "total": len(list_sops(limit=limit))}


@router.get("/sop/{sop_id}")
async def get_sop_endpoint(sop_id: int, user=Depends(get_current_user)):
    from features.compliance.sop_diff import get_sop
    s = get_sop(sop_id)
    if not s:
        raise HTTPException(status_code=404, detail=f"SOP {sop_id} 없음")
    return s


@router.post("/sop")
async def create_sop_endpoint(payload: dict, user=Depends(get_current_user)):
    from features.compliance.sop_diff import create_sop
    if not payload.get("title"):
        raise HTTPException(status_code=400, detail="title 필수")
    new_id = create_sop(
        payload,
        uploaded_by=str(getattr(user, "user_id", "") or getattr(user, "username", "")),
    )
    return {"id": new_id, "ok": True}


@router.post("/changes/{change_id}/sop-diff")
async def sop_diff_for_change(change_id: int, user=Depends(get_current_user)):
    """변경 1건 → 영향 SOP 목록 + diff 블록."""
    import sqlite3 as _sqlite3
    from features.compliance.change_detector import CHANGE_DB_PATH
    from features.compliance.sop_diff import diff_for_change

    conn = _sqlite3.connect(CHANGE_DB_PATH)
    conn.row_factory = _sqlite3.Row
    try:
        r = conn.execute(
            "SELECT * FROM regulation_changes WHERE id=?", (change_id,)
        ).fetchone()
    finally:
        conn.close()
    if not r:
        raise HTTPException(status_code=404, detail=f"변경 {change_id} 없음")

    return diff_for_change(change_id, dict(r))


@router.patch("/tickets/{ticket_id}")
async def patch_collab_ticket(ticket_id: int, payload: dict, user=Depends(get_current_user)):
    """티켓에 assignee/deadline/progress_pct 패치."""
    from features.compliance.collab_ticket import ensure_kanban_columns, patch_ticket
    ensure_kanban_columns()
    if "progress_pct" in payload:
        try:
            v = int(payload["progress_pct"])
            payload["progress_pct"] = max(0, min(100, v))
        except (TypeError, ValueError):
            payload.pop("progress_pct", None)
    ok = patch_ticket(ticket_id, payload)
    if not ok:
        raise HTTPException(status_code=404, detail=f"티켓 {ticket_id} 없음 또는 변경 사항 없음")
    return {"id": ticket_id, "patched": True}


@router.get("/scheduler/jobs")
async def list_scheduler_jobs(user=Depends(get_current_user)):
    """Celery beat 스케줄 + 잡 ID 목록."""
    from backend.services.jobs import list_jobs
    return {"jobs": list_jobs()}


@router.post("/scheduler/trigger/{job_id}")
async def trigger_scheduler_job(job_id: str, user=Depends(require_role_level(3))):
    """단일 잡 동기 트리거 — Celery 미설치 환경에서도 동작."""
    from backend.services.jobs import trigger_job
    try:
        result = trigger_job(job_id)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    return {"job_id": job_id, "result": result}


@router.post("/digest/run-now")
async def run_digest_now(user=Depends(require_role_level(3))):
    """일일 다이제스트 즉시 실행 (전체 사용자 발송)."""
    from features.compliance.alerts.legal_guard import COMPLIANCE_AI_DISCLAIMER
    from backend.services.jobs.digest import run
    result = run(force_all_users=True)
    return {**result, "disclaimer": COMPLIANCE_AI_DISCLAIMER}


@router.get("/glossary")
async def get_compliance_glossary(user=Depends(get_current_user)):
    """규제 용어 사전 (data/glossary/terms.json — 30+ 항목)."""
    p = _glossary_path()
    if not p.exists():
        return {"terms": [], "total": 0}
    try:
        d = json.loads(p.read_text(encoding="utf-8"))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"glossary 읽기 실패: {e}")
    if isinstance(d, dict):
        terms = d.get("terms", [])
        return {"terms": terms, "total": len(terms)}
    if isinstance(d, list):
        return {"terms": d, "total": len(d)}
    return {"terms": [], "total": 0}


# ═══════════════════════════════════════════════════════════
# v4.2 — D 컴플라이언스 실시간 알람 (2 endpoint)
# ═══════════════════════════════════════════════════════════


@router.get("/alarms/recent")
async def recent_compliance_alarms(
    since_ts: int = 0,
    limit: int = 50,
    scope: str = "all",
    user=Depends(get_current_user),
):
    """대시보드 D 알람 피드. SPC `/equipment/spc/violations/recent` 와 대칭.

    - since_ts: epoch sec — 이 시각 이후의 알람만 (0=전체)
    - limit: 최대 50건 권장
    - scope: 'all' (default, 전체) | 'mine' (사용자 부서 매칭, P5 페르소나 필터)
    """
    from backend.schemas.compliance_alarm import ComplianceAlarmsResponse
    from features.compliance.alarm_aggregator import collect_recent_alarms

    department = None
    if scope == "mine":
        department = getattr(user, "department", None)

    items = collect_recent_alarms(
        limit=min(limit, 200),
        since_ts=since_ts,
        department=department,
    )
    return ComplianceAlarmsResponse(items=items, total=len(items))


@router.post("/alarms/{alarm_id}/ack")
async def acknowledge_compliance_alarm(
    alarm_id: str,
    user=Depends(require_role_level(3)),
):
    """알람 ack — role_level ≥ 3 (실무자 이상)."""
    from backend.schemas.compliance_alarm import ComplianceAlarmAckResponse
    from features.compliance.alarm_aggregator import mark_acknowledged

    actor = getattr(user, "username", "") or getattr(user, "email", "") or "unknown"
    try:
        result = mark_acknowledged(alarm_id, actor=actor)
    except RuntimeError as e:
        raise HTTPException(status_code=500, detail=str(e))
    return ComplianceAlarmAckResponse(**result)


@router.get("/alarms/stream")
async def stream_compliance_alarms(
    interval: int = 30,
    scope: str = "all",
    user=Depends(get_current_user),
):
    """
    SSE 스트림 — D 알람 실시간 push.

    v4.2 P4 — 60s polling 대안. 클라이언트가 `EventSource` 로 구독.
    각 이벤트: `data: {items: [...], total: N}\\n\\n`.
    `interval` 초마다 aggregator 재계산 후 push (default 30s).
    `scope=mine` 시 사용자 부서 매칭만 (P5 페르소나 필터).
    """
    import asyncio
    import json as _json

    from fastapi.responses import StreamingResponse

    from features.compliance.alarm_aggregator import collect_recent_alarms

    interval = max(5, min(interval, 300))  # 5~300초 clamp
    department = getattr(user, "department", None) if scope == "mine" else None

    async def event_stream():
        # initial snapshot 즉시
        while True:
            try:
                items = collect_recent_alarms(limit=50, since_ts=0, department=department)
                payload = {
                    "items": [a.model_dump() for a in items],
                    "total": len(items),
                }
                yield f"data: {_json.dumps(payload, ensure_ascii=False)}\n\n"
            except Exception as e:
                yield f"data: {_json.dumps({'error': str(e)})}\n\n"
                break
            try:
                await asyncio.sleep(interval)
            except asyncio.CancelledError:
                break

    return StreamingResponse(
        event_stream(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
            "Connection": "keep-alive",
        },
    )


@router.post("/alarms/refresh-scenario-scores")
async def refresh_scenario_scores(user=Depends(get_current_user)):
    """
    시나리오 영향점수 캐시 전체 강제 재계산.

    v4.2 M3-4 — `scenario_impact_scores` 테이블 갱신. role_level ≥ 5 (관리자).
    일상 트래픽은 1시간 TTL 캐시로 자동 갱신되지만, 시나리오 JSON 변경 직후
    즉시 반영하려면 본 endpoint 사용. 향후 Celery beat 일일 작업으로 통합 예정.
    """
    role_level = getattr(user, "role_level", 0) or 0
    if role_level < 5:
        raise HTTPException(status_code=403, detail="role_level ≥ 5 required")

    from features.compliance.alarm_aggregator import refresh_all_scenario_scores

    try:
        result = refresh_all_scenario_scores()
    except RuntimeError as e:
        raise HTTPException(status_code=500, detail=str(e))
    return result
