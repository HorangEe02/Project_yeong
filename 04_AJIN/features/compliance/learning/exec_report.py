"""P1 D2 — 임원 보고서 자동 생성.

ChangeRecord 누적 → 1페이지 임원 보고서 (markdown 또는 docx).

지원 양식:
  markdown          — 본부장 보고용 (markdown raw text)
  docx_boon_bujang  — 본부장 보고용 (docx, 사내 표준 양식)

이사회 / 감독기관 양식 (docx_board / docx_authority) 은 사내 brand template
제공 후 P1.5 또는 P2 에서 추가.

설계:
  1. DB 에서 기간 내 변경 조회 (등급별 정렬, 최대 N건)
  2. 룰베이스 권고 액션 — Top CRITICAL 변경에서 affected_departments / penalty_extract /
     legal_class 종합해 텍스트 권고 자동 생성. LLM 미사용 — 결정론적, 빠름, 비용 0.
  3. markdown / docx 분기 — python-docx (이미 dep) 활용.
"""
from __future__ import annotations

import io
import json
import logging
import sqlite3
from datetime import date, datetime, timedelta
from typing import Any

from features.compliance.alerts.legal_guard import COMPLIANCE_AI_DISCLAIMER

logger = logging.getLogger(__name__)

# ─────────────────────────────────────────────────────────────
# DB 조회
# ─────────────────────────────────────────────────────────────


def _query_changes(since: date, until: date) -> list[dict]:
    """기간 내 변경 (filtered 제외) 등급순 정렬."""
    from features.compliance.alerts.change_detector import init_change_db, CHANGE_DB_PATH
    init_change_db()
    conn = sqlite3.connect(CHANGE_DB_PATH)
    conn.row_factory = sqlite3.Row

    rows = conn.execute(
        """SELECT id, regulation_type, change_type, item_id, item_title,
                  summary_ko, grade, status, affected_departments, affected_plants,
                  legal_class, penalty_extract, penalty_severity_krw_mn, detected_at
           FROM regulation_changes
           WHERE detected_at >= ? AND detected_at <= ?
             AND status != 'filtered'
           ORDER BY
             CASE grade
               WHEN 'CRITICAL' THEN 1
               WHEN 'HIGH' THEN 2
               WHEN 'MEDIUM' THEN 3
               WHEN 'LOW' THEN 4
               ELSE 5
             END,
             penalty_severity_krw_mn DESC,
             detected_at DESC""",
        (since.isoformat(), (until + timedelta(days=1)).isoformat()),
    ).fetchall()
    conn.close()

    out: list[dict] = []
    for r in rows:
        d = dict(r)
        for k in ("affected_departments", "affected_plants", "legal_class"):
            try:
                d[k] = json.loads(d.get(k) or "[]")
            except (json.JSONDecodeError, TypeError):
                d[k] = []
        out.append(d)
    return out


# ─────────────────────────────────────────────────────────────
# 룰베이스 요약 + 권고 액션
# ─────────────────────────────────────────────────────────────


def _aggregate_stats(changes: list[dict]) -> dict[str, Any]:
    """변경 list → 통계 집계."""
    by_grade = {"CRITICAL": 0, "HIGH": 0, "MEDIUM": 0, "LOW": 0}
    total_penalty_mn = 0
    affected_depts: dict[str, int] = {}
    affected_plants: dict[str, int] = {}
    legal_class_counts: dict[str, int] = {}

    for ch in changes:
        g = ch.get("grade", "MEDIUM")
        if g in by_grade:
            by_grade[g] += 1
        total_penalty_mn += int(ch.get("penalty_severity_krw_mn") or 0)
        for d in ch.get("affected_departments") or []:
            affected_depts[d] = affected_depts.get(d, 0) + 1
        for p in ch.get("affected_plants") or []:
            affected_plants[p] = affected_plants.get(p, 0) + 1
        for lc in ch.get("legal_class") or []:
            legal_class_counts[lc] = legal_class_counts.get(lc, 0) + 1

    return {
        "by_grade": by_grade,
        "total_penalty_mn": total_penalty_mn,
        "top_depts": sorted(affected_depts.items(), key=lambda x: -x[1])[:5],
        "top_plants": sorted(affected_plants.items(), key=lambda x: -x[1])[:5],
        "legal_class_counts": legal_class_counts,
    }


def _build_recommendations(changes: list[dict], stats: dict) -> list[str]:
    """Top CRITICAL 변경 + 통계 → 권고 액션 (룰베이스, 결정론적)."""
    recs: list[str] = []
    crits = [c for c in changes if c.get("grade") == "CRITICAL"]

    if crits:
        recs.append(
            f"CRITICAL {len(crits)}건 24시간 내 부서장 긴급 회의 소집 — "
            f"Top 1: \"{(crits[0].get('item_title') or '')[:40]}\""
        )

    if stats["legal_class_counts"].get("criminal", 0) > 0:
        recs.append(
            f"형사 처벌 가능 {stats['legal_class_counts']['criminal']}건 — "
            "법무팀 자문 의뢰 + 법규 영향 평가서 작성"
        )

    if stats["total_penalty_mn"] >= 100:  # 1억 이상 누적 벌칙
        recs.append(
            f"누적 잠재 벌금 {stats['total_penalty_mn']:,}백만원 — "
            "재무팀 회계 영향 분석 + 임원 보고서 첨부"
        )

    if stats["top_depts"]:
        top_dept = stats["top_depts"][0][0]
        recs.append(
            f"가장 영향이 큰 부서: {top_dept} "
            f"({stats['top_depts'][0][1]}건) — 부서장 직접 검토 권장"
        )

    if not recs:
        recs.append("이번 기간 즉시 대응 필요 변경 없음 — 정기 모니터링 유지")
    return recs


# ─────────────────────────────────────────────────────────────
# 양식 — markdown
# ─────────────────────────────────────────────────────────────


def _render_markdown(changes: list[dict], stats: dict, recommendations: list[str],
                     since: date, until: date) -> str:
    """본부장 보고용 markdown."""
    lines: list[str] = []
    lines.append(f"# 규제 변경 처리 현황 보고서")
    lines.append("")
    lines.append(f"**기간**: {since.isoformat()} ~ {until.isoformat()}")
    lines.append(f"**총 변경**: {len(changes)}건  ·  "
                 f"**누적 잠재 벌금**: {stats['total_penalty_mn']:,}백만원")
    lines.append("")

    if not changes:
        lines.append("> 이번 기간 감지된 규제 변경 없음.")
        return "\n".join(lines)

    # 등급 요약
    g = stats["by_grade"]
    lines.append("## 등급별 분포")
    lines.append("")
    lines.append(f"| 등급 | 건수 |")
    lines.append(f"|---|---|")
    for grade in ("CRITICAL", "HIGH", "MEDIUM", "LOW"):
        lines.append(f"| {grade} | {g[grade]} |")
    lines.append("")

    # Top 3 CRITICAL/HIGH
    lines.append("## 주요 변경 (Top 3)")
    lines.append("")
    for i, ch in enumerate(changes[:3], 1):
        lines.append(f"### {i}. [{ch['grade']}] {ch.get('item_title') or '(제목 없음)'}")
        lines.append("")
        if ch.get("summary_ko"):
            lines.append(f"- **요약**: {ch['summary_ko']}")
        if ch.get("legal_class"):
            lines.append(f"- **법적 리스크**: {', '.join(ch['legal_class'])}")
        if ch.get("penalty_extract"):
            lines.append(f"- **벌칙**: {ch['penalty_extract']}")
        if ch.get("affected_departments"):
            lines.append(f"- **영향 부서**: {', '.join(ch['affected_departments'][:5])}")
        if ch.get("affected_plants"):
            lines.append(f"- **영향 시설**: {', '.join(ch['affected_plants'][:5])}")
        lines.append("")

    # 권고 액션
    lines.append("## 권고 액션")
    lines.append("")
    for rec in recommendations:
        lines.append(f"- {rec}")
    lines.append("")

    # 디스클레이머
    lines.append("---")
    lines.append(f"*{COMPLIANCE_AI_DISCLAIMER}*")
    return "\n".join(lines)


# ─────────────────────────────────────────────────────────────
# 양식 — docx (python-docx)
# ─────────────────────────────────────────────────────────────


def _render_docx_boon_bujang(changes: list[dict], stats: dict, recommendations: list[str],
                              since: date, until: date) -> bytes:
    """본부장 보고용 docx — 표준 양식 (사내 brand 미적용 — Anthropic brand-guidelines 활용 가능)."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading("규제 변경 처리 현황 보고서", level=0)
    p = doc.add_paragraph()
    p.add_run(f"기간: {since.isoformat()} ~ {until.isoformat()}").bold = True
    doc.add_paragraph(
        f"총 변경: {len(changes)}건  ·  누적 잠재 벌금: {stats['total_penalty_mn']:,}백만원"
    )

    if not changes:
        doc.add_paragraph("이번 기간 감지된 규제 변경 없음.")
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    doc.add_heading("등급별 분포", level=1)
    g = stats["by_grade"]
    table = doc.add_table(rows=5, cols=2)
    table.style = "Light Grid Accent 1"
    table.cell(0, 0).text = "등급"
    table.cell(0, 1).text = "건수"
    for i, grade in enumerate(("CRITICAL", "HIGH", "MEDIUM", "LOW"), 1):
        table.cell(i, 0).text = grade
        table.cell(i, 1).text = str(g[grade])

    doc.add_heading("주요 변경 (Top 3)", level=1)
    for i, ch in enumerate(changes[:3], 1):
        doc.add_heading(f"{i}. [{ch['grade']}] {ch.get('item_title') or '(제목 없음)'}", level=2)
        if ch.get("summary_ko"):
            doc.add_paragraph(f"요약: {ch['summary_ko']}")
        if ch.get("legal_class"):
            doc.add_paragraph(f"법적 리스크: {', '.join(ch['legal_class'])}")
        if ch.get("penalty_extract"):
            doc.add_paragraph(f"벌칙: {ch['penalty_extract']}")
        if ch.get("affected_departments"):
            doc.add_paragraph(f"영향 부서: {', '.join(ch['affected_departments'][:5])}")
        if ch.get("affected_plants"):
            doc.add_paragraph(f"영향 시설: {', '.join(ch['affected_plants'][:5])}")

    doc.add_heading("권고 액션", level=1)
    for rec in recommendations:
        doc.add_paragraph(rec, style="List Bullet")

    foot = doc.add_paragraph()
    run = foot.add_run(COMPLIANCE_AI_DISCLAIMER)
    run.italic = True
    run.font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────
# Public API
# ─────────────────────────────────────────────────────────────

# P5 §12 — generic AJIN brand template (Pretendard, 로고 placeholder)
# 회사 brand template 받으면 _AJIN_BRAND 만 교체하면 swap 가능.
_AJIN_BRAND = {
    "company_name": "AJIN",
    "company_name_full": "주식회사 아진산업",
    "primary_color_hex": "0D47A1",   # 진한 파랑
    "accent_color_hex": "F57C00",    # 주의 주황
    "font_korean": "Pretendard",     # 한글 본문 폰트 (system fallback Malgun Gothic)
    "font_korean_fallback": "Malgun Gothic",
    "logo_placeholder": "[AJIN LOGO]",
    "footer_address": "경상남도 양산시 (회사 주소 미설정)",
    "disclaimer": COMPLIANCE_AI_DISCLAIMER,
}


def _apply_brand_to_paragraph(paragraph, *, color_hex: str | None = None) -> None:
    """문단의 모든 run 에 한글 폰트 + (선택) 색상 적용."""
    from docx.shared import RGBColor
    from docx.oxml.ns import qn
    for run in paragraph.runs:
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            from docx.oxml import OxmlElement
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:eastAsia"), _AJIN_BRAND["font_korean"])
        rFonts.set(qn("w:hAnsi"), _AJIN_BRAND["font_korean_fallback"])
        if color_hex:
            run.font.color.rgb = RGBColor.from_string(color_hex)


def _render_docx_brand_header(doc, *, title: str, audience_label: str) -> None:
    """AJIN brand 헤더 — 로고 placeholder + 제목 + 청자 라벨."""
    from docx.shared import Pt
    # 로고 placeholder
    logo_p = doc.add_paragraph()
    logo_run = logo_p.add_run(_AJIN_BRAND["logo_placeholder"])
    logo_run.bold = True
    logo_run.font.size = Pt(11)
    _apply_brand_to_paragraph(
        logo_p, color_hex=_AJIN_BRAND["primary_color_hex"],
    )
    # 제목 (heading 0)
    t_h = doc.add_heading(title, level=0)
    _apply_brand_to_paragraph(t_h, color_hex=_AJIN_BRAND["primary_color_hex"])
    # 청자 라벨
    aud_p = doc.add_paragraph()
    aud_run = aud_p.add_run(f"제출 대상: {audience_label}  ·  발행: {_AJIN_BRAND['company_name_full']}")
    aud_run.italic = True
    aud_run.font.size = Pt(10)
    _apply_brand_to_paragraph(aud_p)


def _render_docx_brand_footer(doc) -> None:
    """공통 footer — 주소 + disclaimer."""
    from docx.shared import Pt
    doc.add_paragraph()  # spacing
    f1 = doc.add_paragraph()
    r1 = f1.add_run(_AJIN_BRAND["footer_address"])
    r1.font.size = Pt(8)
    _apply_brand_to_paragraph(f1)
    f2 = doc.add_paragraph()
    r2 = f2.add_run(_AJIN_BRAND["disclaimer"])
    r2.italic = True
    r2.font.size = Pt(8)
    _apply_brand_to_paragraph(f2, color_hex=_AJIN_BRAND["accent_color_hex"])


def _render_docx_brand_body(
    doc, changes: list[dict], stats: dict,
    recommendations: list[str], since: date, until: date,
) -> None:
    """공통 본문 — _render_docx_boon_bujang 의 본문을 brand 적용 버전으로 재사용."""
    from docx.shared import Pt
    period_p = doc.add_paragraph()
    pr = period_p.add_run(f"기간: {since.isoformat()} ~ {until.isoformat()}")
    pr.bold = True
    _apply_brand_to_paragraph(period_p)

    summary_p = doc.add_paragraph(
        f"총 변경: {len(changes)}건  ·  누적 잠재 벌금: {stats['total_penalty_mn']:,}백만원"
    )
    _apply_brand_to_paragraph(summary_p)

    if not changes:
        empty_p = doc.add_paragraph("이번 기간 감지된 규제 변경 없음.")
        _apply_brand_to_paragraph(empty_p)
        return

    grade_h = doc.add_heading("등급별 분포", level=1)
    _apply_brand_to_paragraph(grade_h, color_hex=_AJIN_BRAND["primary_color_hex"])
    g = stats["by_grade"]
    table = doc.add_table(rows=5, cols=2)
    table.style = "Light Grid Accent 1"
    table.cell(0, 0).text = "등급"
    table.cell(0, 1).text = "건수"
    for i, grade in enumerate(("CRITICAL", "HIGH", "MEDIUM", "LOW"), 1):
        table.cell(i, 0).text = grade
        table.cell(i, 1).text = str(g[grade])
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                _apply_brand_to_paragraph(p)

    top_h = doc.add_heading("주요 변경 (Top 3)", level=1)
    _apply_brand_to_paragraph(top_h, color_hex=_AJIN_BRAND["primary_color_hex"])
    for i, ch in enumerate(changes[:3], 1):
        h2 = doc.add_heading(
            f"{i}. [{ch['grade']}] {ch.get('item_title') or '(제목 없음)'}",
            level=2,
        )
        _apply_brand_to_paragraph(h2)
        for label, key in (
            ("요약", "summary_ko"),
            ("법적 리스크", None),  # legal_class join 별도
            ("벌칙", "penalty_extract"),
            ("영향 부서", None),
            ("영향 시설", None),
        ):
            if label == "법적 리스크" and ch.get("legal_class"):
                p = doc.add_paragraph(f"법적 리스크: {', '.join(ch['legal_class'])}")
                _apply_brand_to_paragraph(p)
            elif label == "영향 부서" and ch.get("affected_departments"):
                p = doc.add_paragraph(
                    f"영향 부서: {', '.join(ch['affected_departments'][:5])}"
                )
                _apply_brand_to_paragraph(p)
            elif label == "영향 시설" and ch.get("affected_plants"):
                p = doc.add_paragraph(
                    f"영향 시설: {', '.join(ch['affected_plants'][:5])}"
                )
                _apply_brand_to_paragraph(p)
            elif key and ch.get(key):
                p = doc.add_paragraph(f"{label}: {ch[key]}")
                _apply_brand_to_paragraph(p)

    rec_h = doc.add_heading("권고 액션", level=1)
    _apply_brand_to_paragraph(rec_h, color_hex=_AJIN_BRAND["primary_color_hex"])
    for rec in recommendations:
        p = doc.add_paragraph(rec, style="List Bullet")
        _apply_brand_to_paragraph(p)


def _render_docx_board(
    changes: list[dict], stats: dict, recommendations: list[str],
    since: date, until: date,
) -> bytes:
    """이사회 보고용 docx — generic AJIN brand template (P5 §12)."""
    from docx import Document
    doc = Document()
    _render_docx_brand_header(
        doc,
        title="규제 변경 처리 현황 보고서",
        audience_label="이사회",
    )
    _render_docx_brand_body(doc, changes, stats, recommendations, since, until)
    _render_docx_brand_footer(doc)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _render_docx_authority(
    changes: list[dict], stats: dict, recommendations: list[str],
    since: date, until: date,
) -> bytes:
    """감독기관 제출용 docx — 동일 brand template, 청자 라벨만 차별 (P5 §12)."""
    from docx import Document
    doc = Document()
    _render_docx_brand_header(
        doc,
        title="규제 준수 현황 통보",
        audience_label="감독 기관",
    )
    _render_docx_brand_body(doc, changes, stats, recommendations, since, until)
    _render_docx_brand_footer(doc)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


VALID_FORMATS = {"markdown", "docx_boon_bujang", "docx_board", "docx_authority"}


def generate_report(format: str, since: date | None = None,
                    until: date | None = None) -> tuple[bytes, str, str]:
    """기간 내 변경 보고서 생성.

    Args:
        format: "markdown" | "docx_boon_bujang"
        since: 기간 시작일 (default = 이번 달 1일)
        until: 기간 종료일 (default = 오늘)

    Returns:
        (content_bytes, content_type, suggested_filename)
    """
    if format not in VALID_FORMATS:
        raise ValueError(f"unsupported format: {format}. Use one of {VALID_FORMATS}")

    today = date.today()
    if since is None:
        since = today.replace(day=1)
    if until is None:
        until = today

    changes = _query_changes(since, until)
    stats = _aggregate_stats(changes)
    recommendations = _build_recommendations(changes, stats)

    date_tag = f"{since.strftime('%Y%m%d')}-{until.strftime('%Y%m%d')}"

    if format == "markdown":
        body = _render_markdown(changes, stats, recommendations, since, until)
        return (body.encode("utf-8"),
                "text/markdown; charset=utf-8",
                f"compliance-report-{date_tag}.md")

    docx_renderers = {
        "docx_boon_bujang": (_render_docx_boon_bujang, "boon_bujang"),
        "docx_board": (_render_docx_board, "board"),               # P5 §12
        "docx_authority": (_render_docx_authority, "authority"),    # P5 §12
    }
    if format in docx_renderers:
        renderer, suffix = docx_renderers[format]
        body_bytes = renderer(changes, stats, recommendations, since, until)
        return (
            body_bytes,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            f"compliance-report-{suffix}-{date_tag}.docx",
        )

    raise ValueError(f"unsupported format: {format}. Use one of {VALID_FORMATS}")
